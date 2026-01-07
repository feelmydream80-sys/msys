from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.section import WD_ORIENT

def create_installation_manual():
    """인프라 관계자용 설치 매뉴얼 생성"""
    doc = Document()

    # 스타일 설정
    title_style = doc.styles.add_style('InstallTitle', WD_STYLE_TYPE.PARAGRAPH)
    title_style.font.size = Pt(28)
    title_style.font.color.rgb = RGBColor(0, 0, 0)
    title_style.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    heading1_style = doc.styles.add_style('InstallH1', WD_STYLE_TYPE.PARAGRAPH)
    heading1_style.font.size = Pt(18)
    heading1_style.font.bold = True

    heading2_style = doc.styles.add_style('InstallH2', WD_STYLE_TYPE.PARAGRAPH)
    heading2_style.font.size = Pt(14)
    heading2_style.font.bold = True

    heading3_style = doc.styles.add_style('InstallH3', WD_STYLE_TYPE.PARAGRAPH)
    heading3_style.font.size = Pt(12)
    heading3_style.font.bold = True

    # 표지
    title = doc.add_paragraph('MSYS 설치 매뉴얼', style='InstallTitle')
    subtitle = doc.add_paragraph('인프라 구축 및 시스템 설치 가이드', style='InstallTitle')
    version = doc.add_paragraph('버전 1.14.2', style='InstallTitle')
    version.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    doc.add_page_break()

    # 목차
    toc_title = doc.add_paragraph('목차', style='InstallH1')
    toc = doc.add_paragraph()
    toc.add_run('1. 시스템 요구사항 ........................ 3\n')
    toc.add_run('2. 네트워크 구성 ........................ 4\n')
    toc.add_run('3. 서버 환경 구축 ....................... 5\n')
    toc.add_run('4. 데이터베이스 설치 .................... 8\n')
    toc.add_run('5. 애플리케이션 배포 .................... 11\n')
    toc.add_run('6. 보안 설정 ........................... 14\n')
    toc.add_run('7. 모니터링 설정 ........................ 16\n')
    toc.add_run('부록 A. 설치 검증 체크리스트 ............. 18\n')

    doc.add_page_break()

    # 1. 시스템 요구사항
    doc.add_paragraph('1. 시스템 요구사항', style='InstallH1')
    reqs = doc.add_paragraph()
    reqs.add_run('1.1 하드웨어 요구사항\n').bold = True
    reqs.add_run('• CPU: Intel Xeon 또는 AMD EPYC, 2.0GHz 이상\n')
    reqs.add_run('• 메모리: 8GB 이상 (16GB 권장)\n')
    reqs.add_run('• 저장소: SSD 50GB 이상\n')
    reqs.add_run('• 네트워크: 1Gbps 이더넷\n\n')

    reqs.add_run('1.2 소프트웨어 요구사항\n').bold = True
    reqs.add_run('• 운영체제: Ubuntu Server 20.04 LTS 이상\n')
    reqs.add_run('• Python: 3.8 이상\n')
    reqs.add_run('• 데이터베이스: PostgreSQL 13 이상\n')
    reqs.add_run('• 웹 서버: Nginx (선택사항)\n\n')

    # 2. 네트워크 구성
    doc.add_paragraph('2. 네트워크 구성', style='InstallH1')
    network = doc.add_paragraph()
    network.add_run('2.1 방화벽 설정\n').bold = True
    network.add_run('필요한 포트 개방:\n')
    network.add_run('• 80 (HTTP)\n')
    network.add_run('• 443 (HTTPS)\n')
    network.add_run('• 5432 (PostgreSQL, 내부 통신만)\n')
    network.add_run('• 22 (SSH)\n\n')

    network.add_run('2.2 DNS 설정\n').bold = True
    network.add_run('도메인 및 SSL 인증서 구성\n\n')

    # 3. 서버 환경 구축
    doc.add_paragraph('3. 서버 환경 구축', style='InstallH1')

    # 3.1 Ubuntu Server 설치
    doc.add_paragraph('3.1 Ubuntu Server 설치', style='InstallH2')
    ubuntu = doc.add_paragraph()
    ubuntu.add_run('1. Ubuntu Server ISO 다운로드\n')
    ubuntu.add_run('2. 부팅 USB 생성\n')
    ubuntu.add_run('3. 서버에 설치 (기본 설정 사용)\n')
    ubuntu.add_run('4. 패키지 업데이트:\n\n')
    ubuntu.add_run('sudo apt update && sudo apt upgrade -y\n\n').font.name = 'Courier New'

    # 3.2 필수 패키지 설치
    doc.add_paragraph('3.2 필수 패키지 설치', style='InstallH2')
    packages = doc.add_paragraph()
    packages.add_run('sudo apt install -y python3 python3-pip python3-venv postgresql postgresql-contrib nginx curl wget\n\n').font.name = 'Courier New'

    # 4. 데이터베이스 설치
    doc.add_paragraph('4. 데이터베이스 설치', style='InstallH1')

    # 4.1 PostgreSQL 설정
    doc.add_paragraph('4.1 PostgreSQL 설정', style='InstallH2')
    pg = doc.add_paragraph()
    pg.add_run('1. PostgreSQL 서비스 시작:\n\n')
    pg.add_run('sudo systemctl start postgresql\n')
    pg.add_run('sudo systemctl enable postgresql\n\n').font.name = 'Courier New'

    pg.add_run('2. 데이터베이스 사용자 생성:\n\n')
    pg.add_run('sudo -u postgres psql\n').font.name = 'Courier New'
    pg.add_run('CREATE USER msys_user WITH PASSWORD \'strong_password\';\n').font.name = 'Courier New'
    pg.add_run('CREATE DATABASE msys_db OWNER msys_user;\n').font.name = 'Courier New'
    pg.add_run('\\q\n\n').font.name = 'Courier New'

    # 5. 애플리케이션 배포
    doc.add_paragraph('5. 애플리케이션 배포', style='InstallH1')

    # 5.1 애플리케이션 사용자 생성
    doc.add_paragraph('5.1 애플리케이션 사용자 생성', style='InstallH2')
    user = doc.add_paragraph()
    user.add_run('sudo useradd -m -s /bin/bash msys\n')
    user.add_run('sudo usermod -aG sudo msys\n\n').font.name = 'Courier New'

    # 5.2 소스 코드 배포
    doc.add_paragraph('5.2 소스 코드 배포', style='InstallH2')
    deploy = doc.add_paragraph()
    deploy.add_run('sudo -u msys mkdir -p /home/msys/app\n')
    deploy.add_run('sudo -u msys git clone <repository> /home/msys/app\n\n').font.name = 'Courier New'

    # 5.3 Python 환경 설정
    doc.add_paragraph('5.3 Python 환경 설정', style='InstallH2')
    python_env = doc.add_paragraph()
    python_env.add_run('cd /home/msys/app\n')
    python_env.add_run('python3 -m venv venv\n')
    python_env.add_run('source venv/bin/activate\n')
    python_env.add_run('pip install -r requirements.txt\n\n').font.name = 'Courier New'

    # 6. 보안 설정
    doc.add_paragraph('6. 보안 설정', style='InstallH1')
    security = doc.add_paragraph()
    security.add_run('6.1 방화벽 설정\n').bold = True
    security.add_run('sudo ufw allow 80\n')
    security.add_run('sudo ufw allow 443\n')
    security.add_run('sudo ufw --force enable\n\n').font.name = 'Courier New'

    security.add_run('6.2 SSL/TLS 설정\n').bold = True
    security.add_run('Let\'s Encrypt를 사용한 무료 SSL 인증서 설정\n\n')

    # 7. 모니터링 설정
    doc.add_paragraph('7. 모니터링 설정', style='InstallH1')
    monitoring = doc.add_paragraph()
    monitoring.add_run('7.1 로그 로테이션\n').bold = True
    monitoring.add_run('logrotate 설정으로 로그 파일 관리\n\n')

    monitoring.add_run('7.2 시스템 모니터링\n').bold = True
    monitoring.add_run('htop, iotop 등 모니터링 도구 설치\n\n')

    doc.save('MSYS_Installation_Manual.docx')
    print('설치 매뉴얼 생성 완료: MSYS_Installation_Manual.docx')

def create_function_manual():
    """개발자용 기능 매뉴얼 생성 - DB→DAO→SERVICE→JS→HTML 구조"""
    doc = Document()

    # 페이지 방향을 가로로 설정
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE

    # 스타일 설정
    title_style = doc.styles.add_style('FuncTitle', WD_STYLE_TYPE.PARAGRAPH)
    title_style.font.size = Pt(28)
    title_style.font.color.rgb = RGBColor(0, 0, 0)
    title_style.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    heading1_style = doc.styles.add_style('FuncH1', WD_STYLE_TYPE.PARAGRAPH)
    heading1_style.font.size = Pt(18)
    heading1_style.font.bold = True

    heading2_style = doc.styles.add_style('FuncH2', WD_STYLE_TYPE.PARAGRAPH)
    heading2_style.font.size = Pt(14)
    heading2_style.font.bold = True

    heading3_style = doc.styles.add_style('FuncH3', WD_STYLE_TYPE.PARAGRAPH)
    heading3_style.font.size = Pt(12)
    heading3_style.font.bold = True

    heading4_style = doc.styles.add_style('FuncH4', WD_STYLE_TYPE.PARAGRAPH)
    heading4_style.font.size = Pt(11)
    heading4_style.font.bold = True

    code_style = doc.styles.add_style('FuncCode', WD_STYLE_TYPE.PARAGRAPH)
    code_style.font.name = 'Courier New'
    code_style.font.size = Pt(10)

    # 표지
    title = doc.add_paragraph('MSYS 기능 매뉴얼', style='FuncTitle')
    subtitle = doc.add_paragraph('개발자용 데이터 흐름 참조 가이드', style='FuncTitle')
    version = doc.add_paragraph('버전 1.14.2', style='FuncTitle')
    version.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    doc.add_page_break()

    # 목차
    toc_title = doc.add_paragraph('목차', style='FuncH1')
    toc = doc.add_paragraph()
    toc.add_run('1. 시스템 아키텍처 ........................ 3\n')
    toc.add_run('2. 인터페이스별 데이터 흐름 ................ 5\n')
    toc.add_run('   2.1 대시보드 인터페이스 .................. 5\n')
    toc.add_run('   2.2 데이터 분석 인터페이스 ................ 15\n')
    toc.add_run('   2.3 차트 분석 인터페이스 .................. 25\n')
    toc.add_run('   2.4 잔디 현황 인터페이스 .................. 35\n')
    toc.add_run('   2.5 매핑 관리 인터페이스 .................. 45\n')
    toc.add_run('   2.6 데이터 명세서 인터페이스 ................ 55\n')
    toc.add_run('3. 공통 컴포넌트 .......................... 65\n')
    toc.add_run('4. 설정 및 배포 .......................... 70\n')
    toc.add_run('부록 A. 데이터 흐름 다이어그램 ................ 75\n')

    doc.add_page_break()

    # 1. 시스템 아키텍처
    doc.add_paragraph('1. 시스템 아키텍처', style='FuncH1')
    arch = doc.add_paragraph()
    arch.add_run('1.1 데이터 흐름 개요\n').bold = True
    arch.add_run('MSYS는 다음과 같은 데이터 흐름을 따르는 계층형 아키텍처를 사용합니다:\n\n')
    arch.add_run('HTML → JavaScript → Service → DAO → Database\n\n')

    arch.add_run('1.2 각 계층의 역할\n').bold = True
    arch.add_run('• HTML: 사용자 인터페이스 및 데이터 표시\n')
    arch.add_run('• JavaScript: 사용자 상호작용 및 API 호출\n')
    arch.add_run('• Service: 비즈니스 로직 및 데이터 가공\n')
    arch.add_run('• DAO: 데이터베이스 쿼리 실행\n')
    arch.add_run('• Database: 데이터 저장 및 검색\n\n')

    # 2. 인터페이스별 데이터 흐름
    doc.add_paragraph('2. 인터페이스별 데이터 흐름', style='FuncH1')

    # 2.1 대시보드 인터페이스
    doc.add_paragraph('2.1 대시보드 인터페이스', style='FuncH2')
    dash_intro = doc.add_paragraph()
    dash_intro.add_run('대시보드는 시스템의 전체 현황을 실시간으로 모니터링하는 메인 인터페이스입니다.\n\n')

    # 대시보드 스크린샷 추가
    try:
        doc.add_picture('scrennshot/dashboard_screenshot.PNG', width=Inches(6))
        doc.add_paragraph('그림 2.1: 대시보드 인터페이스 화면', style='FuncH3')
    except:
        doc.add_paragraph('[스크린샷: dashboard_screenshot.PNG를 찾을 수 없습니다]', style='FuncH3')

    doc.add_paragraph('', style='FuncH3')  # 빈 줄

    # 대시보드 인터페이스 표
    dash_table = doc.add_table(rows=1, cols=5)
    dash_table.style = 'Table Grid'
    hdr_cells = dash_table.rows[0].cells
    hdr_cells[0].text = '레벨'
    hdr_cells[1].text = '파일/클래스'
    hdr_cells[2].text = '주요 메소드/함수'
    hdr_cells[3].text = '복잡도'
    hdr_cells[4].text = '설명'

    # DB 레벨 행
    db_row = dash_table.add_row().cells
    db_row[0].text = 'DB'
    db_row[1].text = 'DDL/*.sql\nsql/dashboard/*'
    db_row[2].text = '-'
    db_row[3].text = '고: 트리거\n서브쿼리 포함'
    db_row[4].text = '복합키, 타임존\nautomatic 로그'

    # DAO 레벨 메소드별 행들 (첫 번째 행에 파일 정보 표시)
    dao_methods = [
        ('get_summary(start_date, end_date, all_data, job_ids)', '중', '대시보드 요약 데이터 조회 - 지정된 기간과 Job에 대한 수집 통계 계산\n입력: start_date="2025-12-01", end_date="2025-12-31", all_data=False, job_ids=["CD101", "CD102"]\n출력: [{"job_id": "CD101", "cd_nm": "기상청 예보 데이터", "total_count": 45, "overall_success_count": 40, "overall_fail_count": 3, "overall_no_data_count": 2, "week_success": 12, "month_success": 35, "day_success": 2, "day_fail_count": 0, "day_ing_count": 1, "day_no_data_count": 0, "day_total_scheduled": 3, "fail_streak": 0, "settings": {...}}]'),
        ('get_raw_data(start_date, end_date, job_ids, all_data)', '중', '원본 수집 데이터 조회 - 필터링된 수집 이력 데이터 반환\n입력: start_date="2025-12-01", end_date="2025-12-31", job_ids=["CD101"], all_data=False\n출력: [{"job_id": "CD101", "con_id": "CON001", "status": "CD901", "start_dt": "2025-12-18T14:30:15+09", "end_dt": "2025-12-18T14:35:22+09", "rqs_info": "총 요청 수: 10, 실패: 0"}]'),
        ('get_analytics_success_rate_trend(start_date, end_date, job_ids, all_data)', '고', '성공률 추이 데이터 조회 - 시간별 성공률 변화 추이 계산\n입력: start_date="2025-12-01", end_date="2025-12-31", job_ids=["CD101"]\n출력: [{"date": "2025-12-18", "job_id": "CD101", "success_rate": 95.2, "total_count": 20, "success_count": 19}]'),
        ('get_trouble_by_code(start_date, end_date, job_ids, all_data)', '중', '장애 코드별 문제 데이터 조회 - 오류 유형별 발생 현황 분석\n입력: start_date="2025-12-01", end_date="2025-12-31", job_ids=["CD101"]\n출력: [{"status": "CD902", "count": 3, "job_id": "CD101", "error_message": "Connection timeout"}]'),
        ('get_event_log(start_date, end_date, all_data, allowed_job_ids)', '중', '이벤트 로그 조회 - 시스템 이벤트 및 변경 이력 확인\n입력: start_date="2025-12-01", end_date="2025-12-31", all_data=False, allowed_job_ids=["CD101"]\n출력: [{"con_id": "CON001", "job_id": "CD101", "status": "CD901", "start_dt": "2025-12-18T14:30:15+09", "end_dt": "2025-12-18T14:35:22+09", "rqs_info": "총 요청 수: 10, 실패: 0"}]'),
        ('get_daily_job_counts(job_id, start_date, end_date, all_data, job_ids)', '중', '일별 Job 카운트 조회 - Job별 일간 실행 통계 집계\n입력: job_id="CD101", start_date="2025-12-01", end_date="2025-12-31", all_data=False, job_ids=["CD101"]\n출력: [{"date": "2025-12-18", "job_id": "CD101", "success_count": 2, "fail_count": 0, "total_count": 2}]'),
        ('save_event(con_id, job_id, status, rqs_info)', '낮음', '이벤트 로그 저장 - 수집 이벤트 정보 기록\n입력: con_id="CON001", job_id="CD101", status="CD901", rqs_info="총 요청 수: 10, 실패: 0"\n출력: None (DB 저장 완료)'),
        ('get_distinct_job_ids(job_ids)', '낮음', '고유 Job ID 목록 조회 - 중복 제거된 Job ID 리스트 반환\n입력: job_ids=["CD101", "CD102"]\n출력: ["CD101", "CD102"]')
    ]

    # 첫 번째 메소드 행 (파일 정보 포함)
    first_dao_row = dash_table.add_row().cells
    first_dao_row[0].text = 'DAO'
    first_dao_row[1].text = 'mapper/dashboard_mapper.py\nsql/dashboard/*'
    first_dao_row[2].text = dao_methods[0][0]
    first_dao_row[3].text = dao_methods[0][1]
    first_dao_row[4].text = dao_methods[0][2]

    # 나머지 메소드 행들 (파일 정보 생략)
    for method, complexity, description in dao_methods[1:]:
        dao_row = dash_table.add_row().cells
        dao_row[0].text = ''  # 레벨 생략
        dao_row[1].text = ''  # 파일 정보 생략
        dao_row[2].text = method
        dao_row[3].text = complexity
        dao_row[4].text = description

    # 이벤트 로그 상세 분석 행 추가
    event_log_row = dash_table.add_row().cells
    event_log_row[0].text = 'DAO'
    event_log_row[1].text = 'mapper/dashboard_mapper.py\nsql/dashboard/get_event_log.sql'
    event_log_row[2].text = 'get_event_log(start_date, end_date, all_data, allowed_job_ids)'
    event_log_row[3].text = '중'
    event_log_row[4].text = '이벤트 로그 조회 - TB_CON_HIST_EVNT_LOG에서 JSON 데이터 추출 및 변환'

    # 이벤트 로그 JavaScript 렌더링 행 추가
    event_log_js_row = dash_table.add_row().cells
    event_log_js_row[0].text = 'JS'
    event_log_js_row[1].text = 'static/js/modules/dashboard/eventLog.js'
    event_log_js_row[2].text = 'renderEventLogToasts()\nloadEventLogPage()\nfilterEventLogData()'
    event_log_js_row[3].text = '중'
    event_log_js_row[4].text = '이벤트 로그 렌더링 - 상태별 아이콘, 시간 포맷팅, 검색 필터링'

    # SERVICE 레벨 메소드별 행들 (첫 번째 행에 파일 정보 표시)
    svc_methods = [
        ('get_summary(start_date, end_date, user)', '고', '대시보드 데이터 처리 메인 메소드 - 6단계 비즈니스 로직 실행'),
        ('get_raw_data(start_date, end_date, job_ids, all_data)', '중', '원본 수집 데이터 조회 및 기본 가공'),
        ('get_analytics_success_rate_trend(start_date, end_date, job_ids, all_data)', '고', '성공률 추이 데이터 계산 및 반환'),
        ('get_trouble_by_code(start_date, end_date, job_ids, all_data)', '중', '장애 코드별 문제 현황 분석'),
        ('get_event_log(start_date, end_date, all_data, allowed_job_ids)', '중', '시스템 이벤트 로그 조회'),
        ('get_daily_job_counts(job_id, start_date, end_date, all_data, job_ids)', '중', '일별 Job 실행 통계 계산'),
        ('_calculate_fail_streak(job_id)', '중', '특정 Job의 연속 실패 횟수 계산'),
        ('_get_allowed_job_ids(user)', '중', '사용자 권한에 따른 허용 Job ID 필터링'),
        ('_fetch_manager_settings_with_icons()', '중', '관리자 설정과 아이콘 정보 조회'),
        ('_combine_historical_and_today_data()', '고', '과거 데이터와 오늘 데이터 결합'),
        ('_apply_settings_and_filters()', '중', '설정 적용 및 추가 필터링'),
        ('_log_final_data_counts(processed_data)', '낮음', '최종 처리 데이터 로깅')
    ]

    # 첫 번째 메소드 행 (파일 정보 포함)
    first_svc_row = dash_table.add_row().cells
    first_svc_row[0].text = 'SERVICE'
    first_svc_row[1].text = 'service/dashboard_service.py'
    first_svc_row[2].text = svc_methods[0][0]
    first_svc_row[3].text = svc_methods[0][1]
    first_svc_row[4].text = svc_methods[0][2]

    # 나머지 메소드 행들 (파일 정보 생략)
    for method, complexity, description in svc_methods[1:]:
        svc_row = dash_table.add_row().cells
        svc_row[0].text = ''  # 레벨 생략
        svc_row[1].text = ''  # 파일 정보 생략
        svc_row[2].text = method
        svc_row[3].text = complexity
        svc_row[4].text = description

    # JS 레벨 행
    js_row = dash_table.add_row().cells
    js_row[0].text = 'JS'
    js_row[1].text = 'static/js/modules/dashboard/*'
    js_row[2].text = 'initializeDashboardData\nloadDashboardData\nrenderSummaryCards\nrenderStatusTable\nhandleDateFilterChange'
    js_row[3].text = '중: Promise.all\nAPI 호출\nDOM 조작'
    js_row[4].text = '병렬 로딩\n상태 추적\n이벤트 핸들링'

    # HTML 레벨 행
    html_row = dash_table.add_row().cells
    html_row[0].text = 'HTML'
    html_row[1].text = 'templates/dashboard.html'
    html_row[2].text = '-'
    html_row[3].text = '낮음'
    html_row[4].text = '템플릿 변수\nJinja2 필터'

    doc.add_paragraph('', style='FuncH3')  # 빈 줄

    # 실제 데이터 예시
    doc.add_paragraph('실제 데이터 예시:', style='FuncH3')
    example = doc.add_paragraph()
    example.add_run('• DB: job_id=\'CD101\', status=\'CD901\', start_dt=\'2025-12-18 14:30:15+09\'\n')
    example.add_run('• DAO: params=[\'2025-12-01\', \'2025-12-31\'], job_ids=[\'CD101\', \'CD102\']\n')
    example.add_run('• SERVICE: allowed_job_ids=[\'CD101\'], fail_streak=2, status=\'warning\'\n')
    example.add_run('• JS: dataFlowStatus.apiCallSuccess=true, loadedRecords=45\n')
    example.add_run('• HTML: {{ user.is_admin }} = true, {{ current_date }} = \'2025-12-18\'\n\n')

    # 이벤트 로그 상세 분석 및 실제 데이터 예시
    doc.add_paragraph('이벤트 로그 데이터 상세 분석:', style='FuncH3')
    event_log_analysis = doc.add_paragraph()
    event_log_analysis.add_run('이벤트 로그는 TB_CON_HIST_EVNT_LOG 테이블에서 JSON 형태의 데이터를 추출하여 표시합니다.\n\n')

    event_log_analysis.add_run('SQL 쿼리 분석 (sql/dashboard/get_event_log.sql):\n').bold = True
    event_log_analysis.add_run('SELECT\n')
    event_log_analysis.add_run('    (EVNT_CHG_ROW ->> \'con_id\')::text AS con_id,\n')
    event_log_analysis.add_run('    EVNT_OCCR_TIME AS start_dt,\n')
    event_log_analysis.add_run('    (EVNT_CHG_ROW ->> \'end_dt\')::timestamptz AS end_dt,\n')
    event_log_analysis.add_run('    (EVNT_CHG_ROW ->> \'job_id\')::text AS job_id,\n')
    event_log_analysis.add_run('    (EVNT_CHG_ROW ->> \'rqs_info\')::text AS rqs_info,\n')
    event_log_analysis.add_run('    (EVNT_CHG_ROW ->> \'status\')::text AS status\n')
    event_log_analysis.add_run('FROM TB_CON_HIST_EVNT_LOG\n')
    event_log_analysis.add_run('{where_clause}\n')
    event_log_analysis.add_run('ORDER BY EVNT_OCCR_TIME DESC\n\n')

    event_log_analysis.add_run('JavaScript 렌더링 로직 (static/js/modules/dashboard/eventLog.js):\n').bold = True
    event_log_analysis.add_run('• renderEventLogToasts(): 이벤트 로그를 그리드 형태로 표시\n')
    event_log_analysis.add_run('• formatDurationHr(): 수집 시간 계산 (시작~종료 시간)\n')
    event_log_analysis.add_run('• filterEventLogData(): 검색어 기반 필터링\n')
    event_log_analysis.add_run('• 상태별 아이콘 매핑: CD901(🟢), CD902(🔴), CD903(🟠), CD904(🔵)\n\n')

    event_log_analysis.add_run('실제 입력/출력 예시:\n').bold = True
    event_log_analysis.add_run('• 함수 입력: get_event_log(start_date=\'2025-12-01\', end_date=\'2025-12-31\', all_data=False, allowed_job_ids=[\'CD101\', \'CD102\'])\n')
    event_log_analysis.add_run('• SQL 파라미터: params = [\'2025-12-01\', \'2025-12-31\', [\'CD101\', \'CD102\']]\n')
    event_log_analysis.add_run('• SQL 결과: [{\'con_id\': \'CON001\', \'start_dt\': \'2025-12-18T14:30:15+09\', \'job_id\': \'CD101\', \'status\': \'CD901\', \'rqs_info\': \'총 요청 수: 10, 실패: 0\'}, ...]\n')
    event_log_analysis.add_run('• JavaScript 렌더링: 시간(2025-12-18 14:30:15), 아이콘(🟢), Job ID(CD101), 상태(정상 수집), 수집 건수(10/10), 성공률(100%), 수집 시간(수집시간: 0.5hr)\n\n')

    # 2.2 데이터 분석 인터페이스
    doc.add_paragraph('2.2 데이터 분석 인터페이스', style='FuncH2')
    analysis_intro = doc.add_paragraph()
    analysis_intro.add_run('데이터 수집 결과를 심층 분석하고 필터링하여 조회하는 인터페이스입니다.\n\n')

    # 데이터 분석 인터페이스 스크린샷 추가
    try:
        doc.add_picture('scrennshot/data_analysis_screenshot.PNG', width=Inches(6))
        doc.add_paragraph('그림 2.2: 데이터 분석 인터페이스 화면', style='FuncH3')
    except:
        doc.add_paragraph('[스크린샷: data_analysis_screenshot.PNG를 찾을 수 없습니다]', style='FuncH3')

    doc.add_paragraph('', style='FuncH3')  # 빈 줄

    # 데이터 분석 인터페이스 표
    analysis_table = doc.add_table(rows=1, cols=5)
    analysis_table.style = 'Table Grid'
    hdr_cells = analysis_table.rows[0].cells
    hdr_cells[0].text = '레벨'
    hdr_cells[1].text = '파일/클래스'
    hdr_cells[2].text = '주요 메소드/함수'
    hdr_cells[3].text = '복잡도'
    hdr_cells[4].text = '설명'

    # DB 레벨 행
    analysis_db_row = analysis_table.add_row().cells
    analysis_db_row[0].text = 'DB'
    analysis_db_row[1].text = 'DDL/*.sql\nsql/analytics/*'
    analysis_db_row[2].text = '-'
    analysis_db_row[3].text = '고: 복합 조인\n다중 집계'
    analysis_db_row[4].text = '시간대 변환\n통계 계산'

    # DAO 레벨 메소드별 행들 (첫 번째 행에 파일 정보 표시)
    analysis_dao_methods = [
        ('get_user_access_stats(start_date, end_date, menu_id)', '중', '사용자별 메뉴 접근 통계 조회'),
        ('get_menu_access_stats(start_date, end_date, menu_id)', '중', '메뉴 접근 통계 조회'),
        ('get_menu_access_stats_weekly(start_date, end_date, menu_id)', '중', '주별 메뉴 접근 통계 조회'),
        ('get_yearly_total_stats(year, menu_id)', '중', '연간 총 통계 조회'),
        ('get_most_recent_data_date()', '낮음', '가장 최근 데이터 날짜 조회'),
        ('get_available_years_months()', '중', '사용 가능한 연월 목록 조회'),
        ('get_distinct_menu_names()', '낮음', '고유 메뉴 이름 목록 조회'),
        ('get_total_unique_users_by_week(start_date, end_date)', '중', '주별 고유 사용자 수 조회'),
        ('get_menu_access_stats_monthly(start_date, end_date, menu_id)', '중', '월별 메뉴 접근 통계 조회'),
        ('insert_user_access_log(user_id, menu_id, access_time)', '낮음', '사용자 접근 로그 기록'),
        ('get_menu_name_by_menu_id(menu_id)', '낮음', '메뉴 ID로 메뉴 이름 조회')
    ]

    # 첫 번째 메소드 행 (파일 정보 포함)
    first_analysis_dao_row = analysis_table.add_row().cells
    first_analysis_dao_row[0].text = 'DAO'
    first_analysis_dao_row[1].text = 'dao/analytics_dao.py'
    first_analysis_dao_row[2].text = analysis_dao_methods[0][0]
    first_analysis_dao_row[3].text = analysis_dao_methods[0][1]
    first_analysis_dao_row[4].text = analysis_dao_methods[0][2]

    # 나머지 메소드 행들 (파일 정보 생략)
    for method, complexity, description in analysis_dao_methods[1:]:
        analysis_dao_row = analysis_table.add_row().cells
        analysis_dao_row[0].text = ''  # 레벨 생략
        analysis_dao_row[1].text = ''  # 파일 정보 생략
        analysis_dao_row[2].text = method
        analysis_dao_row[3].text = complexity
        analysis_dao_row[4].text = description

    # SERVICE 레벨 메소드별 행들 (첫 번째 행에 파일 정보 표시)
    analysis_svc_methods = [
        ('get_analytics_data(start_date, end_date, filters)', '고', '분석 데이터 조회 및 가공'),
        ('process_chart_data(raw_data, chart_type)', '중', '차트 데이터 처리 및 변환'),
        ('calculate_statistics(data, stat_type)', '고', '통계 계산 수행'),
        ('generate_reports(data, report_format)', '고', '보고서 생성'),
        ('validate_date_range(start_date, end_date)', '낮음', '날짜 범위 유효성 검증'),
        ('apply_user_permissions(data, user)', '중', '사용자 권한 적용'),
        ('format_analytics_response(data)', '중', '분석 결과 포맷팅')
    ]

    # 첫 번째 메소드 행 (파일 정보 포함)
    first_analysis_svc_row = analysis_table.add_row().cells
    first_analysis_svc_row[0].text = 'SERVICE'
    first_analysis_svc_row[1].text = 'service/analysis_service.py'
    first_analysis_svc_row[2].text = analysis_svc_methods[0][0]
    first_analysis_svc_row[3].text = analysis_svc_methods[0][1]
    first_analysis_svc_row[4].text = analysis_svc_methods[0][2]

    # 나머지 메소드 행들 (파일 정보 생략)
    for method, complexity, description in analysis_svc_methods[1:]:
        analysis_svc_row = analysis_table.add_row().cells
        analysis_svc_row[0].text = ''  # 레벨 생략
        analysis_svc_row[1].text = ''  # 파일 정보 생략
        analysis_svc_row[2].text = method
        analysis_svc_row[3].text = complexity
        analysis_svc_row[4].text = description

    # JS 레벨 메소드별 행들 (첫 번째 행에 파일 정보 표시)
    analysis_js_methods = [
        ('initializeAnalysisPage()', '중', '데이터 분석 페이지 초기화'),
        ('fetchAndRenderAll()', '고', '전체 데이터 조회 및 렌더링'),
        ('fetchSummaryData(startDate, endDate)', '중', '요약 데이터 조회'),
        ('fetchTrendData(startDate, endDate)', '중', '추이 데이터 조회'),
        ('fetchRawData(startDate, endDate)', '중', '원본 데이터 조회'),
        ('renderSummaryCards(data)', '중', '요약 카드 렌더링'),
        ('renderTrendChart(data)', '중', '추이 차트 렌더링'),
        ('renderRawTable(data)', '중', '원본 데이터 테이블 렌더링'),
        ('renderJobInfoTable()', '중', 'Job 정보 테이블 렌더링'),
        ('initializeEventListeners()', '중', '이벤트 리스너 초기화'),
        ('askGeminiAnalysis()', '고', 'AI 분석 요청 및 결과 표시')
    ]

    # 첫 번째 메소드 행 (파일 정보 포함)
    first_analysis_js_row = analysis_table.add_row().cells
    first_analysis_js_row[0].text = 'JS'
    first_analysis_js_row[1].text = 'static/js/modules/data_analysis/*'
    first_analysis_js_row[2].text = analysis_js_methods[0][0]
    first_analysis_js_row[3].text = analysis_js_methods[0][1]
    first_analysis_js_row[4].text = analysis_js_methods[0][2]

    # 나머지 메소드 행들 (파일 정보 생략)
    for method, complexity, description in analysis_js_methods[1:]:
        analysis_js_row = analysis_table.add_row().cells
        analysis_js_row[0].text = ''  # 레벨 생략
        analysis_js_row[1].text = ''  # 파일 정보 생략
        analysis_js_row[2].text = method
        analysis_js_row[3].text = complexity
        analysis_js_row[4].text = description

    # HTML 레벨 행
    analysis_html_row = analysis_table.add_row().cells
    analysis_html_row[0].text = 'HTML'
    analysis_html_row[1].text = 'templates/data_analysis.html'
    analysis_html_row[2].text = '-'
    analysis_html_row[3].text = '중: 복합 템플릿'
    analysis_html_row[4].text = '필터 폼\n데이터 테이블\n차트 컨테이너\n탭 메뉴'

    doc.add_paragraph('', style='FuncH3')  # 빈 줄

    # 3. 공통 컴포넌트
    doc.add_paragraph('3. 공통 컴포넌트', style='FuncH1')
    common = doc.add_paragraph()
    common.add_run('3.1 데이터베이스 연결\n').bold = True
    common.add_run('파일: msys/database.py\n')
    common.add_run('함수: get_db_connection()\n\n')

    common.add_run('3.2 공통 DAO\n').bold = True
    common.add_run('파일: dao/__init__.py\n')
    common.add_run('기본 CRUD 메소드 제공\n\n')

    # 4. 설정 및 배포
    doc.add_paragraph('4. 설정 및 배포', style='FuncH1')
    deploy = doc.add_paragraph()
    deploy.add_run('4.1 환경 설정\n').bold = True
    deploy.add_run('.env 파일을 통한 환경별 설정 분리\n\n')

    deploy.add_run('4.2 배포 스크립트\n').bold = True
    deploy.add_run('Docker Compose를 활용한 컨테이너화 배포\n\n')

    try:
        doc.save('MSYS_Function_Manual_v8.docx')
        print('기능 매뉴얼 생성 완료: MSYS_Function_Manual_v8.docx')
    except PermissionError:
        doc.save('MSYS_Function_Manual_v9.docx')
        print('기능 매뉴얼 생성 완료: MSYS_Function_Manual_v9.docx')

def create_operation_manual():
    """운영자용 운영 매뉴얼 생성"""
    doc = Document()

    # 스타일 설정
    title_style = doc.styles.add_style('OpTitle', WD_STYLE_TYPE.PARAGRAPH)
    title_style.font.size = Pt(28)
    title_style.font.color.rgb = RGBColor(0, 0, 0)
    title_style.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    heading1_style = doc.styles.add_style('OpH1', WD_STYLE_TYPE.PARAGRAPH)
    heading1_style.font.size = Pt(18)
    heading1_style.font.bold = True

    heading2_style = doc.styles.add_style('OpH2', WD_STYLE_TYPE.PARAGRAPH)
    heading2_style.font.size = Pt(14)
    heading2_style.font.bold = True

    heading3_style = doc.styles.add_style('OpH3', WD_STYLE_TYPE.PARAGRAPH)
    heading3_style.font.size = Pt(12)
    heading3_style.font.bold = True

    # 표지
    title = doc.add_paragraph('MSYS 운영 매뉴얼', style='OpTitle')
    subtitle = doc.add_paragraph('시스템 운영 및 유지보수 가이드', style='OpTitle')
    version = doc.add_paragraph('버전 1.14.2', style='OpTitle')
    version.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    doc.add_page_break()

    # 목차
    toc_title = doc.add_paragraph('목차', style='OpH1')
    toc = doc.add_paragraph()
    toc.add_run('1. 일상 운영 절차 ........................ 3\n')
    toc.add_run('2. 모니터링 ............................. 5\n')
    toc.add_run('   2.1 대시보드 모니터링 .................. 5\n')
    toc.add_run('   2.2 데이터 분석 모니터링 ................ 15\n')
    toc.add_run('   2.3 차트 분석 모니터링 .................. 25\n')
    toc.add_run('   2.4 잔디 현황 모니터링 .................. 35\n')
    toc.add_run('   2.5 매핑 관리 모니터링 .................. 45\n')
    toc.add_run('   2.6 데이터 명세서 모니터링 ................ 55\n')
    toc.add_run('   2.7 관리자 설정 모니터링 ................ 65\n')
    toc.add_run('   2.8 수집 일정 모니터링 .................. 75\n')
    toc.add_run('   2.9 카드 요약 모니터링 .................. 85\n')
    toc.add_run('   2.10 시스템 리소스 모니터링 .............. 95\n')
    toc.add_run('3. 백업 및 복원 ........................ 105\n')
    toc.add_run('4. 장애 대응 ........................... 108\n')
    toc.add_run('5. 성능 최적화 ........................ 111\n')
    toc.add_run('6. 인증 시스템 ........................ 114\n')
    toc.add_run('7. 로그 관리 ........................... 125\n')
    toc.add_run('8. 운영 시나리오 ........................ 140\n')
    toc.add_run('부록 A. 점검 체크리스트 .................. 200\n')

    doc.add_page_break()

    # 1. 일상 운영 절차
    doc.add_paragraph('1. 일상 운영 절차', style='OpH1')

    # 1.1 시스템 시작
    doc.add_paragraph('1.1 시스템 시작', style='OpH2')
    start_sys = doc.add_paragraph()
    start_sys.add_run('1.1.1 서버 접속\n').bold = True
    start_sys.add_run('SSH를 사용하여 서버에 접속합니다:\n\n')
    start_sys.add_run('ssh msys@server_ip_address\n\n').font.name = 'Courier New'

    start_sys.add_run('1.1.2 애플리케이션 디렉토리 이동\n').bold = True
    start_sys.add_run('MSYS 애플리케이션 디렉토리로 이동합니다:\n\n')
    start_sys.add_run('cd /home/msys/app\n\n').font.name = 'Courier New'

    start_sys.add_run('1.1.3 가상환경 활성화\n').bold = True
    start_sys.add_run('Python 가상환경을 활성화합니다:\n\n')
    start_sys.add_run('source venv/bin/activate\n\n').font.name = 'Courier New'

    start_sys.add_run('1.1.4 애플리케이션 실행\n').bold = True
    start_sys.add_run('MSYS 애플리케이션을 실행합니다:\n\n')
    start_sys.add_run('python msys_app.py\n\n').font.name = 'Courier New'
    start_sys.add_run('※ 애플리케이션이 백그라운드에서 실행되도록 하려면 nohup을 사용합니다:\n\n')
    start_sys.add_run('nohup python msys_app.py > /dev/null 2>&1 &\n\n').font.name = 'Courier New'

    # 1.2 시스템 중지
    doc.add_paragraph('1.2 시스템 중지', style='OpH2')
    stop_sys = doc.add_paragraph()
    stop_sys.add_run('1.2.1 정상적인 종료\n').bold = True
    stop_sys.add_run('실행 중인 프로세스를 찾아 종료합니다:\n\n')
    stop_sys.add_run('ps aux | grep python\n').font.name = 'Courier New'
    stop_sys.add_run('kill -TERM <process_id>\n\n').font.name = 'Courier New'

    stop_sys.add_run('1.2.2 강제 종료 (비상시)\n').bold = True
    stop_sys.add_run('응답이 없는 경우 강제 종료합니다:\n\n')
    stop_sys.add_run('kill -9 <process_id>\n\n').font.name = 'Courier New'

    # 1.3 서비스 상태 확인
    doc.add_paragraph('1.3 서비스 상태 확인', style='OpH2')
    check_svc = doc.add_paragraph()
    check_svc.add_run('1.3.1 프로세스 상태 확인\n').bold = True
    check_svc.add_run('Python 애플리케이션 프로세스 상태를 확인합니다:\n\n')
    check_svc.add_run('ps aux | grep python\n\n').font.name = 'Courier New'

    check_svc.add_run('1.3.2 데이터베이스 상태 확인\n').bold = True
    check_svc.add_run('PostgreSQL 서비스 상태를 확인합니다:\n\n')
    check_svc.add_run('systemctl status postgresql\n\n').font.name = 'Courier New'

    check_svc.add_run('1.3.3 웹 서비스 상태 확인\n').bold = True
    check_svc.add_run('웹 서버가 정상 응답하는지 확인합니다:\n\n')
    check_svc.add_run('curl -I http://localhost:5000\n\n').font.name = 'Courier New'

    # 2. 모니터링
    doc.add_paragraph('2. 모니터링', style='OpH1')

    # 2.1 대시보드 모니터링
    doc.add_paragraph('2.1 대시보드 모니터링', style='OpH2')
    dash_monitor = doc.add_paragraph()
    dash_monitor.add_run('웹 인터페이스를 통해 실시간으로 다음 항목을 모니터링합니다:\n')
    dash_monitor.add_run('• 데이터 수집 성공률\n')
    dash_monitor.add_run('• 시스템 상태\n')
    dash_monitor.add_run('• 오류 발생 현황\n\n')

    # 대시보드 스크린샷 삽입
    try:
        doc.add_picture('scrennshot/dashboard_screenshot.PNG', width=Inches(6))
        doc.add_paragraph('그림 2.1: 대시보드 모니터링 화면', style='OpH3')
    except:
        doc.add_paragraph('[스크린샷: dashboard_screenshot.PNG를 찾을 수 없습니다]', style='OpH3')

    doc.add_paragraph('', style='OpH3')  # 빈 줄

    dash_desc = doc.add_paragraph()
    dash_desc.add_run('2.1.1 대시보드 화면 구성 요소 상세 설명\n').bold = True

    dash_desc.add_run('1. 요약 카드 영역:\n').bold = True
    dash_desc.add_run('• 전체 Job ID 개수: 시스템에 등록된 데이터 수집 작업의 총 개수\n')
    dash_desc.add_run('• 총 수집 건수: 선택된 기간 동안 성공적으로 수집된 데이터의 총 건수\n')
    dash_desc.add_run('• 평균 성공률: 모든 Job ID의 평균 데이터 수집 성공률 (백분율)\n\n')

    dash_desc.add_run('2. 상세 테이블:\n').bold = True
    dash_desc.add_run('• Job ID: 각 데이터 수집 작업의 고유 식별자 (예: "CD101", "CD102", "CD103", "CD104")\n')
    dash_desc.add_run('  - 시스템에서 자동으로 부여되는 작업 식별 번호\n')
    dash_desc.add_run('  - 각 작업을 구분하고 추적하는 데 사용\n')
    dash_desc.add_run('  - CD + 3자리 숫자 형식으로 구성\n\n')

    dash_desc.add_run('• 데이터명: 작업의 실제 데이터 수집 대상 이름\n')
    dash_desc.add_run('  - forecasts: 예보 데이터 수집\n')
    dash_desc.add_run('  - forecasts_o: 예보 데이터 수집 (외부)\n')
    dash_desc.add_run('  - live: 실시간 데이터 수집\n')
    dash_desc.add_run('  - live_o: 실시간 데이터 수집 (외부)\n\n')

    dash_desc.add_run('• 주기(cron): 데이터 수집 실행 주기\n')
    dash_desc.add_run('  - "0 6,10 * * *": 매일 6시, 10시에 실행\n')
    dash_desc.add_run('  - "0 16 */6 * *": 6일마다 16시에 실행\n')
    dash_desc.add_run('  - cron 표현식으로 자동화된 스케줄링\n\n')

    dash_desc.add_run('• 성공률: 해당 기간 동안의 데이터 수집 성공률 (백분율)\n')
    dash_desc.add_run('  - 기간별, Job ID별 통계를 모두 계산하여 산출\n')
    dash_desc.add_run('  - 성공한 수집 건수 ÷ 전체 수집 시도 건수 × 100\n')
    dash_desc.add_run('  - 오늘/주간/월간 등 다양한 기간 단위로 계산 가능\n\n')

    dash_desc.add_run('• 연속 실패 횟수: 최근 연속으로 실패한 데이터 수집 횟수\n')
    dash_desc.add_run('  - 마지막 성공 이후부터 현재까지 연속 실패한 시도 횟수\n')
    dash_desc.add_run('  - 0: 최근에 성공한 기록이 있음\n')
    dash_desc.add_run('  - 높을수록: 장기간 데이터 수집 실패 상태 지속\n\n')

    dash_desc.add_run('• 임계값: 설정된 성공률 기준값 (백분율)\n')
    dash_desc.add_run('  - 이 값 이상이면 정상, 미만이면 경고/위험 상태로 판단\n')
    dash_desc.add_run('  - 관리자가 각 Job별로 개별 설정 가능\n\n')

    dash_desc.add_run('• 상태: 현재 작업의 상태 표시 (정상/경고/위험)\n')
    dash_desc.add_run('  - 정상: 성공률 ≥ 임계값 AND 연속 실패 < 3회\n')
    dash_desc.add_run('  - 경고: (성공률 80-95% 사이) OR (연속 실패 3-5회)\n')
    dash_desc.add_run('  - 위험: (성공률 < 80%) OR (연속 실패 ≥ 5회)\n\n')

    dash_desc.add_run('2.1.2 관리자 설정과의 연동\n').bold = True
    dash_desc.add_run('대시보드에 표시되는 모든 값들은 관리자 설정에서 정의된 임계값과 설정을 기반으로 계산됩니다:\n\n')

    dash_desc.add_run('• 연속 실패 임계값 (cnn_failr_thrs_val)\n').bold = True
    dash_desc.add_run('  - 기본값: 3회\n')
    dash_desc.add_run('  - 의미: 이 횟수만큼 연속으로 실패하면 위험 상태로 표시\n')
    dash_desc.add_run('  - 설정 위치: 관리자 설정 → 기본 설정 → 연속 실패(CNT)\n\n')

    dash_desc.add_run('• 경고 임계값 (cnn_warn_thrs_val)\n').bold = True
    dash_desc.add_run('  - 기본값: 2회\n')
    dash_desc.add_run('  - 의미: 이 횟수만큼 연속으로 실패하면 경고 상태로 표시\n')
    dash_desc.add_run('  - 설정 위치: 관리자 설정 → 기본 설정 → 경고(CNT)\n\n')

    dash_desc.add_run('• 일별 성공률 임계값 (dly_sucs_rt_thrs_val)\n').bold = True
    dash_desc.add_run('  - 기본값: 80%\n')
    dash_desc.add_run('  - 의미: 일별 성공률이 이 값 미만이면 위험 상태로 표시\n')
    dash_desc.add_run('  - 설정 위치: 관리자 설정 → 기본 설정 → 일간 임계값\n\n')

    dash_desc.add_run('• 주간 성공률 임계값 (dd7_sucs_rt_thrs_val)\n').bold = True
    dash_desc.add_run('  - 기본값: 75%\n')
    dash_desc.add_run('  - 의미: 7일간 성공률이 이 값 미만이면 위험 상태로 표시\n')
    dash_desc.add_run('  - 설정 위치: 관리자 설정 → 기본 설정 → 주간 임계값\n\n')

    dash_desc.add_run('• 월별 성공률 임계값 (mthl_sucs_rt_thrs_val)\n').bold = True
    dash_desc.add_run('  - 기본값: 70%\n')
    dash_desc.add_run('  - 의미: 월별 성공률이 이 값 미만이면 위험 상태로 표시\n')
    dash_desc.add_run('  - 설정 위치: 관리자 설정 → 기본 설정 → 월간 임계값\n\n')

    dash_desc.add_run('• 상태별 아이콘 설정\n').bold = True
    dash_desc.add_run('  - 성공 아이콘: 정상 상태 표시 (기본: ✅)\n')
    dash_desc.add_run('  - 경고 아이콘: 경고 상태 표시 (기본: ⚠️)\n')
    dash_desc.add_run('  - 실패 아이콘: 위험 상태 표시 (기본: ❌)\n')
    dash_desc.add_run('  - 설정 위치: 관리자 설정 → 기본 설정 → 각 상태별 아이콘 선택\n\n')

    dash_desc.add_run('• 상태별 색상 설정\n').bold = True
    dash_desc.add_run('  - 성공 색상: 정상 상태 배경색 (기본: #008000)\n')
    dash_desc.add_run('  - 경고 색상: 경고 상태 배경색 (기본: #FFA500)\n')
    dash_desc.add_run('  - 실패 색상: 위험 상태 배경색 (기본: #FF0000)\n')
    dash_desc.add_run('  - 설정 위치: 관리자 설정 → 기본 설정 → 각 상태별 색상 선택\n\n')

    dash_desc.add_run('• 대시보드 표시 여부 (chrt_dsp_yn)\n').bold = True
    dash_desc.add_run('  - 기본값: true (표시)\n')
    dash_desc.add_run('  - 의미: 해당 Job을 대시보드에 표시할지 여부 제어\n')
    dash_desc.add_run('  - 설정 위치: 관리자 설정 → 기본 설정 → 대시보드 표시 여부\n\n')

    dash_desc.add_run('3. 날짜 필터 컨트롤:\n').bold = True
    dash_desc.add_run('• 시작일/종료일: 모니터링할 기간을 선택\n')
    dash_desc.add_run('• 조회 버튼: 선택된 기간의 데이터를 다시 로드\n')
    dash_desc.add_run('• 기간별 보기: 일간/주간/월간 데이터 집계 방식 선택\n\n')

    dash_desc.add_run('4. 상태 표시 색상 코드 및 조건:\n').bold = True
    dash_desc.add_run('• 녹색 (●): 정상 상태 - (성공률 ≥ 임계값) AND (연속 실패 < 3)\n')
    dash_desc.add_run('• 노랑 (●): 경고 상태 - (성공률 80-95%) OR (연속 실패 3-5)\n')
    dash_desc.add_run('• 빨강 (●): 위험 상태 - (성공률 < 80%) OR (연속 실패 ≥ 5)\n\n')

    dash_desc.add_run('2.1.2 데이터 계산 방식\n').bold = True
    dash_desc.add_run('• 성공률 계산: (성공 건수 ÷ 총 시도 건수) × 100\n')
    dash_desc.add_run('• 연속 실패: 마지막 성공 이후 연속 실패 횟수\n')
    dash_desc.add_run('• 기간별 집계: 일별(당일)/주별(최근7일)/월별(최근30일)\n\n')

    dash_desc.add_run('2.1.3 실제 데이터 예시\n').bold = True
    dash_desc.add_run('• 정상 케이스: CD101 (성공률 98%, 연속 실패 0회) → 녹색 표시\n')
    dash_desc.add_run('• 경고 케이스: CD102 (성공률 85%, 연속 실패 2회) → 노랑 표시\n')
    dash_desc.add_run('• 위험 케이스: CD103 (성공률 65%, 연속 실패 7회) → 빨강 표시\n\n')

    dash_desc.add_run('※ 해당 값들은 관리자 설정에 의해 변경이 가능하며 현재 값들은 기본값을 기준으로 작성되었다.\n\n')

    # 2.2 데이터 분석 모니터링
    doc.add_paragraph('2.2 데이터 분석 모니터링', style='OpH2')
    data_analysis_monitor = doc.add_paragraph()
    data_analysis_monitor.add_run('데이터 수집 결과를 심층 분석하고 필터링을 통한 데이터 조회를 수행합니다.\n\n')

    # 데이터 분석 스크린샷 삽입
    try:
        doc.add_picture('d:/dev/msys/scrennshot/data_analysis_screenshot.PNG', width=Inches(6))
        doc.add_paragraph('그림 2.2: 데이터 분석 화면', style='OpH3')
    except:
        doc.add_paragraph('[스크린샷: data_analysis_screenshot.PNG를 찾을 수 없습니다]', style='OpH3')

    doc.add_paragraph('', style='OpH3')  # 빈 줄

    data_analysis_desc = doc.add_paragraph()
    data_analysis_desc.add_run('2.2.1 데이터 분석 화면 구성 요소 상세 설명\n').bold = True

    data_analysis_desc.add_run('1. 필터 컨트롤 영역:\n').bold = True
    data_analysis_desc.add_run('• 기간 선택: 분석할 데이터의 시작일과 종료일 설정\n')
    data_analysis_desc.add_run('• Job ID 선택: 분석할 특정 Job ID 선택 (다중 선택 가능)\n')
    data_analysis_desc.add_run('• 장애코드 필터: 특정 오류 코드로 데이터 필터링\n\n')

    data_analysis_desc.add_run('2. 데이터 조회 영역:\n').bold = True
    data_analysis_desc.add_run('• 요약 데이터: 선택된 기간의 전체 통계 요약\n')
    data_analysis_desc.add_run('• 추이 데이터: 시간에 따른 데이터 수집 추이 그래프\n')
    data_analysis_desc.add_run('• 원천 데이터: 실제 수집된 데이터의 상세 목록\n\n')

    data_analysis_desc.add_run('3. 데이터 분석 영역:\n').bold = True
    data_analysis_desc.add_run('• 데이터 필터링: Job ID, 장애코드, 기간별 필터링 기능\n')
    data_analysis_desc.add_run('• 분석 결과 표시: 필터링된 데이터의 통계 및 추이 분석\n\n')

    data_analysis_desc.add_run('2.2.2 데이터 구조 및 의미\n').bold = True
    data_analysis_desc.add_run('• 요약 데이터: 기간 내 총 수집 건수, 성공률, 실패 건수 등 집계 정보\n')
    data_analysis_desc.add_run('• 추이 데이터: 일별/시간별 데이터 수집 성공률 변화 추이\n')
    data_analysis_desc.add_run('• 원천 데이터: 각 수집 시도의 상세 정보 (시간, 상태, 오류 코드 등)\n\n')

    data_analysis_desc.add_run('2.2.3 사용자 인터랙션\n').bold = True
    data_analysis_desc.add_run('1. 필터 설정 후 [조회] 버튼 클릭하여 데이터 로드\n')
    data_analysis_desc.add_run('2. [AI 분석] 버튼 클릭하여 자동 분석 실행\n')
    data_analysis_desc.add_run('3. 분석 결과를 검토하고 필요한 조치 수행\n\n')

    data_analysis_desc.add_run('2.2.4 실제 데이터 예시\n').bold = True
    data_analysis_desc.add_run('• 정상 분석 케이스: CD101 (2025-12-01~2025-12-07) - 성공률 95%, 안정적인 수집 패턴\n')
    data_analysis_desc.add_run('• 문제 분석 케이스: CD102 (2025-12-01~2025-12-07) - 성공률 75%, 연속 실패 증가 추세\n\n')

    # 2.3 차트 분석 모니터링
    doc.add_paragraph('2.3 차트 분석 모니터링', style='OpH2')
    chart_analysis_monitor = doc.add_paragraph()
    chart_analysis_monitor.add_run('데이터 수집 성공률의 시간적 추이와 장애 코드별 현황을 시각적인 차트로 분석합니다.\n\n')

    # 차트 분석 스크린샷 삽입
    try:
        doc.add_picture('d:/dev/msys/scrennshot/chart_analysis_screenshot.PNG', width=Inches(6))
        doc.add_paragraph('그림 2.3: 차트 분석 화면', style='OpH3')
    except:
        doc.add_paragraph('[스크린샷: chart_analysis_screenshot.PNG를 찾을 수 없습니다]', style='OpH3')

    # 잔디 현황 스크린샷 삽입
    try:
        doc.add_picture('d:/dev/msys/scrennshot/jandi_screenshot.PNG', width=Inches(6))
        doc.add_paragraph('그림 2.4: 잔디 현황 화면', style='OpH3')
    except:
        doc.add_paragraph('[스크린샷: jandi_screenshot.PNG를 찾을 수 없습니다]', style='OpH3')

    # 매핑 관리 스크린샷 삽입
    try:
        doc.add_picture('d:/dev/msys/scrennshot/mapping_screenshot.PNG', width=Inches(6))
        doc.add_paragraph('그림 2.5: 매핑 관리 화면', style='OpH3')
    except:
        doc.add_paragraph('[스크린샷: mapping_screenshot.PNG를 찾을 수 없습니다]', style='OpH3')

    # 데이터 명세서 스크린샷 삽입
    try:
        doc.add_picture('d:/dev/msys/scrennshot/data_spec_screenshot.PNG', width=Inches(6))
        doc.add_paragraph('그림 2.6: 데이터 명세서 화면', style='OpH3')
    except:
        doc.add_paragraph('[스크린샷: data_spec_screenshot.PNG를 찾을 수 없습니다]', style='OpH3')

    doc.add_paragraph('', style='OpH3')  # 빈 줄

    chart_analysis_desc = doc.add_paragraph()
    chart_analysis_desc.add_run('2.3.1 차트 분석 화면 구성 요소 상세 설명\n').bold = True

    chart_analysis_desc.add_run('1. 필터 컨트롤 영역:\n').bold = True
    chart_analysis_desc.add_run('• 기간 선택: 차트에 표시할 데이터의 시작일과 종료일 설정\n')
    chart_analysis_desc.add_run('• Job ID 선택: 분석할 Job ID 체크박스 (다중 선택 가능)\n')
    chart_analysis_desc.add_run('• 조회 버튼: 선택된 조건으로 차트 데이터 재로딩\n\n')

    chart_analysis_desc.add_run('2. 성공률 추이 차트:\n').bold = True
    chart_analysis_desc.add_run('• 시간에 따른 각 Job의 성공률 변화 그래프\n')
    chart_analysis_desc.add_run('• X축: 날짜/시간, Y축: 성공률 (백분율)\n')
    chart_analysis_desc.add_run('• 범례: 각 Job ID별 색상 구분\n\n')

    chart_analysis_desc.add_run('3. 장애 코드별 현황 차트:\n').bold = True
    chart_analysis_desc.add_run('• 선택된 기간 내 장애 코드 발생 빈도 파이 차트 또는 바 차트\n')
    chart_analysis_desc.add_run('• 각 장애 코드별 발생 건수 및 백분율 표시\n')
    chart_analysis_desc.add_run('• 상세 정보: 마우스 오버 시 구체적인 수치 표시\n\n')

    chart_analysis_desc.add_run('2.3.2 데이터 구조 및 의미\n').bold = True
    chart_analysis_desc.add_run('• 성공률 추이: 일별/시간별 데이터 수집 성공률의 시계열 데이터\n')
    chart_analysis_desc.add_run('• 장애 코드 현황: 기간 내 각 오류 코드의 발생 빈도 통계\n')
    chart_analysis_desc.add_run('• 차트 데이터: Chart.js 라이브러리를 사용한 인터랙티브 시각화\n\n')

    chart_analysis_desc.add_run('2.3.3 사용자 인터랙션\n').bold = True
    chart_analysis_desc.add_run('1. 기간과 Job ID 선택 후 [조회] 버튼 클릭\n')
    chart_analysis_desc.add_run('2. 차트 위에 마우스 오버하여 상세 데이터 확인\n')
    chart_analysis_desc.add_run('3. 범례 클릭으로 특정 Job의 표시/숨김 토글\n\n')

    chart_analysis_desc.add_run('2.3.4 실제 데이터 예시\n').bold = True
    chart_analysis_desc.add_run('• 성공률 추이: CD101의 12월 한 달간 성공률 95-98% 안정적 유지\n')
    chart_analysis_desc.add_run('• 장애 코드 현황: ERR001 (50%), ERR002 (30%), 기타 (20%)\n\n')

    # 2.4 잔디 현황 모니터링
    doc.add_paragraph('2.4 잔디 현황 모니터링', style='OpH2')
    jandi_monitor = doc.add_paragraph()
    jandi_monitor.add_run('Github의 활동 그래프처럼 각 Job의 일별 데이터 수집 현황을 히트맵 형태로 시각화하여 보여줍니다.\n\n')

    # 잔디 스크린샷 삽입
    try:
        doc.add_picture('scrennshot/jandi_screenshot.PNG', width=Inches(6))
        doc.add_paragraph('그림 2.4: 잔디 현황 화면', style='OpH3')
    except:
        doc.add_paragraph('[스크린샷: jandi_screenshot.PNG를 찾을 수 없습니다]', style='OpH3')

    doc.add_paragraph('', style='OpH3')  # 빈 줄

    jandi_desc = doc.add_paragraph()
    jandi_desc.add_run('2.4.1 잔디 현황 화면 구성 요소 상세 설명\n').bold = True

    jandi_desc.add_run('1. 기간 선택 컨트롤:\n').bold = True
    jandi_desc.add_run('• 시작일/종료일: 히트맵에 표시할 기간 설정\n')
    jandi_desc.add_run('• 조회 버튼: 선택된 기간의 데이터 재로딩\n\n')

    jandi_desc.add_run('2. Job 목록 테이블:\n').bold = True
    jandi_desc.add_run('• Job ID: 각 데이터 수집 작업의 고유 식별자\n')
    jandi_desc.add_run('• Job 이름: 작업의 표시 이름\n')
    jandi_desc.add_run('• 펼치기/접기 버튼: 히트맵 표시 토글\n\n')

    jandi_desc.add_run('3. 히트맵 영역:\n').bold = True
    jandi_desc.add_run('• 일별 활동 표시: 각 날짜의 수집 건수에 따른 색상 표시\n')
    jandi_desc.add_run('• 색상 범례: 수집 건수별 색상 구분 (낮음 → 높음)\n')
    jandi_desc.add_run('• 마우스 오버: 구체적인 날짜와 수집 건수 표시\n\n')

    jandi_desc.add_run('2.4.2 데이터 구조 및 의미\n').bold = True
    jandi_desc.add_run('• 히트맵 데이터: 일별 수집 성공 건수 기반 색상 매핑\n')
    jandi_desc.add_run('• 색상 등급: 0건(흰색) → 1-5건(연한 녹색) → 6-10건(진한 녹색) → 10건+(어두운 녹색)\n')
    jandi_desc.add_run('• 데이터 집계: 선택된 기간 내 일별 총 수집 건수\n\n')

    jandi_desc.add_run('2.4.3 상태 표시 기준 및 색상 코드\n').bold = True
    jandi_desc.add_run('• 활동 없음: 흰색 (수집 건수 = 0)\n')
    jandi_desc.add_run('• 낮은 활동: 연한 녹색 (수집 건수 1-5)\n')
    jandi_desc.add_run('• 보통 활동: 중간 녹색 (수집 건수 6-10)\n')
    jandi_desc.add_run('• 높은 활동: 진한 녹색 (수집 건수 11+)\n\n')

    jandi_desc.add_run('2.4.4 실제 데이터 예시\n').bold = True
    jandi_desc.add_run('• 정상 패턴: CD101의 12월 히트맵 - 매일 8-12건 수집 (진한 녹색)\n')
    jandi_desc.add_run('• 불규칙 패턴: CD102의 12월 히트맵 - 주말 낮은 활동, 평일 높은 활동\n')
    jandi_desc.add_run('• 문제 패턴: CD103의 12월 히트맵 - 중간에 연속 흰색 구간 (수집 실패)\n\n')

    jandi_desc.add_run('2.4.5 사용자 인터랙션\n').bold = True
    jandi_desc.add_run('1. 기간 선택 후 [조회] 버튼 클릭하여 Job 목록 로드\n')
    jandi_desc.add_run('2. 관심 있는 Job의 [펼치기] 버튼 클릭하여 히트맵 표시\n')
    jandi_desc.add_run('3. 히트맵 위 마우스 오버하여 상세 데이터 확인\n')
    jandi_desc.add_run('4. 색상 범례로 활동 강도 파악\n\n')

    # 2.5 매핑 관리 모니터링
    doc.add_paragraph('2.5 매핑 관리 모니터링', style='OpH2')
    mapping_monitor = doc.add_paragraph()
    mapping_monitor.add_run('데이터베이스 테이블 간의 컬럼 매핑 정보를 관리하고 조회하는 기능을 제공합니다.\n\n')

    # 매핑 스크린샷 삽입
    try:
        doc.add_picture('scrennshot/mapping_screenshot.PNG', width=Inches(6))
        doc.add_paragraph('그림 2.5: 매핑 관리 화면', style='OpH3')
    except:
        doc.add_paragraph('[스크린샷: mapping_screenshot.PNG를 찾을 수 없습니다]', style='OpH3')

    doc.add_paragraph('', style='OpH3')  # 빈 줄

    mapping_desc = doc.add_paragraph()
    mapping_desc.add_run('2.5.1 매핑 관리 화면 구성 요소 상세 설명\n').bold = True

    mapping_desc.add_run('1. 매핑 조회 영역:\n').bold = True
    mapping_desc.add_run('• 모든 매핑 보기: 시스템의 모든 매핑 정보 표시\n')
    mapping_desc.add_run('• 미매핑 컬럼 보기: 아직 매핑되지 않은 컬럼 목록\n')
    mapping_desc.add_run('• 새 매핑 추가: 새로운 매핑 관계 생성\n\n')

    mapping_desc.add_run('2. 매핑 테이블:\n').bold = True
    mapping_desc.add_run('• 소스 테이블: 원본 데이터 테이블 이름\n')
    mapping_desc.add_run('• 소스 컬럼: 원본 테이블의 컬럼명\n')
    mapping_desc.add_run('• 타겟 테이블: 대상 데이터 테이블 이름\n')
    mapping_desc.add_run('• 타겟 컬럼: 대상 테이블의 컬럼명\n')
    mapping_desc.add_run('• 매핑 유형: 매핑 관계의 종류 (직접, 변환, 계산 등)\n\n')

    mapping_desc.add_run('3. 매핑 작업 버튼:\n').bold = True
    mapping_desc.add_run('• 추가: 새로운 매핑 관계 생성\n')
    mapping_desc.add_run('• 수정: 기존 매핑 관계 편집\n')
    mapping_desc.add_run('• 삭제: 매핑 관계 제거\n')
    mapping_desc.add_run('• 내보내기: 매핑 정보를 파일로 저장\n\n')

    mapping_desc.add_run('2.5.2 데이터 구조 및 의미\n').bold = True
    mapping_desc.add_run('• 매핑 정보: 테이블 간 컬럼 관계 정의 데이터\n')
    mapping_desc.add_run('• 소스/타겟: 데이터 변환의 입력과 출력 관계\n')
    mapping_desc.add_run('• 매핑 유형: 데이터 변환 방식 (1:1, 1:N, 계산식 등)\n\n')

    mapping_desc.add_run('2.5.3 사용자 인터랙션\n').bold = True
    mapping_desc.add_run('1. [모든 매핑 보기] 또는 [미매핑 컬럼 보기] 선택\n')
    mapping_desc.add_run('2. 매핑할 컬럼 선택 후 [새 매핑 추가] 클릭\n')
    mapping_desc.add_run('3. 소스와 타겟 컬럼 관계 설정\n')
    mapping_desc.add_run('4. [저장] 버튼으로 매핑 관계 저장\n\n')

    mapping_desc.add_run('2.5.4 실제 데이터 예시\n').bold = True
    mapping_desc.add_run('• 직접 매핑: source_table.id → target_table.user_id\n')
    mapping_desc.add_run('• 변환 매핑: source_table.status_code → target_table.status (코드 변환)\n')
    mapping_desc.add_run('• 계산 매핑: source_table.price * 1.1 → target_table.final_price\n\n')

    # 2.6 데이터 명세서 모니터링
    doc.add_paragraph('2.6 데이터 명세서 모니터링', style='OpH2')
    data_spec_monitor = doc.add_paragraph()
    data_spec_monitor.add_run('시스템에서 사용하는 외부 데이터(API 등)의 명세를 관리하는 기능을 제공합니다.\n\n')

    # 데이터 명세서 스크린샷 삽입
    try:
        doc.add_picture('scrennshot/data_spec_screenshot.PNG', width=Inches(6))
        doc.add_paragraph('그림 2.6: 데이터 명세서 화면', style='OpH3')
    except:
        doc.add_paragraph('[스크린샷: data_spec_screenshot.PNG를 찾을 수 없습니다]', style='OpH3')

    doc.add_paragraph('', style='OpH3')  # 빈 줄

    data_spec_desc = doc.add_paragraph()
    data_spec_desc.add_run('2.6.1 데이터 명세서 화면 구성 요소 상세 설명\n').bold = True

    data_spec_desc.add_run('1. 명세서 목록:\n').bold = True
    data_spec_desc.add_run('• 명세서 이름: 각 데이터 소스의 식별자\n')
    data_spec_desc.add_run('• 설명: 명세서에 대한 간단한 설명\n')
    data_spec_desc.add_run('• URL: 데이터 소스의 접근 주소\n')
    data_spec_desc.add_run('• 최종 수정일: 명세서 마지막 업데이트 일시\n\n')

    data_spec_desc.add_run('2. 명세서 작업 버튼:\n').bold = True
    data_spec_desc.add_run('• 새 명세서: 새로운 데이터 명세서 생성\n')
    data_spec_desc.add_run('• URL에서 가져오기: 웹 주소로부터 명세서 자동 생성\n')
    data_spec_desc.add_run('• 수정: 기존 명세서 편집\n')
    data_spec_desc.add_run('• 삭제: 명세서 제거\n\n')

    data_spec_desc.add_run('3. 명세서 상세 정보:\n').bold = True
    data_spec_desc.add_run('• API 엔드포인트: 데이터 접근 주소\n')
    data_spec_desc.add_run('• 요청 방식: HTTP 메소드 (GET, POST 등)\n')
    data_spec_desc.add_run('• 파라미터: API 호출에 필요한 매개변수\n')
    data_spec_desc.add_run('• 응답 형식: 반환 데이터 구조 (JSON, XML 등)\n\n')

    data_spec_desc.add_run('2.6.2 데이터 구조 및 의미\n').bold = True
    data_spec_desc.add_run('• 명세서 메타데이터: API 이름, 설명, 접근 권한 등\n')
    data_spec_desc.add_run('• API 스펙: 엔드포인트, 메소드, 파라미터, 응답 형식\n')
    data_spec_desc.add_run('• 데이터 스키마: 요청/응답 데이터 구조 정의\n\n')

    data_spec_desc.add_run('2.6.3 사용자 인터랙션\n').bold = True
    data_spec_desc.add_run('1. [새 명세서] 클릭하여 새로운 명세서 생성 시작\n')
    data_spec_desc.add_run('2. 또는 [URL에서 가져오기]로 자동 생성\n')
    data_spec_desc.add_run('3. 명세서 정보 입력 (이름, URL, 설명 등)\n')
    data_spec_desc.add_run('4. [저장]으로 명세서 등록 완료\n\n')

    data_spec_desc.add_run('2.6.4 실제 데이터 예시\n').bold = True
    data_spec_desc.add_run('• 날씨 API: 기상청 API 명세서 (GET /weather, 파라미터: location, date)\n')
    data_spec_desc.add_run('• 사용자 API: 외부 시스템 사용자 정보 (POST /users, 인증 토큰 필요)\n')
    data_spec_desc.add_run('• 금융 API: 주식 정보 조회 (GET /stocks/{symbol}, 실시간 데이터)\n\n')

    # 2.7 관리자 설정 모니터링
    doc.add_paragraph('2.7 관리자 설정 모니터링', style='OpH2')
    admin_monitor = doc.add_paragraph()
    admin_monitor.add_run('시스템의 각종 설정을 관리하고 모니터링하는 기능을 제공합니다.\n\n')

    # 관리자 설정 스크린샷 삽입
    try:
        doc.add_picture('d:/dev/msys/scrennshot/admin_basic_screenshot.PNG', width=Inches(6))
        doc.add_paragraph('그림 2.7: 관리자 설정 기본 탭 화면', style='OpH3')
    except:
        doc.add_paragraph('[스크린샷: admin_basic_screenshot.PNG를 찾을 수 없습니다]', style='OpH3')

    doc.add_paragraph('', style='OpH3')  # 빈 줄

    admin_desc = doc.add_paragraph()
    admin_desc.add_run('2.7.1 관리자 설정 화면 구성 요소 상세 설명\n').bold = True

    admin_desc.add_run('1. 탭 메뉴:\n').bold = True
    admin_desc.add_run('화면 상단의 탭을 클릭하여 다른 설정 화면으로 이동할 수 있습니다.\n')
    admin_desc.add_run('• 기본 설정 (현재 선택됨)\n')
    admin_desc.add_run('• 수집 스케줄 설정\n')
    admin_desc.add_run('• Icon 관리\n')
    admin_desc.add_run('• 차트/시각화 설정\n')
    admin_desc.add_run('• 사용자 관리\n')
    admin_desc.add_run('• 데이터 접근 권한\n')
    admin_desc.add_run('• 통계\n\n')

    admin_desc.add_run('2. 설정 테이블:\n').bold = True
    admin_desc.add_run('Job ID별 설정을 표 형식으로 표시합니다.\n')
    admin_desc.add_run('• Job ID: 작업 고유 식별자\n')
    admin_desc.add_run('• Job 이름: 작업의 표시 이름\n')
    admin_desc.add_run('• Job 설명: 작업에 대한 설명\n')
    admin_desc.add_run('• 비고: 추가 메모\n\n')

    admin_desc.add_run('3. 연속 실패시 표시방법:\n').bold = True
    admin_desc.add_run('데이터 수집이 연속으로 실패할 경우의 표시 설정입니다.\n')
    admin_desc.add_run('• 실패(CNT): 연속 실패 횟수 임계값\n')
    admin_desc.add_run('• 실패(Icon): 실패 상태 아이콘 선택\n')
    admin_desc.add_run('• 실패(Color): 실패 상태 색상 선택\n')
    admin_desc.add_run('• 경고(CNT): 경고 상태 임계값\n')
    admin_desc.add_run('• 경고(Icon): 경고 상태 아이콘 선택\n')
    admin_desc.add_run('• 경고(Color): 경고 상태 색상 선택\n')
    admin_desc.add_run('• 정상(Icon): 정상 상태 아이콘 선택\n')
    admin_desc.add_run('• 정상(Color): 정상 상태 색상 선택\n\n')

    admin_desc.add_run('4. 성공률 기준 임계치 설정:\n').bold = True
    admin_desc.add_run('각 기간별 성공률 임계값을 설정합니다.\n')
    admin_desc.add_run('• 일간 임계값: 일일 성공률 기준\n')
    admin_desc.add_run('• 주간 임계값: 주간 성공률 기준\n')
    admin_desc.add_run('• 월간 임계값: 월간 성공률 기준\n')
    admin_desc.add_run('• 반기 임계값: 반기 성공률 기준\n')
    admin_desc.add_run('• 연간 임계값: 연간 성공률 기준\n')
    admin_desc.add_run('• 성공 아이콘/색상: 성공 상태 표시\n')
    admin_desc.add_run('• 경고 아이콘/색상: 경고 상태 표시\n\n')

    admin_desc.add_run('5. 대시보드 표시 여부:\n').bold = True
    admin_desc.add_run('해당 Job ID의 데이터를 대시보드에 표시할지 여부를 설정합니다.\n\n')

    admin_desc.add_run('6. 컬러 팔레트:\n').bold = True
    admin_desc.add_run('색상 선택 시 사용할 수 있는 미리 정의된 색상 팔레트입니다.\n\n')

    admin_desc.add_run('7. 작업 버튼:\n').bold = True
    admin_desc.add_run('• 행 추가: 새로운 Job ID 설정 행을 추가합니다.\n')
    admin_desc.add_run('• 기본 설정 저장: 현재 설정을 저장합니다.\n')
    admin_desc.add_run('• 설정 내보내기: 설정을 JSON 파일로 내보냅니다.\n')
    admin_desc.add_run('• 설정 가져오기: JSON 파일로부터 설정을 불러옵니다.\n\n')

    admin_desc.add_run('2.7.2 사용자 인터랙션\n').bold = True
    admin_desc.add_run('1. Job ID 선택: 설정할 Job ID의 행을 선택합니다.\n')
    admin_desc.add_run('2. 임계값 설정: 각 필드에 적절한 임계값을 입력합니다.\n')
    admin_desc.add_run('3. 설정 저장: [기본 설정 저장] 버튼을 클릭하여 변경사항을 저장합니다.\n')
    admin_desc.add_run('4. 설정 관리: 내보내기/가져오기 기능으로 설정을 백업/복원합니다.\n\n')

    admin_desc.add_run('2.7.3 실제 데이터 예시\n').bold = True
    admin_desc.add_run('• CD101 설정: 연속 실패 3회, 성공률 80%, 빨강/노랑/초록 색상\n')
    admin_desc.add_run('• CD102 설정: 연속 실패 2회, 성공률 75%, 사용자 정의 색상\n\n')

    # 기본 설정 탭
    doc.add_paragraph('2.7.4.1 기본 설정 탭', style='OpH3')

    # 기본 설정 탭 스크린샷 삽입
    try:
        doc.add_picture('d:/dev/msys/scrennshot/admin_basic_screenshot.PNG', width=Inches(6))
        doc.add_paragraph('그림 2.13: 관리자 설정 - 기본 설정 탭 화면', style='OpH3')
    except:
        doc.add_paragraph('[스크린샷: admin_basic_screenshot.PNG를 찾을 수 없습니다]', style='OpH3')

    doc.add_paragraph('', style='OpH3')  # 빈 줄

    # 기본 설정 탭 설명
    basic_desc = doc.add_paragraph()
    basic_desc.add_run('2.7.4.1.1 기본 설정 탭 카드 및 변수 설명\n').bold = True
    basic_desc.add_run('기본 설정 탭에서 설정하는 각 변수의 의미와 사용법:\n\n')

    basic_desc.add_run('• Job ID (cd): 모니터링할 데이터 수집 작업의 고유 식별자\n').bold = True
    basic_desc.add_run('  - 형식: CD + 3자리 숫자 (예: CD101, CD102)\n')
    basic_desc.add_run('  - 용도: 대시보드와 각종 모니터링에서 작업을 구분하는 키\n\n')

    basic_desc.add_run('• Job 이름 (cd_nm): Job ID에 대한 사람이 읽기 쉬운 이름\n').bold = True
    basic_desc.add_run('  - 예시: "기상청 예보 데이터", "실시간 기상 데이터"\n')
    basic_desc.add_run('  - 용도: UI 표시 및 사용자 이해를 위한 레이블\n\n')

    basic_desc.add_run('• 연속 실패 임계값 (cnn_failr_thrs_val): 기본값 3회\n').bold = True
    basic_desc.add_run('  - 의미: 이 횟수만큼 연속 실패 시 위험 상태로 표시\n')
    basic_desc.add_run('  - 영향: 대시보드 카드의 상태 결정 (빨강 표시)\n\n')

    basic_desc.add_run('• 연속 경고 임계값 (cnn_warn_thrs_val): 기본값 2회\n').bold = True
    basic_desc.add_run('  - 의미: 이 횟수만큼 연속 실패 시 경고 상태로 표시\n')
    basic_desc.add_run('  - 영향: 대시보드 카드의 상태 결정 (노랑 표시)\n\n')

    basic_desc.add_run('• 일간 성공률 임계값 (dly_sucs_rt_thrs_val): 기본값 80%\n').bold = True
    basic_desc.add_run('  - 의미: 일별 성공률이 이 값 미만이면 위험 상태\n')
    basic_desc.add_run('  - 계산: (일별 성공 건수 ÷ 일별 총 건수) × 100\n\n')

    basic_desc.add_run('• 주간 성공률 임계값 (dd7_sucs_rt_thrs_val): 기본값 75%\n').bold = True
    basic_desc.add_run('  - 의미: 7일간 성공률이 이 값 미만이면 위험 상태\n\n')

    basic_desc.add_run('• 월간 성공률 임계값 (mthl_sucs_rt_thrs_val): 기본값 70%\n').bold = True
    basic_desc.add_run('  - 의미: 월별 성공률이 이 값 미만이면 위험 상태\n\n')

    basic_desc.add_run('• 상태별 아이콘/색상: 성공/경고/실패 상태 표시 설정\n').bold = True
    basic_desc.add_run('  - 성공: ✅ 아이콘, #008000 색상\n')
    basic_desc.add_run('  - 경고: ⚠️ 아이콘, #FFA500 색상\n')
    basic_desc.add_run('  - 실패: ❌ 아이콘, #FF0000 색상\n\n')

    basic_desc.add_run('• 대시보드 표시 여부 (chrt_dsp_yn): 해당 Job 표시 여부\n').bold = True
    basic_desc.add_run('  - true: 대시보드에 표시\n')
    basic_desc.add_run('  - false: 대시보드에서 숨김\n\n')

    # 수집 스케줄 설정 탭
    doc.add_paragraph('2.7.4.2 수집 스케줄 설정 탭', style='OpH3')

    # 수집 스케줄 설정 탭 스크린샷 삽입
    try:
        doc.add_picture('d:/dev/msys/scrennshot/admin_schedule_screenshot.PNG', width=Inches(6))
        doc.add_paragraph('그림 2.14: 관리자 설정 - 수집 스케줄 설정 탭 화면', style='OpH3')
    except:
        doc.add_paragraph('[스크린샷: admin_schedule_screenshot.PNG를 찾을 수 없습니다]', style='OpH3')

    doc.add_paragraph('', style='OpH3')  # 빈 줄

    # 수집 스케줄 설정 탭 설명
    doc.add_paragraph('2.7.4.2.1 수집 스케줄 설정 탭 카드 및 변수 설명', style='OpH3')
    doc.add_paragraph('수집 스케줄 설정 탭의 각 설정 항목 설명:', style='OpH3')
    doc.add_paragraph('', style='OpH3')  # 빈 줄

    doc.add_paragraph('• 그룹 최소 개수 (grp_min_cnt): 기본값 3', style='OpH3')
    doc.add_paragraph('  - 의미: 그룹화할 최소 Job 개수', style='OpH3')
    doc.add_paragraph('  - 영향: 카드 요약에서 그룹 생성 조건', style='OpH3')
    doc.add_paragraph('', style='OpH3')  # 빈 줄

    doc.add_paragraph('• 그룹 외곽선 스타일 (grp_brdr_styl): 기본값 solid', style='OpH3')
    doc.add_paragraph('  - 옵션: none, solid, dashed, dotted, double', style='OpH3')
    doc.add_paragraph('  - 의미: 그룹 카드의 테두리 스타일', style='OpH3')
    doc.add_paragraph('', style='OpH3')  # 빈 줄

    doc.add_paragraph('• 그룹 색상 기준 (grp_colr_crtr): prgr/succ', style='OpH3')
    doc.add_paragraph('  - prgr: 진행률 기준 색상', style='OpH3')
    doc.add_paragraph('  - succ: 성공률 기준 색상', style='OpH3')
    doc.add_paragraph('', style='OpH3')  # 빈 줄

    doc.add_paragraph('• 진행률 임계값: prgs_rt_red_thrsval(30%), prgs_rt_org_thrsval(60%)', style='OpH3')
    doc.add_paragraph('  - 문제점/경고 기준값 설정', style='OpH3')
    doc.add_paragraph('', style='OpH3')  # 빈 줄

    doc.add_paragraph('• 성공률 임계값: succ_rt_red_thrsval(30%), succ_rt_org_thrsval(60%)', style='OpH3')
    doc.add_paragraph('  - 문제점/경고 기준값 설정', style='OpH3')
    doc.add_paragraph('', style='OpH3')  # 빈 줄

    doc.add_paragraph('• 그룹 아이콘: grp_prgs_icon_id, grp_sucs_icon_id', style='OpH3')
    doc.add_paragraph('  - 그룹별 대표 아이콘 설정', style='OpH3')
    doc.add_paragraph('', style='OpH3')  # 빈 줄

    # 아이콘 관리 탭
    doc.add_paragraph('2.7.4.3 아이콘 관리 탭', style='OpH3')

    # 아이콘 관리 탭 스크린샷 삽입
    try:
        doc.add_picture('d:/dev/msys/scrennshot/admin_icon_screenshot.PNG', width=Inches(6))
        doc.add_paragraph('그림 2.15: 관리자 설정 - 아이콘 관리 탭 화면', style='OpH3')
    except:
        doc.add_paragraph('[스크린샷: admin_icon_screenshot.PNG를 찾을 수 없습니다]', style='OpH3')

    doc.add_paragraph('', style='OpH3')  # 빈 줄

    # 아이콘 관리 탭 설명
    doc.add_paragraph('2.7.4.3.1 아이콘 관리 탭 카드 및 변수 설명', style='OpH3')
    doc.add_paragraph('아이콘 관리 탭의 각 필드 설명:', style='OpH3')
    doc.add_paragraph('', style='OpH3')  # 빈 줄

    doc.add_paragraph('• 아이콘 코드 (icon_cd): 이모지 또는 텍스트 심볼', style='OpH3')
    doc.add_paragraph('  - 예: ✅, ❌, ⚠️, 🔄', style='OpH3')
    doc.add_paragraph('', style='OpH3')  # 빈 줄

    doc.add_paragraph('• 아이콘 이름 (icon_nm): 아이콘의 의미 설명', style='OpH3')
    doc.add_paragraph('  - 예: "성공", "실패", "경고", "처리중"', style='OpH3')
    doc.add_paragraph('', style='OpH3')  # 빈 줄

    doc.add_paragraph('• 아이콘 설명 (icon_expl): 상세 용도 설명', style='OpH3')
    doc.add_paragraph('  - 아이콘의 사용 예시와 목적', style='OpH3')
    doc.add_paragraph('', style='OpH3')  # 빈 줄

    doc.add_paragraph('• 표시 여부 (icon_dsp_yn): Y/N', style='OpH3')
    doc.add_paragraph('  - Y: 아이콘 사용 활성화', style='OpH3')
    doc.add_paragraph('  - N: 아이콘 사용 비활성화', style='OpH3')
    doc.add_paragraph('', style='OpH3')  # 빈 줄

    # 차트/시각화 설정 탭
    doc.add_paragraph('2.7.4.4 차트/시각화 설정 탭', style='OpH3')

    # 차트/시각화 설정 탭 스크린샷 삽입
    try:
        doc.add_picture('d:/dev/msys/scrennshot/admin_chart_screenshot.PNG', width=Inches(6))
        doc.add_paragraph('그림 2.16: 관리자 설정 - 차트/시각화 설정 탭 화면', style='OpH3')
    except:
        doc.add_paragraph('[스크린샷: admin_chart_screenshot.PNG를 찾을 수 없습니다]', style='OpH3')

    doc.add_paragraph('', style='OpH3')  # 빈 줄

    # 차트/시각화 설정 탭 설명
    doc.add_paragraph('2.7.4.4.1 차트/시각화 설정 탭 카드 및 변수 설명', style='OpH3')
    doc.add_paragraph('차트 설정 탭의 각 설정 항목:', style='OpH3')
    doc.add_paragraph('', style='OpH3')  # 빈 줄

    doc.add_paragraph('• 차트 색상 (chrt_colr): 16진수 색상 코드', style='OpH3')
    doc.add_paragraph('  - 용도: 차트에서 해당 Job의 선/바 색상', style='OpH3')
    doc.add_paragraph('', style='OpH3')  # 빈 줄

    doc.add_paragraph('• 잔디 시작 색상 (grass_chrt_min_colr): #9be9a8', style='OpH3')
    doc.add_paragraph('  - 용도: 잔디 차트의 최소 활동 색상', style='OpH3')
    doc.add_paragraph('', style='OpH3')  # 빈 줄

    doc.add_paragraph('• 잔디 끝 색상 (grass_chrt_max_colr): #216e39', style='OpH3')
    doc.add_paragraph('  - 용도: 잔디 차트의 최대 활동 색상', style='OpH3')
    doc.add_paragraph('', style='OpH3')  # 빈 줄

    # 사용자 관리 탭
    doc.add_paragraph('2.7.4.5 사용자 관리 탭', style='OpH3')

    # 사용자 관리 탭 스크린샷 삽입
    try:
        doc.add_picture('d:/dev/msys/scrennshot/admin_user_screenshot.PNG', width=Inches(6))
        doc.add_paragraph('그림 2.17: 관리자 설정 - 사용자 관리 탭 화면', style='OpH3')
    except:
        doc.add_paragraph('[스크린샷: admin_user_screenshot.PNG를 찾을 수 없습니다]', style='OpH3')

    doc.add_paragraph('', style='OpH3')  # 빈 줄

    # 사용자 관리 탭 설명
    doc.add_paragraph('2.7.4.5.1 사용자 관리 탭 카드 및 변수 설명', style='OpH3')
    doc.add_paragraph('사용자 관리 탭의 각 필드 설명:', style='OpH3')
    doc.add_paragraph('', style='OpH3')  # 빈 줄

    doc.add_paragraph('• 사용자 ID (user_id): 영문/숫자 조합 식별자', style='OpH3')
    doc.add_paragraph('  - 시스템 로그인 및 식별용', style='OpH3')
    doc.add_paragraph('', style='OpH3')  # 빈 줄

    doc.add_paragraph('• 계정 상태 (acc_sts): PENDING/APPROVED/REJECTED/SUSPENDED', style='OpH3')
    doc.add_paragraph('  - PENDING: 승인 대기 상태', style='OpH3')
    doc.add_paragraph('  - APPROVED: 정상 사용 가능', style='OpH3')
    doc.add_paragraph('  - REJECTED: 승인 거부', style='OpH3')
    doc.add_paragraph('  - SUSPENDED: 일시 정지', style='OpH3')
    doc.add_paragraph('', style='OpH3')  # 빈 줄

    doc.add_paragraph('• 메뉴 권한: dashboard, data_analysis, chart_analysis, mngr_sett', style='OpH3')
    doc.add_paragraph('  - 각 메뉴에 대한 접근 권한 설정', style='OpH3')
    doc.add_paragraph('', style='OpH3')  # 빈 줄

    # 데이터 접근 권한 탭
    doc.add_paragraph('2.7.4.6 데이터 접근 권한 탭', style='OpH3')

    # 데이터 접근 권한 탭 스크린샷 삽입
    try:
        doc.add_picture('d:/dev/msys/scrennshot/admin_permission_screenshot.PNG', width=Inches(6))
        doc.add_paragraph('그림 2.18: 관리자 설정 - 데이터 접근 권한 탭 화면', style='OpH3')
    except:
        doc.add_paragraph('[스크린샷: admin_permission_screenshot.PNG를 찾을 수 없습니다]', style='OpH3')

    doc.add_paragraph('', style='OpH3')  # 빈 줄

    # 데이터 접근 권한 탭 설명
    doc.add_paragraph('2.7.4.6.1 데이터 접근 권한 탭 카드 및 변수 설명', style='OpH3')
    doc.add_paragraph('데이터 접근 권한 설정 항목:', style='OpH3')
    doc.add_paragraph('', style='OpH3')  # 빈 줄

    doc.add_paragraph('• Job ID (job_id): CD로 시작하는 Job 식별자', style='OpH3')
    doc.add_paragraph('  - 접근을 허용할 데이터 수집 작업', style='OpH3')
    doc.add_paragraph('', style='OpH3')  # 빈 줄

    doc.add_paragraph('• 접근 권한 (perm_yn): true/false', style='OpH3')
    doc.add_paragraph('  - true: 데이터 접근 허용', style='OpH3')
    doc.add_paragraph('  - false: 데이터 접근 거부', style='OpH3')
    doc.add_paragraph('', style='OpH3')  # 빈 줄

    # 통계 탭
    doc.add_paragraph('2.7.4.7 통계 탭', style='OpH3')

    # 통계 탭 스크린샷 삽입 (일별 현황)
    try:
        doc.add_picture('d:/dev/msys/scrennshot/admin_stats_daily_screenshot.PNG', width=Inches(6))
        doc.add_paragraph('그림 2.19: 관리자 설정 - 통계 탭 (일별 현황) 화면', style='OpH3')
    except:
        doc.add_paragraph('[스크린샷: admin_stats_daily_screenshot.PNG를 찾을 수 없습니다]', style='OpH3')

    doc.add_paragraph('', style='OpH3')  # 빈 줄

    try:
        doc.add_picture('d:/dev/msys/scrennshot/admin_stats_weekly_screenshot.PNG', width=Inches(6))
        doc.add_paragraph('그림 2.20: 관리자 설정 - 통계 탭 (주별 현황) 화면', style='OpH3')
    except:
        doc.add_paragraph('[스크린샷: admin_stats_weekly_screenshot.PNG를 찾을 수 없습니다]', style='OpH3')

    doc.add_paragraph('', style='OpH3')  # 빈 줄

    try:
        doc.add_picture('d:/dev/msys/scrennshot/admin_stats_comparison_screenshot.PNG', width=Inches(6))
        doc.add_paragraph('그림 2.21: 관리자 설정 - 통계 탭 (비교 현황) 화면', style='OpH3')
    except:
        doc.add_paragraph('[스크린샷: admin_stats_comparison_screenshot.PNG를 찾을 수 없습니다]', style='OpH3')

    doc.add_paragraph('', style='OpH3')  # 빈 줄

    # 통계 탭 설명
    doc.add_paragraph('2.7.4.7.1 통계 탭 카드 및 변수 설명', style='OpH3')
    doc.add_paragraph('통계 분석에 사용되는 필터 및 표시 옵션:', style='OpH3')
    doc.add_paragraph('', style='OpH3')  # 빈 줄

    doc.add_paragraph('• 연도 선택 (year): 현재 연도 ± 2년 범위', style='OpH3')
    doc.add_paragraph('  - 분석할 기간의 기준 연도', style='OpH3')
    doc.add_paragraph('', style='OpH3')  # 빈 줄

    doc.add_paragraph('• 메뉴 선택 (menu_nm): all 또는 특정 메뉴', style='OpH3')
    doc.add_paragraph('  - all: 전체 메뉴 통계', style='OpH3')
    doc.add_paragraph('  - 특정 메뉴: 개별 메뉴 통계', style='OpH3')
    doc.add_paragraph('', style='OpH3')  # 빈 줄

    doc.add_paragraph('• 날짜 범위: start_date, end_date (YYYY-MM-DD)', style='OpH3')
    doc.add_paragraph('  - 일별 현황용 기간 설정', style='OpH3')
    doc.add_paragraph('', style='OpH3')  # 빈 줄

    admin_desc.add_run('2.7.4 관리자 설정 하위 탭 상세 설명\n').bold = True

    # 수집 스케줄 설정 탭
    admin_desc.add_run('2.7.4.1 수집 스케줄 설정 탭\n').bold = True
    admin_desc.add_run('데이터 수집 작업의 스케줄을 설정하고 관리하는 탭입니다.\n\n')

    admin_desc.add_run('1. 탭 기능 개요:\n').bold = True
    admin_desc.add_run('• Job별 자동 실행 주기 설정 (cron 표현식 기반)\n')
    admin_desc.add_run('• 수집 작업 활성화/비활성화 제어\n')
    admin_desc.add_run('• 실시간 스케줄 상태 모니터링\n')
    admin_desc.add_run('• 다음 실행 예정 시간 예측 표시\n\n')

    admin_desc.add_run('2. 주요 설정 항목:\n').bold = True
    admin_desc.add_run('• 그룹 최소 개수: 그룹화할 최소 Job 수 (기본값: 3)\n')
    admin_desc.add_run('• 그룹 외곽선 스타일: 카드 그룹 테두리 스타일 (solid/dashed/dotted/double)\n')
    admin_desc.add_run('• 그룹 색상 기준: 그룹 색상 결정 기준 (진행률/성공률)\n')
    admin_desc.add_run('• 진행률 임계값: 진행률 경고/문제점 기준값 (%)\n')
    admin_desc.add_run('• 성공률 임계값: 성공률 경고/문제점 기준값 (%)\n')
    admin_desc.add_run('• 그룹 아이콘: 그룹별 대표 아이콘 설정\n\n')

    admin_desc.add_run('3. 상태별 표시 설정:\n').bold = True
    admin_desc.add_run('• 성공 상태: 정상 완료된 작업 표시 (아이콘, 배경색, 글자색)\n')
    admin_desc.add_run('• 실패 상태: 오류로 실패한 작업 표시\n')
    admin_desc.add_run('• 진행중 상태: 현재 실행 중인 작업 표시\n')
    admin_desc.add_run('• 미수집 상태: 예정 시간 초과된 작업 표시\n')
    admin_desc.add_run('• 예정 상태: 아직 실행 예정인 작업 표시\n\n')

    admin_desc.add_run('4. 사용자 인터랙션:\n').bold = True
    admin_desc.add_run('• 설정 값 변경 후 [스케줄 표시 설정 저장] 버튼 클릭\n')
    admin_desc.add_run('• 실시간으로 카드 요약 화면에 반영\n')
    admin_desc.add_run('• 색상 팔레트에서 색상 선택 가능\n\n')

    admin_desc.add_run('5. 실제 데이터 예시:\n').bold = True
    admin_desc.add_run('• 그룹 최소 개수: 3 (3개 이상 Job이 있을 때 그룹화)\n')
    admin_desc.add_run('• 진행률 임계값: 문제점 30%, 경고 60%\n')
    admin_desc.add_run('• 성공 상태: 녹색 배경 (#EBF8FF), 파란색 글자 (#3182CE)\n\n')

    # 아이콘 관리 탭
    admin_desc.add_run('2.7.4.2 아이콘 관리 탭\n').bold = True
    admin_desc.add_run('시스템에서 사용하는 아이콘을 관리하는 탭입니다.\n\n')

    admin_desc.add_run('1. 탭 기능 개요:\n').bold = True
    admin_desc.add_run('• 시스템 아이콘의 중앙 집중 관리\n')
    admin_desc.add_run('• 이모지 기반 아이콘 코드 사용\n')
    admin_desc.add_run('• 아이콘 표시 여부 제어\n')
    admin_desc.add_run('• CSV 기반 일괄 관리 지원\n\n')

    admin_desc.add_run('2. 아이콘 추가 폼:\n').bold = True
    admin_desc.add_run('• 아이콘 코드: 이모지 또는 텍스트 코드 입력\n')
    admin_desc.add_run('• 아이콘 이름: 아이콘의 의미 설명\n')
    admin_desc.add_run('• 설명: 상세한 아이콘 용도 설명\n')
    admin_desc.add_run('• 표시 여부: 아이콘 사용 활성화/비활성화\n\n')

    admin_desc.add_run('3. 아이콘 관리 테이블:\n').bold = True
    admin_desc.add_run('• ID: 고유 식별자\n')
    admin_desc.add_run('• 아이콘: 실제 아이콘 표시\n')
    admin_desc.add_run('• 이름: 아이콘 이름\n')
    admin_desc.add_run('• 설명: 상세 설명\n')
    admin_desc.add_run('• 표시 여부: 사용 상태\n')
    admin_desc.add_run('• 작업: 수정/삭제 버튼\n\n')

    admin_desc.add_run('4. 일괄 관리 기능:\n').bold = True
    admin_desc.add_run('• CSV 내보내기: 모든 아이콘 정보를 파일로 저장\n')
    admin_desc.add_run('• CSV 가져오기: 파일로부터 아이콘 정보 일괄 등록\n')
    admin_desc.add_run('• 데이터 검증: 중복 코드 및 형식 검증\n\n')

    admin_desc.add_run('5. 실제 데이터 예시:\n').bold = True
    admin_desc.add_run('• ✅ 성공: 정상 완료 상태 표시\n')
    admin_desc.add_run('• ❌ 실패: 오류 발생 상태 표시\n')
    admin_desc.add_run('• ⚠️ 경고: 주의 필요 상태 표시\n')
    admin_desc.add_run('• 🔄 처리중: 작업 진행 중 표시\n\n')

    # 차트/시각화 설정 탭
    admin_desc.add_run('2.7.4.3 차트/시각화 설정 탭\n').bold = True
    admin_desc.add_run('차트 및 시각화 요소의 설정을 관리하는 탭입니다.\n\n')

    admin_desc.add_run('1. 탭 기능 개요:\n').bold = True
    admin_desc.add_run('• Job별 차트 색상 및 잔디 차트 색상 설정\n')
    admin_desc.add_run('• 시각화 요소의 색상 테마 관리\n')
    admin_desc.add_run('• 차트 표시 옵션 제어\n')
    admin_desc.add_run('• 데이터 포맷 및 레이아웃 설정\n\n')

    admin_desc.add_run('2. 차트 설정 항목:\n').bold = True
    admin_desc.add_run('• Job ID: 설정할 작업의 식별자\n')
    admin_desc.add_run('• Job 이름: 작업의 표시 이름 (읽기 전용)\n')
    admin_desc.add_run('• 차트 색상: 해당 Job의 차트 색상\n')
    admin_desc.add_run('• 잔디 시작 색상: 잔디 차트의 최소 색상\n')
    admin_desc.add_run('• 잔디 끝 색상: 잔디 차트의 최대 색상\n\n')

    admin_desc.add_run('3. 색상 선택 방식:\n').bold = True
    admin_desc.add_run('• 색상 팔레트에서 미리 정의된 색상 선택\n')
    admin_desc.add_run('• 사용자 정의 색상 직접 입력\n')
    admin_desc.add_run('• 색상 미리보기 기능\n\n')

    admin_desc.add_run('4. 실제 데이터 예시:\n').bold = True
    admin_desc.add_run('• CD101: 파란색 계열 (#007bff, #9be9a8, #216e39)\n')
    admin_desc.add_run('• CD102: 녹색 계열 (#28a745, #c3e6cb, #155724)\n')
    admin_desc.add_run('• CD103: 주황색 계열 (#fd7e14, #ffeaa7, #d68910)\n\n')

    # 사용자 관리 탭 스크린샷 삽입
    try:
        doc.add_picture('d:/dev/msys/scrennshot/admin_user_screenshot.PNG', width=Inches(6))
        doc.add_paragraph('그림 2.8: 관리자 설정 - 사용자 관리 탭 화면', style='OpH3')
    except:
        doc.add_paragraph('[스크린샷: admin_user_screenshot.PNG를 찾을 수 없습니다]', style='OpH3')

    doc.add_paragraph('', style='OpH3')  # 빈 줄

    admin_desc.add_run('2.7.4.4 사용자 관리 탭\n').bold = True
    admin_desc.add_run('시스템 사용자를 관리하는 탭입니다.\n\n')

    admin_desc.add_run('1. 탭 기능 개요:\n').bold = True
    admin_desc.add_run('• 사용자 계정 상태 관리 (승인/거부/정지)\n')
    admin_desc.add_run('• 메뉴 접근 권한 설정\n')
    admin_desc.add_run('• 비밀번호 초기화 및 재설정\n')
    admin_desc.add_run('• 사용자 활동 로그 조회\n\n')

    admin_desc.add_run('2. 사용자 목록 테이블:\n').bold = True
    admin_desc.add_run('• 사용자 ID: 고유 식별자\n')
    admin_desc.add_run('• 상태: 계정 상태 (PENDING/APPROVED/REJECTED/SUSPENDED)\n')
    admin_desc.add_run('• 가입일: 계정 생성 일시\n')
    admin_desc.add_run('• 메뉴 권한: 접근 가능한 메뉴 목록\n')
    admin_desc.add_run('• 작업: 승인/거부/비밀번호 초기화/삭제 버튼\n\n')

    admin_desc.add_run('3. 권한 설정 방식:\n').bold = True
    admin_desc.add_run('• 개별 메뉴 권한 체크박스로 제어\n')
    admin_desc.add_run('• mngr_sett 권한 부여 시 관리자 권한 획득\n')
    admin_desc.add_run('• 실시간 권한 적용\n\n')

    # 데이터 접근 권한 탭 스크린샷 삽입
    try:
        doc.add_picture('d:/dev/msys/scrennshot/admin_permission_screenshot.PNG', width=Inches(6))
        doc.add_paragraph('그림 2.9: 관리자 설정 - 데이터 접근 권한 탭 화면', style='OpH3')
    except:
        doc.add_paragraph('[스크린샷: admin_permission_screenshot.PNG를 찾을 수 없습니다]', style='OpH3')

    doc.add_paragraph('', style='OpH3')  # 빈 줄

    admin_desc.add_run('2.7.4.5 데이터 접근 권한 탭\n').bold = True
    admin_desc.add_run('사용자별 데이터 접근 권한을 설정하는 탭입니다.\n\n')

    admin_desc.add_run('1. 탭 기능 개요:\n').bold = True
    admin_desc.add_run('• Job ID별 데이터 접근 제어\n')
    admin_desc.add_run('• 민감한 데이터에 대한 보안 관리\n')
    admin_desc.add_run('• 사용자별 데이터 가시성 설정\n')
    admin_desc.add_run('• 권한 변경 이력 추적\n\n')

    admin_desc.add_run('2. 권한 설정 인터페이스:\n').bold = True
    admin_desc.add_run('• 미할당 Job ID 목록: 권한 없는 Job들\n')
    admin_desc.add_run('• 할당된 Job ID 목록: 권한 있는 Job들\n')
    admin_desc.add_run('• 이동 버튼: Job ID 권한 부여/해제\n')
    admin_desc.add_run('• 일괄 이동: 전체 권한 부여/해제\n\n')

    # 통계 탭 스크린샷 삽입
    try:
        doc.add_picture('d:/dev/msys/scrennshot/admin_stats_daily_screenshot.PNG', width=Inches(6))
        doc.add_paragraph('그림 2.10: 관리자 설정 - 통계 탭 (일별 현황) 화면', style='OpH3')
    except:
        doc.add_paragraph('[스크린샷: admin_stats_daily_screenshot.PNG를 찾을 수 없습니다]', style='OpH3')

    doc.add_paragraph('', style='OpH3')  # 빈 줄

    try:
        doc.add_picture('d:/dev/msys/scrennshot/admin_stats_weekly_screenshot.PNG', width=Inches(6))
        doc.add_paragraph('그림 2.11: 관리자 설정 - 통계 탭 (주별 현황) 화면', style='OpH3')
    except:
        doc.add_paragraph('[스크린샷: admin_stats_weekly_screenshot.PNG를 찾을 수 없습니다]', style='OpH3')

    doc.add_paragraph('', style='OpH3')  # 빈 줄

    try:
        doc.add_picture('d:/dev/msys/scrennshot/admin_stats_comparison_screenshot.PNG', width=Inches(6))
        doc.add_paragraph('그림 2.12: 관리자 설정 - 통계 탭 (비교 현황) 화면', style='OpH3')
    except:
        doc.add_paragraph('[스크린샷: admin_stats_comparison_screenshot.PNG를 찾을 수 없습니다]', style='OpH3')

    doc.add_paragraph('', style='OpH3')  # 빈 줄

    admin_desc.add_run('2.7.4.6 통계 탭\n').bold = True
    admin_desc.add_run('시스템 사용 통계를 분석하는 탭들입니다.\n\n')

    admin_desc.add_run('• 일별 현황: 일별 메뉴 접근 통계\n').bold = True
    admin_desc.add_run('  - 각 메뉴의 일별 클릭 수\n')
    admin_desc.add_run('  - 사용자별 활동량 분석\n')
    admin_desc.add_run('  - 시간대별 사용 패턴\n\n')

    admin_desc.add_run('• 주별/월별 현황: 장기간 사용 통계\n').bold = True
    admin_desc.add_run('  - 주별/월별 메뉴 사용량 추이\n')
    admin_desc.add_run('  - 사용자 유지율 분석\n')
    admin_desc.add_run('  - 계절별/월별 패턴 분석\n\n')

    admin_desc.add_run('• 비교 현황: 기간별 비교 분석\n').bold = True
    admin_desc.add_run('  - 올해 vs 작년 비교\n')
    admin_desc.add_run('  - 분기별 성장률 분석\n')
    admin_desc.add_run('  - 메뉴 사용량 변화 추이\n\n')

    # 사용자 관리 탭
    admin_desc.add_run('2.7.4.4 사용자 관리 탭\n').bold = True
    admin_desc.add_run('시스템 사용자를 관리하는 탭입니다.\n')
    admin_desc.add_run('• 사용자 목록 조회\n')
    admin_desc.add_run('• 사용자 권한 설정\n')
    admin_desc.add_run('• 계정 활성화/비활성화\n')
    admin_desc.add_run('• 비밀번호 초기화\n\n')

    # 데이터 접근 권한 탭
    admin_desc.add_run('2.7.4.5 데이터 접근 권한 탭\n').bold = True
    admin_desc.add_run('사용자별 데이터 접근 권한을 설정하는 탭입니다.\n')
    admin_desc.add_run('• Job ID별 권한 부여/해제\n')
    admin_desc.add_run('• 사용자 그룹 관리\n')
    admin_desc.add_run('• 권한 템플릿 적용\n')
    admin_desc.add_run('• 권한 변경 이력 조회\n\n')

    # 통계 탭들
    admin_desc.add_run('2.7.4.6 통계 탭\n').bold = True
    admin_desc.add_run('시스템 사용 통계를 분석하는 탭들입니다.\n\n')

    admin_desc.add_run('• 일별 현황: 일별 메뉴 접근 통계\n').bold = True
    admin_desc.add_run('  - 각 메뉴의 일별 클릭 수\n')
    admin_desc.add_run('  - 사용자별 활동량\n')
    admin_desc.add_run('  - 시간대별 사용 패턴\n\n')

    admin_desc.add_run('• 주별/월별 현황: 장기간 사용 통계\n').bold = True
    admin_desc.add_run('  - 주별/월별 메뉴 사용량 추이\n')
    admin_desc.add_run('  - 사용자 유지율 분석\n')
    admin_desc.add_run('  - 계절별/월별 패턴 분석\n\n')

    admin_desc.add_run('• 비교 현황: 기간별 비교 분석\n').bold = True
    admin_desc.add_run('  - 올해 vs 작년 비교\n')
    admin_desc.add_run('  - 분기별 성장률\n')
    admin_desc.add_run('  - 메뉴 사용량 변화 추이\n\n')

    # 2.8 시스템 리소스 모니터링
    doc.add_paragraph('2.8 시스템 리소스 모니터링', style='OpH2')
    sys_monitor = doc.add_paragraph()
    sys_monitor.add_run('시스템 리소스 사용량을 정기적으로 모니터링합니다:\n')
    sys_monitor.add_run('• CPU 사용률: 80% 미만 유지\n')
    sys_monitor.add_run('• 메모리 사용률: 85% 미만 유지\n')
    sys_monitor.add_run('• 디스크 사용률: 90% 미만 유지\n\n')

    sys_monitor.add_run('2.8.1 리소스 모니터링 명령어\n').bold = True
    sys_monitor.add_run('CPU 및 메모리 사용량 확인:\n\n')
    sys_monitor.add_run('top\n').font.name = 'Courier New'
    sys_monitor.add_run('htop\n\n').font.name = 'Courier New'

    sys_monitor.add_run('디스크 사용량 확인:\n\n')
    sys_monitor.add_run('df -h\n').font.name = 'Courier New'
    sys_monitor.add_run('du -sh /home/msys/app\n\n').font.name = 'Courier New'

    # 2.9 수집 일정 모니터링
    doc.add_paragraph('2.9 수집 일정 모니터링', style='OpH2')
    schedule_monitor = doc.add_paragraph()
    schedule_monitor.add_run('데이터 수집 작업의 스케줄과 실행 현황을 모니터링하는 메뉴입니다.\n\n')

    schedule_monitor.add_run('2.9.1 화면 구성 요소 상세 설명\n').bold = True
    schedule_monitor.add_run('1. 기간 선택 컨트롤:\n').bold = True
    schedule_monitor.add_run('• 주간/월간 뷰 전환: 탭 메뉴로 기간 단위 선택\n')
    schedule_monitor.add_run('• 날짜 범위 설정: 조회할 기간의 시작일과 종료일 지정\n\n')

    schedule_monitor.add_run('2. 스케줄 테이블:\n').bold = True
    schedule_monitor.add_run('• Job ID: 데이터 수집 작업의 고유 식별자\n')
    schedule_monitor.add_run('• Job 이름: 작업의 표시 이름\n')
    schedule_monitor.add_run('• 스케줄 시간: 예정된 실행 시간\n')
    schedule_monitor.add_run('• 상태: 현재 실행 상태 (예정/수집중/성공/실패/미수집)\n')
    schedule_monitor.add_run('• 실행 결과: 실제 실행 결과 및 소요 시간\n\n')

    schedule_monitor.add_run('3. 상태 표시 색상 코드:\n').bold = True
    schedule_monitor.add_run('• 회색: 예정 - 아직 실행되지 않은 작업\n')
    schedule_monitor.add_run('• 파랑: 수집중 - 현재 실행 중인 작업\n')
    schedule_monitor.add_run('• 녹색: 성공 - 정상적으로 완료된 작업\n')
    schedule_monitor.add_run('• 빨강: 실패 - 오류로 실패한 작업\n')
    schedule_monitor.add_run('• 노랑: 미수집 - 예정된 시간에 실행되지 않은 작업\n\n')

    schedule_monitor.add_run('2.9.2 데이터 구조 및 의미\n').bold = True
    schedule_monitor.add_run('• 스케줄 데이터: crontab 형식의 실행 주기 정보\n')
    schedule_monitor.add_run('• 실행 이력: 각 Job의 과거 실행 결과 기록\n')
    schedule_monitor.add_run('• 상태 계산: 실시간 실행 상태와 예정 상태의 조합\n\n')

    schedule_monitor.add_run('2.9.3 사용자 인터랙션\n').bold = True
    schedule_monitor.add_run('1. 기간 선택: 주간 또는 월간 탭 클릭\n')
    schedule_monitor.add_run('2. 날짜 설정: 시작일과 종료일 입력\n')
    schedule_monitor.add_run('3. 조회 실행: [조회] 버튼으로 데이터 새로고침\n')
    schedule_monitor.add_run('4. 상세 정보: Job 행 클릭으로 실행 상세 정보 확인\n\n')

    schedule_monitor.add_run('2.9.4 실제 데이터 예시\n').bold = True
    schedule_monitor.add_run('• 정상 스케줄: CD101 (06:00 예정 → 06:02 성공, 소요 2분)\n')
    schedule_monitor.add_run('• 지연 스케줄: CD102 (07:00 예정 → 07:15 성공, 소요 15분)\n')
    schedule_monitor.add_run('• 실패 스케줄: CD103 (08:00 예정 → 실패, 네트워크 오류)\n\n')

    # 2.10 카드 요약 모니터링
    doc.add_paragraph('2.10 카드 요약 모니터링', style='OpH2')
    card_monitor = doc.add_paragraph()
    card_monitor.add_run('오늘의 데이터 수집 현황을 그룹별로 요약하여 카드 형태로 표시하는 메뉴입니다. 각 카드는 Job ID의 첫 3자리로 그룹화되어 표시되며, 실시간으로 오늘의 수집 상태를 집계하여 보여줍니다.\n\n')

    card_monitor.add_run('2.10.1 화면 구성 요소 상세 설명\n').bold = True
    card_monitor.add_run('1. 표시 옵션 컨트롤:\n').bold = True
    card_monitor.add_run('• 명칭: Job의 한글명만 표시 (예: "기상청 예보 데이터")\n')
    card_monitor.add_run('• 코드: Job ID만 표시 (예: "CD101")\n')
    card_monitor.add_run('• 명칭+코드: 둘 다 표시 (예: "CD101 (기상청 예보 데이터)")\n\n')

    card_monitor.add_run('2. 그룹별 카드 배치:\n').bold = True
    card_monitor.add_run('• CD100 그룹: CD101-CD199 범위의 Job들 (기상청, 기상정보 등)\n')
    card_monitor.add_run('• CD200 그룹: CD201-CD299 범위의 Job들 (외부 API, 타 시스템 등)\n')
    card_monitor.add_run('• CD300 그룹: CD301-CD399 범위의 Job들 (추가 확장 그룹)\n')
    card_monitor.add_run('• 기타 그룹: 분류되지 않은 Job들\n\n')

    card_monitor.add_run('3. 각 카드의 상태 표시 및 의미:\n').bold = True
    card_monitor.add_run('각 카드는 오늘의 수집 스케줄과 실행 결과를 실시간으로 집계하여 5가지 상태별 카운트를 표시합니다:\n\n')

    card_monitor.add_run('• 성공 (Success): status = \'성공\'인 오늘의 Job 수\n').bold = True
    card_monitor.add_run('  - 의미: 예정된 시간에 정상적으로 데이터 수집이 완료된 작업 수\n')
    card_monitor.add_run('  - 계산: 오늘 스케줄 중 status 필드가 \'성공\'으로 기록된 Job의 개수\n')
    card_monitor.add_run('  - 표시 색상: 녹색 배경\n')
    card_monitor.add_run('  - 세부 정보: 성공한 Job ID 목록과 실행 시간 (예: "CD101(06시)")\n\n')

    card_monitor.add_run('• 수집중 (Progress): status = \'수집중\'인 현재 실행 중인 Job 수\n').bold = True
    card_monitor.add_run('  - 의미: 현재 데이터 수집 작업이 진행 중인 상태\n')
    card_monitor.add_run('  - 계산: 오늘 스케줄 중 status 필드가 \'수집중\'인 Job의 개수\n')
    card_monitor.add_run('  - 표시 색상: 파랑 배경\n')
    card_monitor.add_run('  - 세부 정보: 진행 중인 Job ID 목록과 시작 시간\n\n')

    card_monitor.add_run('• 실패 (Fail): status = \'실패\'인 오류 발생 Job 수\n').bold = True
    card_monitor.add_run('  - 의미: 데이터 수집 중 오류가 발생하여 실패한 작업 수\n')
    card_monitor.add_run('  - 계산: 오늘 스케줄 중 status 필드가 \'실패\'인 Job의 개수\n')
    card_monitor.add_run('  - 표시 색상: 빨강 배경\n')
    card_monitor.add_run('  - 세부 정보: 실패한 Job ID 목록과 실패 원인 (가능한 경우)\n\n')

    card_monitor.add_run('• 미수집 (Uncollected): status = \'미수집\'인 누락된 Job 수\n').bold = True
    card_monitor.add_run('  - 의미: 예정된 수집 시간이 지났으나 아직 실행되지 않은 작업 수\n')
    card_monitor.add_run('  - 계산: 오늘 스케줄 중 예정 시간 경과 + status가 \'미수집\'인 Job의 개수\n')
    card_monitor.add_run('  - 표시 색상: 노랑 배경\n')
    card_monitor.add_run('  - 세부 정보: 미수집된 Job ID 목록과 예정 시간\n\n')

    card_monitor.add_run('• 예정 (Scheduled): status = \'예정\'인 앞으로 실행할 Job 수\n').bold = True
    card_monitor.add_run('  - 의미: 오늘 예정된 수집 시간에 아직 도달하지 않은 작업 수\n')
    card_monitor.add_run('  - 계산: 오늘 스케줄 중 예정 시간 미도달 + status가 \'예정\'인 Job의 개수\n')
    card_monitor.add_run('  - 표시 색상: 회색 배경\n')
    card_monitor.add_run('  - 세부 정보: 예정된 Job ID 목록과 실행 예정 시간\n\n')

    card_monitor.add_run('4. 카드 헤더 정보:\n').bold = True
    card_monitor.add_run('• 그룹명: CD100, CD200 등의 그룹 식별자\n')
    card_monitor.add_run('• 총 건수: 해당 그룹의 전체 Job 수 (성공 + 수집중 + 실패 + 미수집 + 예정)\n\n')

    card_monitor.add_run('2.10.2 데이터 구조 및 의미\n').bold = True
    card_monitor.add_run('• 그룹 분류: Job ID의 첫 3자리로 자동 그룹화 (CD100, CD200 등)\n')
    card_monitor.add_run('• 실시간 집계: 오늘의 모든 Job 실행 상태를 실시간으로 집계\n')
    card_monitor.add_run('• 권한 필터링: 사용자 권한에 따른 데이터 접근 제한 (mngr_sett 권한 제외)\n')
    card_monitor.add_run('• 시간 기준: Asia/Seoul 타임존 기준 오늘 00:00 ~ 현재 시간까지\n')
    card_monitor.add_run('• 데이터 소스: collection_schedule_service.get_schedule_and_history() 결과\n\n')

    card_monitor.add_run('2.10.3 사용자 인터랙션\n').bold = True
    card_monitor.add_run('1. 표시 옵션 선택: 라디오 버튼으로 Job 표시 방식 변경\n')
    card_monitor.add_run('2. 그룹별 현황 파악: 각 카드의 숫자와 색상으로 상태 파악\n')
    card_monitor.add_run('3. 세부 정보 확인: 각 상태별 Job ID 목록과 시간 정보 확인\n')
    card_monitor.add_run('4. 엑셀 양식 다운로드: 수집 요청서 양식 다운로드 버튼 사용\n')
    card_monitor.add_run('5. 실시간 업데이트: 페이지 방문 시 자동 새로고침\n\n')

    card_monitor.add_run('2.10.4 실제 데이터 예시\n').bold = True
    card_monitor.add_run('• CD100 그룹 카드:\n')
    card_monitor.add_run('  - 총 11건 (성공 8건, 실패 1건, 예정 2건)\n')
    card_monitor.add_run('  - 성공 Job: CD101(06시), CD102(07시), CD103(08시), CD104(09시)\n')
    card_monitor.add_run('  - 실패 Job: CD105(10시) - 네트워크 타임아웃\n')
    card_monitor.add_run('  - 예정 Job: CD106(14시), CD107(16시)\n\n')

    card_monitor.add_run('• CD200 그룹 카드:\n')
    card_monitor.add_run('  - 총 14건 (성공 12건, 수집중 1건, 미수집 1건)\n')
    card_monitor.add_run('  - 수집중 Job: CD201(현재 11시 시작)\n')
    card_monitor.add_run('  - 미수집 Job: CD202(10시 예정, 아직 실행 안됨)\n\n')

    card_monitor.add_run('2.10.5 상태별 색상 코드 및 조건\n').bold = True
    card_monitor.add_run('• 녹색 (●): 성공 - status == \'성공\'\n')
    card_monitor.add_run('• 파랑 (●): 수집중 - status == \'수집중\'\n')
    card_monitor.add_run('• 빨강 (●): 실패 - status == \'실패\'\n')
    card_monitor.add_run('• 노랑 (●): 미수집 - status == \'미수집\'\n')
    card_monitor.add_run('• 회색 (●): 예정 - status == \'예정\'\n\n')

    card_monitor.add_run('2.10.6 데이터 계산 방식\n').bold = True
    card_monitor.add_run('• 상태 판별: collection_schedule 테이블의 status 필드 값 기반\n')
    card_monitor.add_run('• 시간 비교: 현재 시간 vs 예정 실행 시간\n')
    card_monitor.add_run('• 그룹 집계: Job ID 패턴 (CD1xx → CD100)으로 그룹화\n')
    card_monitor.add_run('• 실시간성: 매 페이지 로드 시 최신 데이터 조회\n\n')

    # 3. 백업 및 복원
    doc.add_paragraph('3. 백업 및 복원', style='OpH1')
    backup = doc.add_paragraph()
    backup.add_run('3.1 자동 백업 설정\n').bold = True
    backup.add_run('crontab에 다음 작업을 추가:\n\n')
    backup.add_run('0 2 * * * pg_dump -U msys_user msys_db > /backup/msys_$(date +\\%Y\\%m\\%d).sql\n\n').font.name = 'Courier New'

    backup.add_run('3.2 수동 백업\n').bold = True
    backup.add_run('pg_dump -U msys_user -h localhost msys_db > backup.sql\n\n').font.name = 'Courier New'

    backup.add_run('3.3 복원 절차\n').bold = True
    backup.add_run('1. 서비스 중지\n')
    backup.add_run('2. 데이터베이스 드롭 및 재생성\n')
    backup.add_run('3. 백업 파일로부터 복원:\n\n')
    backup.add_run('psql -U msys_user msys_db < backup.sql\n\n').font.name = 'Courier New'

    # 4. 장애 대응
    doc.add_paragraph('4. 장애 대응', style='OpH1')
    incident = doc.add_paragraph()
    incident.add_run('4.1 장애 유형별 대응\n\n').bold = True

    incident.add_run('웹 서비스 장애:\n').bold = True
    incident.add_run('1. 프로세스 상태 확인\n')
    incident.add_run('2. 로그 파일 검토\n')
    incident.add_run('3. 서비스 재시작\n\n')

    incident.add_run('데이터베이스 장애:\n').bold = True
    incident.add_run('1. PostgreSQL 서비스 상태 확인\n')
    incident.add_run('2. 연결 설정 검토\n')
    incident.add_run('3. 데이터베이스 복구\n\n')

    # 5. 성능 최적화
    doc.add_paragraph('5. 성능 최적화', style='OpH1')
    performance = doc.add_paragraph()
    performance.add_run('5.1 데이터베이스 최적화\n').bold = True
    performance.add_run('• 인덱스 생성 및 유지\n')
    performance.add_run('• 쿼리 최적화\n')
    performance.add_run('• 테이블 파티셔닝\n\n')

    performance.add_run('5.2 캐싱 전략\n').bold = True
    performance.add_run('• Redis 캐시 서버 도입 고려\n')
    performance.add_run('• 정적 파일 캐싱\n\n')

    # 6. 인증 시스템
    doc.add_paragraph('6. 인증 시스템', style='OpH1')

    # 6.1 로그인 화면
    doc.add_paragraph('6.1 로그인 화면', style='OpH2')
    login_desc = doc.add_paragraph()
    login_desc.add_run('시스템에 접근하기 위한 사용자 인증 화면입니다.\n\n')

    # 로그인 스크린샷 삽입
    try:
        doc.add_picture('scrennshot/login_screenshot.PNG', width=Inches(6))
        doc.add_paragraph('그림 6.1: 로그인 화면', style='OpH3')
    except:
        doc.add_paragraph('[스크린샷: login_screenshot.PNG를 찾을 수 없습니다]', style='OpH3')

    doc.add_paragraph('', style='OpH3')  # 빈 줄

    login_elements = doc.add_paragraph()
    login_elements.add_run('6.1.1 화면 구성 요소\n').bold = True

    login_elements.add_run('1. 탭 메뉴:\n').bold = True
    login_elements.add_run('• 로그인: 사용자 인증을 위한 기본 탭\n')
    login_elements.add_run('• 회원가입: 신규 사용자 등록을 위한 탭\n\n')

    login_elements.add_run('2. 로그인 폼:\n').bold = True
    login_elements.add_run('• 사용자 ID: 시스템에 등록된 사용자 식별자 입력\n')
    login_elements.add_run('• 비밀번호: 사용자 인증을 위한 비밀번호 입력\n')
    login_elements.add_run('• 로그인 버튼: 인증 요청 실행\n')
    login_elements.add_run('• 비밀번호 초기화 요청: 비밀번호 분실 시 재설정 요청\n\n')

    login_elements.add_run('3. 시스템 안내:\n').bold = True
    login_elements.add_run('• 회원가입 후 관리자 승인이 필요하다는 안내 메시지\n')
    login_elements.add_run('• 문의사항 연락처 정보 표시\n\n')

    login_usage = doc.add_paragraph()
    login_usage.add_run('6.1.2 사용 방법\n').bold = True
    login_usage.add_run('1. 사용자 ID 입력: 시스템에 등록된 ID를 입력합니다.\n')
    login_usage.add_run('2. 비밀번호 입력: 해당 ID의 비밀번호를 입력합니다.\n')
    login_usage.add_run('3. 로그인 버튼 클릭: 인증을 요청합니다.\n')
    login_usage.add_run('4. 성공 시 대시보드로 자동 이동, 실패 시 오류 메시지 표시\n\n')

    # 6.2 회원가입 화면
    doc.add_paragraph('6.2 회원가입 화면', style='OpH2')
    register_desc = doc.add_paragraph()
    register_desc.add_run('신규 사용자가 시스템에 등록하기 위한 화면입니다.\n\n')

    # 회원가입 스크린샷 삽입
    try:
        doc.add_picture('scrennshot/register_screenshot.PNG', width=Inches(6))
        doc.add_paragraph('그림 6.2: 회원가입 화면', style='OpH3')
    except:
        doc.add_paragraph('[스크린샷: register_screenshot.PNG를 찾을 수 없습니다]', style='OpH3')

    doc.add_paragraph('', style='OpH3')  # 빈 줄

    register_elements = doc.add_paragraph()
    register_elements.add_run('6.2.1 화면 구성 요소\n').bold = True

    register_elements.add_run('1. 회원가입 폼:\n').bold = True
    register_elements.add_run('• 사용자 ID: 영문/숫자 조합의 고유 식별자\n')
    register_elements.add_run('• 비밀번호: 보안 정책을 만족하는 비밀번호\n')
    register_elements.add_run('• 비밀번호 확인: 비밀번호 재입력으로 검증\n')
    register_elements.add_run('• 가입 신청 버튼: 회원가입 요청 실행\n\n')

    register_elements.add_run('2. 비밀번호 정책 체크리스트:\n').bold = True
    register_elements.add_run('실시간으로 비밀번호가 정책에 맞는지 검증하여 표시:\n')
    register_elements.add_run('• 8자 이상: 최소 길이 요구사항\n')
    register_elements.add_run('• 특수문자 1개 이상 포함: 보안 강도 향상\n')
    register_elements.add_run('• 연속된 숫자 사용 불가: 123, 456 등 패턴 방지\n')
    register_elements.add_run('• 동일 숫자 반복 불가: 111, 222 등 반복 방지\n\n')

    register_elements.add_run('3. 실시간 검증:\n').bold = True
    register_elements.add_run('• 입력 시 정책 준수 여부 실시간 표시 (✅/❌)\n')
    register_elements.add_run('• 비밀번호 확인 일치 여부 검증\n')
    register_elements.add_run('• 모든 조건 만족 시에만 가입 신청 버튼 활성화\n\n')

    register_usage = doc.add_paragraph()
    register_usage.add_run('6.2.2 사용 방법\n').bold = True
    register_usage.add_run('1. 사용자 ID 입력: 영문/숫자 조합으로 원하는 ID 입력\n')
    register_usage.add_run('2. 비밀번호 입력: 정책을 만족하는 비밀번호 입력\n')
    register_usage.add_run('3. 비밀번호 확인: 동일한 비밀번호 재입력\n')
    register_usage.add_run('4. 정책 체크리스트 확인: 모든 항목이 ✅ 표시되는지 확인\n')
    register_usage.add_run('5. 가입 신청 버튼 클릭: 회원가입 요청 (관리자 승인 대기)\n\n')

    # 6.4 회원가입 승인 및 권한 부여 흐름
    doc.add_paragraph('6.4 회원가입 승인 및 권한 부여 흐름', style='OpH2')
    approval_flow = doc.add_paragraph()
    approval_flow.add_run('신규 사용자의 시스템 접근 권한을 부여하는 전체 프로세스입니다.\n\n')

    approval_flow.add_run('6.4.1 권한 부여 단계별 흐름\n').bold = True
    approval_flow.add_run('1. 사용자 회원가입 신청\n').bold = True
    approval_flow.add_run('   - 사용자 ID, 비밀번호 입력\n')
    approval_flow.add_run('   - DB에 PENDING 상태로 저장\n')
    approval_flow.add_run('   - 관리자에게 승인 요청 알림\n\n')

    approval_flow.add_run('2. 관리자 승인 처리\n').bold = True
    approval_flow.add_run('   - 관리자 설정 → 사용자 관리 탭 접근\n')
    approval_flow.add_run('   - PENDING 상태 사용자 목록 확인\n')
    approval_flow.add_run('   - 승인 버튼 클릭 → APPROVED 상태로 변경\n')
    approval_flow.add_run('   - 비밀번호를 사용자 ID와 동일하게 초기화 (보안 정책 준수를 위해)\n')
    approval_flow.add_run('   - 사용자에게 초기 비밀번호 안내\n\n')

    approval_flow.add_run('2.1 사용자 첫 로그인 및 비밀번호 변경 강제\n').bold = True
    approval_flow.add_run('   - 사용자가 초기 비밀번호(사용자 ID와 동일)로 로그인 시도\n')
    approval_flow.add_run('   - 시스템이 보안 정책 위반 감지하여 비밀번호 변경 강제\n')
    approval_flow.add_run('   - "비밀번호를 변경해야 합니다. 초기화된 비밀번호는 안전하지 않습니다." 메시지 표시\n')
    approval_flow.add_run('   - 비밀번호 변경 화면으로 자동 리디렉션\n\n')

    approval_flow.add_run('3. 메뉴 권한 설정\n').bold = True
    approval_flow.add_run('   - 기본 권한 자동 부여: dashboard, collection_schedule\n')
    approval_flow.add_run('   - 추가 권한 수동 설정: data_analysis, chart_analysis 등\n')
    approval_flow.add_run('   - 권한별 메뉴 접근 제어\n\n')

    approval_flow.add_run('4. 데이터 접근 권한 설정\n').bold = True
    approval_flow.add_run('   - Job ID별 데이터 접근 권한 부여\n')
    approval_flow.add_run('   - 데이터 보안 및 프라이버시 제어\n')
    approval_flow.add_run('   - 권한에 따른 데이터 필터링 적용\n\n')

    approval_flow.add_run('6.4.2 권한 종류 및 의미\n').bold = True
    approval_flow.add_run('• dashboard: 대시보드 메뉴 접근 권한\n')
    approval_flow.add_run('• collection_schedule: 수집 일정 조회 권한\n')
    approval_flow.add_run('• data_analysis: 데이터 분석 메뉴 접근 권한\n')
    approval_flow.add_run('• chart_analysis: 차트 분석 메뉴 접근 권한\n')
    approval_flow.add_run('• jandi: 잔디 현황 메뉴 접근 권한\n')
    approval_flow.add_run('• mapping: 매핑 메뉴 접근 권한\n')
    approval_flow.add_run('• data_spec: 데이터 명세서 메뉴 접근 권한\n')
    approval_flow.add_run('• mngr_sett: 관리자 설정 메뉴 접근 권한\n\n')

    approval_flow.add_run('6.4.3 권한 설정 예시\n').bold = True
    approval_flow.add_run('• 일반 사용자: dashboard, collection_schedule, data_analysis\n')
    approval_flow.add_run('• 분석 담당자: 위 권한 + chart_analysis, jandi\n')
    approval_flow.add_run('• 관리자: 모든 권한 + mngr_sett\n\n')

    # 6.3 비밀번호 변경 화면
    doc.add_paragraph('6.3 비밀번호 변경 화면', style='OpH2')
    pw_change_desc = doc.add_paragraph()
    pw_change_desc.add_run('로그인된 사용자가 비밀번호를 변경하기 위한 화면입니다.\n\n')

    # 비밀번호 변경 스크린샷 삽입
    try:
        doc.add_picture('scrennshot/change_password_screenshot.PNG', width=Inches(6))
        doc.add_paragraph('그림 6.3: 비밀번호 변경 화면', style='OpH3')
    except:
        doc.add_paragraph('[스크린샷: change_password_screenshot.PNG를 찾을 수 없습니다]', style='OpH3')

    doc.add_paragraph('', style='OpH3')  # 빈 줄

    pw_change_elements = doc.add_paragraph()
    pw_change_elements.add_run('6.3.1 화면 구성 요소\n').bold = True
    pw_change_elements.add_run('• 현재 비밀번호: 기존 비밀번호 확인용 입력 필드\n')
    pw_change_elements.add_run('• 새 비밀번호: 변경할 새 비밀번호 입력\n')
    pw_change_elements.add_run('• 새 비밀번호 확인: 새 비밀번호 재입력 검증\n')
    pw_change_elements.add_run('• 변경 버튼: 비밀번호 변경 실행\n\n')

    pw_change_usage = doc.add_paragraph()
    pw_change_usage.add_run('6.3.2 사용 방법\n').bold = True
    pw_change_usage.add_run('1. 현재 비밀번호 입력: 기존 비밀번호 확인\n')
    pw_change_usage.add_run('2. 새 비밀번호 입력: 새로운 비밀번호 입력\n')
    pw_change_usage.add_run('3. 새 비밀번호 확인: 동일한 새 비밀번호 재입력\n')
    pw_change_usage.add_run('4. 변경 버튼 클릭: 비밀번호 변경 완료\n\n')

    # 6.4 로그 관리
    doc.add_paragraph('7. 로그 관리', style='OpH1')

    # 6.1 로그 파일 위치 및 종류
    doc.add_paragraph('6.1 로그 파일 위치 및 종류', style='OpH2')
    log_files = doc.add_paragraph()
    log_files.add_run('6.1.1 로그 파일 종류\n').bold = True
    log_files.add_run('• 애플리케이션 로그 (app.log): 애플리케이션 실행 중 발생하는 이벤트 및 오류 기록\n')
    log_files.add_run('• 접근 로그 (access.log): 웹 서버 접근 기록\n')
    log_files.add_run('• 시스템 로그 (/var/log/syslog): 운영체제 수준의 시스템 이벤트\n\n')

    log_files.add_run('6.1.2 로그 파일 위치\n').bold = True
    log_files.add_run('• /home/msys/app/log/app.log\n')
    log_files.add_run('• /home/msys/app/log/access.log\n')
    log_files.add_run('• /var/log/syslog (시스템 로그)\n\n')

    # 6.2 로그 확인 방법
    doc.add_paragraph('6.2 로그 확인 방법', style='OpH2')
    log_check = doc.add_paragraph()
    log_check.add_run('6.2.1 실시간 로그 모니터링\n').bold = True
    log_check.add_run('로그 파일을 실시간으로 모니터링합니다:\n\n')
    log_check.add_run('tail -f /home/msys/app/log/app.log\n\n').font.name = 'Courier New'

    log_check.add_run('6.2.2 최근 로그 확인\n').bold = True
    log_check.add_run('최근 50줄의 로그를 확인합니다:\n\n')
    log_check.add_run('tail -50 /home/msys/app/log/app.log\n\n').font.name = 'Courier New'

    log_check.add_run('6.2.3 특정 시간대 로그 검색\n').bold = True
    log_check.add_run('특정 시간대의 로그를 검색합니다:\n\n')
    log_check.add_run('grep "2025-12-18" /home/msys/app/log/app.log\n\n').font.name = 'Courier New'

    log_check.add_run('6.2.4 오류 로그만 필터링\n').bold = True
    log_check.add_run('ERROR 레벨의 로그만 확인합니다:\n\n')
    log_check.add_run('grep "ERROR" /home/msys/app/log/app.log\n\n').font.name = 'Courier New'

    # 6.3 로그 데이터 해석
    doc.add_paragraph('6.3 로그 데이터 해석', style='OpH2')
    log_interpret = doc.add_paragraph()
    log_interpret.add_run('6.3.1 로그 레벨\n').bold = True
    log_interpret.add_run('• DEBUG: 상세한 디버깅 정보\n')
    log_interpret.add_run('• INFO: 일반적인 정보 메시지\n')
    log_interpret.add_run('• WARNING: 경고 상황\n')
    log_interpret.add_run('• ERROR: 오류 발생\n')
    log_interpret.add_run('• CRITICAL: 심각한 오류\n\n')

    log_interpret.add_run('6.3.2 일반적인 로그 메시지 예시\n').bold = True
    log_interpret.add_run('[2025-12-18 14:30:15] INFO: 애플리케이션 시작됨\n').font.name = 'Courier New'
    log_interpret.add_run('[2025-12-18 14:30:16] INFO: 데이터베이스 연결 성공\n').font.name = 'Courier New'
    log_interpret.add_run('[2025-12-18 14:35:22] ERROR: 데이터베이스 연결 실패 - Connection timeout\n').font.name = 'Courier New'
    log_interpret.add_run('[2025-12-18 14:35:23] WARNING: 메모리 사용률 85% 초과\n\n').font.name = 'Courier New'

    # 6.4 오류 발생 시 조치 방법
    doc.add_paragraph('6.4 오류 발생 시 조치 방법', style='OpH2')
    error_handling = doc.add_paragraph()
    error_handling.add_run('6.4.1 데이터베이스 연결 오류\n').bold = True
    error_handling.add_run('증상: "Connection timeout" 또는 "Connection refused" 오류\n')
    error_handling.add_run('조치:\n')
    error_handling.add_run('1. PostgreSQL 서비스 상태 확인: systemctl status postgresql\n')
    error_handling.add_run('2. 데이터베이스 서버 재시작: systemctl restart postgresql\n')
    error_handling.add_run('3. 네트워크 연결 확인: ping database_server\n')
    error_handling.add_run('4. .env 파일의 데이터베이스 설정 확인\n\n')

    error_handling.add_run('6.4.2 메모리 부족 오류\n').bold = True
    error_handling.add_run('증상: "MemoryError" 또는 메모리 사용률 90% 이상\n')
    error_handling.add_run('조치:\n')
    error_handling.add_run('1. 현재 메모리 사용량 확인: free -h\n')
    error_handling.add_run('2. 메모리 사용 프로세스 확인: ps aux --sort=-%mem | head\n')
    error_handling.add_run('3. 불필요한 프로세스 종료\n')
    error_handling.add_run('4. 서버 메모리 증설 고려\n\n')

    error_handling.add_run('6.4.3 디스크 공간 부족\n').bold = True
    error_handling.add_run('증상: "No space left on device" 오류\n')
    error_handling.add_run('조치:\n')
    error_handling.add_run('1. 디스크 사용량 확인: df -h\n')
    error_handling.add_run('2. 큰 파일 찾기: find / -type f -size +100M\n')
    error_handling.add_run('3. 로그 파일 정리: 로그 로테이션 실행\n')
    error_handling.add_run('4. 불필요한 파일 삭제\n\n')

    error_handling.add_run('6.4.4 애플리케이션 응답 없음\n').bold = True
    error_handling.add_run('증상: 웹 페이지가 로드되지 않음\n')
    error_handling.add_run('조치:\n')
    error_handling.add_run('1. 프로세스 상태 확인: ps aux | grep python\n')
    error_handling.add_run('2. 애플리케이션 로그 확인\n')
    error_handling.add_run('3. 애플리케이션 재시작\n')
    error_handling.add_run('4. 시스템 리소스 확인\n\n')

    # 6.5 로그 로테이션
    doc.add_paragraph('6.5 로그 로테이션', style='OpH2')
    log_rotate = doc.add_paragraph()
    log_rotate.add_run('6.5.1 로그 로테이션 설정\n').bold = True
    log_rotate.add_run('/etc/logrotate.d/msys 파일을 생성합니다:\n\n')
    log_rotate.add_run('/home/msys/app/log/*.log {\n').font.name = 'Courier New'
    log_rotate.add_run('    daily\n').font.name = 'Courier New'
    log_rotate.add_run('    rotate 30\n').font.name = 'Courier New'
    log_rotate.add_run('    compress\n').font.name = 'Courier New'
    log_rotate.add_run('    missingok\n').font.name = 'Courier New'
    log_rotate.add_run('    postrotate\n').font.name = 'Courier New'
    log_rotate.add_run('        systemctl reload msys\n').font.name = 'Courier New'
    log_rotate.add_run('    endscript\n').font.name = 'Courier New'
    log_rotate.add_run('}\n\n').font.name = 'Courier New'

    log_rotate.add_run('6.5.2 수동 로테이션 실행\n').bold = True
    log_rotate.add_run('logrotate -f /etc/logrotate.d/msys\n\n').font.name = 'Courier New'

    # 8. 운영 시나리오
    doc.add_paragraph('8. 운영 시나리오', style='OpH1')

    # 8.1 정상 운영 시나리오
    doc.add_paragraph('8.1 정상 운영 시나리오', style='OpH2')
    normal_scenario = doc.add_paragraph()
    normal_scenario.add_run('8.1.1 일일 모니터링 루틴\n').bold = True
    normal_scenario.add_run('운영자가 매일 수행하는 기본 모니터링 절차입니다.\n\n')

    normal_scenario.add_run('단계별 절차:\n').bold = True
    normal_scenario.add_run('1. 시스템 로그인 및 대시보드 확인\n').bold = True
    normal_scenario.add_run('   - 웹 브라우저에서 MSYS 시스템 접속\n')
    normal_scenario.add_run('   - 사용자 ID/비밀번호로 로그인\n')
    normal_scenario.add_run('   - 대시보드 페이지로 이동\n\n')

    normal_scenario.add_run('2. 대시보드 현황 파악\n').bold = True
    normal_scenario.add_run('   - 전체 Job ID 개수 및 총 수집 건수 확인\n')
    normal_scenario.add_run('   - 평균 성공률이 80% 이상인지 확인\n')
    normal_scenario.add_run('   - 빨강/노랑 표시된 위험/경고 Job 확인\n')
    normal_scenario.add_run('   - 연속 실패 횟수가 3회 이상인 Job 식별\n\n')

    normal_scenario.add_run('3. 상세 모니터링 메뉴 점검\n').bold = True
    normal_scenario.add_run('   - 데이터 분석 메뉴에서 최근 7일간 추이 확인\n')
    normal_scenario.add_run('   - 차트 분석 메뉴에서 성공률 그래프 패턴 분석\n')
    normal_scenario.add_run('   - 수집 일정 메뉴에서 오늘 예정된 작업 상태 확인\n\n')

    normal_scenario.add_run('4. 시스템 리소스 점검\n').bold = True
    normal_scenario.add_run('   - CPU 사용률 80% 미만 유지 확인\n')
    normal_scenario.add_run('   - 메모리 사용률 85% 미만 유지 확인\n')
    normal_scenario.add_run('   - 디스크 여유 공간 20% 이상 유지 확인\n\n')

    normal_scenario.add_run('5. 로그 및 이벤트 확인\n').bold = True
    normal_scenario.add_run('   - 최근 1시간 이내 ERROR 레벨 로그 존재 여부 확인\n')
    normal_scenario.add_run('   - WARNING 레벨 로그의 주요 내용 파악\n')
    normal_scenario.add_run('   - 시스템 이벤트 로그에서 비정상 패턴 식별\n\n')

    normal_scenario.add_run('6. 일일 보고서 작성\n').bold = True
    normal_scenario.add_run('   - 이상 징후 발견 시 관련 팀에 보고\n')
    normal_scenario.add_run('   - 정상 운영 확인 시 일일 운영 로그 기록\n\n')

    # 8.2 장애 발생 및 대응 시나리오
    doc.add_paragraph('8.2 장애 발생 및 대응 시나리오', style='OpH2')

    # 8.2.1 데이터 수집 실패 장애
    doc.add_paragraph('8.2.1 데이터 수집 실패 장애', style='OpH3')
    data_failure = doc.add_paragraph()
    data_failure.add_run('증상: 대시보드에서 특정 Job의 성공률 급락 또는 연속 실패 발생\n\n')

    data_failure.add_run('대응 절차:\n').bold = True
    data_failure.add_run('1. 장애 범위 파악\n').bold = True
    data_failure.add_run('   - 실패한 Job ID 및 시간대 확인\n')
    data_failure.add_run('   - 다른 Job에도 영향이 있는지 확인\n')
    data_failure.add_run('   - 데이터 분석 메뉴에서 실패 패턴 분석\n\n')

    data_failure.add_run('2. 로그 분석\n').bold = True
    data_failure.add_run('   - 애플리케이션 로그에서 ERROR 메시지 확인\n')
    data_failure.add_run('   - 데이터베이스 연결 상태 점검\n')
    data_failure.add_run('   - 외부 API 응답 상태 확인\n\n')

    data_failure.add_run('3. 임시 조치\n').bold = True
    data_failure.add_run('   - 해당 Job 일시 중지 (관리자 설정에서 비활성화)\n')
    data_failure.add_run('   - 수동 데이터 수집 시도\n')
    data_failure.add_run('   - 백업 데이터 활용 검토\n\n')

    data_failure.add_run('4. 근본 원인 분석 및 해결\n').bold = True
    data_failure.add_run('   - 네트워크 연결 문제 해결\n')
    data_failure.add_run('   - API 키 또는 인증 정보 갱신\n')
    data_failure.add_run('   - 데이터 소스 변경 사항 대응\n\n')

    data_failure.add_run('5. 복구 및 재개\n').bold = True
    data_failure.add_run('   - Job 재활성화\n')
    data_failure.add_run('   - 누락 데이터 수동 수집\n')
    data_failure.add_run('   - 모니터링 강화\n\n')

    # 8.2.2 시스템 성능 저하 장애
    doc.add_paragraph('8.2.2 시스템 성능 저하 장애', style='OpH3')
    performance_issue = doc.add_paragraph()
    performance_issue.add_run('증상: 웹 페이지 응답 지연, CPU/메모리 사용률 급증\n\n')

    performance_issue.add_run('대응 절차:\n').bold = True
    performance_issue.add_run('1. 성능 모니터링\n').bold = True
    performance_issue.add_run('   - top/htop 명령으로 프로세스별 리소스 사용량 확인\n')
    performance_issue.add_run('   - CPU 모니터링 로그에서 급격한 증가 패턴 분석\n')
    performance_issue.add_run('   - 메모리 누수 여부 확인\n\n')

    performance_issue.add_run('2. 프로세스 분석\n').bold = True
    performance_issue.add_run('   - 고CPU 사용 프로세스 식별\n')
    performance_issue.add_run('   - 비정상 프로세스 강제 종료 검토\n')
    performance_issue.add_run('   - Python 애플리케이션 스레드 덤프 분석\n\n')

    performance_issue.add_run('3. 임시 조치\n').bold = True
    performance_issue.add_run('   - 불필요한 서비스 일시 중지\n')
    performance_issue.add_run('   - 캐시 클리어 및 재시작\n')
    performance_issue.add_run('   - 로드 밸런서에서 해당 서버 제외\n\n')

    performance_issue.add_run('4. 근본 해결\n').bold = True
    performance_issue.add_run('   - 애플리케이션 메모리 누수 수정\n')
    performance_issue.add_run('   - 데이터베이스 쿼리 최적화\n')
    performance_issue.add_run('   - 서버 리소스 증설 검토\n\n')

    # 8.2.3 데이터베이스 장애
    doc.add_paragraph('8.2.3 데이터베이스 장애', style='OpH3')
    db_failure = doc.add_paragraph()
    db_failure.add_run('증상: 데이터베이스 연결 실패, 쿼리 타임아웃\n\n')

    db_failure.add_run('대응 절차:\n').bold = True
    db_failure.add_run('1. 연결 상태 확인\n').bold = True
    db_failure.add_run('   - PostgreSQL 서비스 상태 확인: systemctl status postgresql\n')
    db_failure.add_run('   - 데이터베이스 접속 시도: psql -h host -U user dbname\n')
    db_failure.add_run('   - 네트워크 연결 확인\n\n')

    db_failure.add_run('2. 로그 분석\n').bold = True
    db_failure.add_run('   - PostgreSQL 로그 파일 확인\n')
    db_failure.add_run('   - 디스크 공간 부족 여부 확인\n')
    db_failure.add_run('   - 연결 풀 고갈 여부 분석\n\n')

    db_failure.add_run('3. 복구 조치\n').bold = True
    db_failure.add_run('   - PostgreSQL 서비스 재시작\n')
    db_failure.add_run('   - 긴급 백업으로부터 복원 검토\n')
    db_failure.add_run('   - 읽기 전용 모드로 전환\n\n')

    # 8.3 시스템 점검 시나리오
    doc.add_paragraph('8.3 시스템 점검 시나리오', style='OpH2')
    maintenance_scenario = doc.add_paragraph()
    maintenance_scenario.add_run('8.3.1 주간 점검 절차\n').bold = True
    maintenance_scenario.add_run('매주 수행하는 시스템 정기 점검 절차입니다.\n\n')

    maintenance_scenario.add_run('점검 항목:\n').bold = True
    maintenance_scenario.add_run('1. 백업 상태 확인\n').bold = True
    maintenance_scenario.add_run('   - 자동 백업 스크립트 실행 상태 확인\n')
    maintenance_scenario.add_run('   - 백업 파일 무결성 검증\n')
    maintenance_scenario.add_run('   - 백업 저장 공간 충분성 확인\n\n')

    maintenance_scenario.add_run('2. 로그 파일 정리\n').bold = True
    maintenance_scenario.add_run('   - 오래된 로그 파일 압축 및 아카이브\n')
    maintenance_scenario.add_run('   - 로그 로테이션 설정 확인\n')
    maintenance_scenario.add_run('   - 비정상 로그 패턴 분석\n\n')

    maintenance_scenario.add_run('3. 데이터베이스 유지보수\n').bold = True
    maintenance_scenario.add_run('   - 테이블 및 인덱스 상태 확인\n')
    maintenance_scenario.add_run('   - VACUUM 및 ANALYZE 실행\n')
    maintenance_scenario.add_run('   - 불필요한 데이터 정리\n\n')

    maintenance_scenario.add_run('4. 보안 점검\n').bold = True
    maintenance_scenario.add_run('   - 사용자 계정 상태 확인\n')
    maintenance_scenario.add_run('   - 비밀번호 정책 준수 여부 검토\n')
    maintenance_scenario.add_run('   - 접근 로그에서 이상 징후 분석\n\n')

    # 8.4 데이터 복구 시나리오
    doc.add_paragraph('8.4 데이터 복구 시나리오', style='OpH2')
    recovery_scenario = doc.add_paragraph()
    recovery_scenario.add_run('8.4.1 데이터 유실 복구 절차\n').bold = True
    recovery_scenario.add_run('데이터 유실 발생 시 복구를 위한 절차입니다.\n\n')

    recovery_scenario.add_run('복구 단계:\n').bold = True
    recovery_scenario.add_run('1. 유실 범위 파악\n').bold = True
    recovery_scenario.add_run('   - 어떤 데이터가 유실되었는지 확인\n')
    recovery_scenario.add_run('   - 유실 기간 및 범위 식별\n')
    recovery_scenario.add_run('   - 영향받는 Job 및 사용자 파악\n\n')

    recovery_scenario.add_run('2. 백업 데이터 확인\n').bold = True
    recovery_scenario.add_run('   - 가장 최근 백업 파일 식별\n')
    recovery_scenario.add_run('   - 백업 데이터 무결성 검증\n')
    recovery_scenario.add_run('   - 복구에 필요한 데이터 범위 확인\n\n')

    recovery_scenario.add_run('3. 복구 실행\n').bold = True
    recovery_scenario.add_run('   - 테스트 환경에서 복구 시뮬레이션\n')
    recovery_scenario.add_run('   - 운영 환경으로 복구 적용\n')
    recovery_scenario.add_run('   - 데이터 일관성 검증\n\n')

    recovery_scenario.add_run('4. 검증 및 보고\n').bold = True
    recovery_scenario.add_run('   - 복구된 데이터 정확성 확인\n')
    recovery_scenario.add_run('   - 관련 시스템 정상 작동 확인\n')
    recovery_scenario.add_run('   - 복구 결과 보고서 작성\n\n')

    # 8.5 사용자 관리 시나리오
    doc.add_paragraph('8.5 사용자 관리 시나리오', style='OpH2')
    user_scenario = doc.add_paragraph()
    user_scenario.add_run('8.5.1 신규 사용자 등록 및 권한 부여\n').bold = True
    user_scenario.add_run('새로운 사용자가 시스템을 사용할 수 있도록 등록하고 권한을 부여하는 절차입니다.\n\n')

    user_scenario.add_run('등록 절차:\n').bold = True
    user_scenario.add_run('1. 사용자 회원가입\n').bold = True
    user_scenario.add_run('   - 사용자가 시스템에 접속하여 회원가입 진행\n')
    user_scenario.add_run('   - ID, 비밀번호 입력 및 정책 준수 확인\n')
    user_scenario.add_run('   - 가입 신청 완료 (PENDING 상태로 저장)\n\n')

    user_scenario.add_run('2. 관리자 승인\n').bold = True
    user_scenario.add_run('   - 관리자가 사용자 관리 메뉴 접근\n')
    user_scenario.add_run('   - PENDING 상태 사용자 목록 확인\n')
    user_scenario.add_run('   - 사용자 정보 및 가입 목적 검토\n')
    user_scenario.add_run('   - 승인 버튼 클릭 (APPROVED 상태로 변경)\n\n')

    user_scenario.add_run('3. 초기 비밀번호 설정\n').bold = True
    user_scenario.add_run('   - 시스템이 사용자 ID와 동일한 초기 비밀번호 설정\n')
    user_scenario.add_run('   - 사용자에게 초기 비밀번호 안내\n')
    user_scenario.add_run('   - 보안 정책 위반으로 인한 강제 변경 유도\n\n')

    user_scenario.add_run('4. 메뉴 권한 설정\n').bold = True
    user_scenario.add_run('   - 기본 권한 자동 부여 (dashboard, collection_schedule)\n')
    user_scenario.add_run('   - 업무에 필요한 추가 권한 수동 설정\n')
    user_scenario.add_run('   - 권한별 메뉴 접근 제어 적용\n\n')

    user_scenario.add_run('5. 데이터 접근 권한 설정\n').bold = True
    user_scenario.add_run('   - Job ID별 데이터 접근 권한 부여\n')
    user_scenario.add_run('   - 민감한 데이터에 대한 접근 제한\n')
    user_scenario.add_run('   - 권한 변경 이력 기록\n\n')

    user_scenario.add_run('6. 사용자 교육 및 안내\n').bold = True
    user_scenario.add_run('   - 시스템 사용 방법 안내\n')
    user_scenario.add_run('   - 담당자 연락처 제공\n')
    user_scenario.add_run('   - 추가 지원 요청 방법 설명\n\n')

    try:
        doc.save('MSYS_Operation_Manual_v4.docx')
        print('운영 매뉴얼 생성 완료: MSYS_Operation_Manual_v4.docx')
    except PermissionError:
        doc.save('MSYS_Operation_Manual_v5.docx')
        print('운영 매뉴얼 생성 완료: MSYS_Operation_Manual_v5.docx')

def create_database_manual():
    """DB 관리자용 데이터베이스 매뉴얼 생성"""
    doc = Document()

    # 스타일 설정
    title_style = doc.styles.add_style('DBTitle', WD_STYLE_TYPE.PARAGRAPH)
    title_style.font.size = Pt(28)
    title_style.font.color.rgb = RGBColor(0, 0, 0)
    title_style.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    heading1_style = doc.styles.add_style('DBH1', WD_STYLE_TYPE.PARAGRAPH)
    heading1_style.font.size = Pt(18)
    heading1_style.font.bold = True

    heading2_style = doc.styles.add_style('DBH2', WD_STYLE_TYPE.PARAGRAPH)
    heading2_style.font.size = Pt(14)
    heading2_style.font.bold = True

    heading3_style = doc.styles.add_style('DBH3', WD_STYLE_TYPE.PARAGRAPH)
    heading3_style.font.size = Pt(12)
    heading3_style.font.bold = True

    code_style = doc.styles.add_style('DBCode', WD_STYLE_TYPE.PARAGRAPH)
    code_style.font.name = 'Courier New'
    code_style.font.size = Pt(10)

    # 표지
    title = doc.add_paragraph('MSYS 데이터베이스 매뉴얼', style='DBTitle')
    subtitle = doc.add_paragraph('DB 스키마 및 쿼리 참조 가이드', style='DBTitle')
    version = doc.add_paragraph('버전 1.14.2', style='DBTitle')
    version.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    doc.add_page_break()

    # 목차
    toc_title = doc.add_paragraph('목차', style='DBH1')
    toc = doc.add_paragraph()
    toc.add_run('1. 데이터베이스 개요 ........................ 3\n')
    toc.add_run('2. 접속 정보 ............................. 4\n')
    toc.add_run('3. 테이블 스키마 ........................ 5\n')
    toc.add_run('4. 주요 쿼리 패턴 ........................ 15\n')
    toc.add_run('5. 인덱스 및 최적화 ...................... 20\n')
    toc.add_run('6. 백업 및 복원 ........................ 23\n')
    toc.add_run('부록 A. DDL 스크립트 ...................... 25\n')

    doc.add_page_break()

    # 1. 데이터베이스 개요
    doc.add_paragraph('1. 데이터베이스 개요', style='DBH1')
    overview = doc.add_paragraph()
    overview.add_run('1.1 시스템 개요\n').bold = True
    overview.add_run('MSYS는 데이터 수집 및 모니터링 시스템으로, PostgreSQL 데이터베이스를 사용하여 다양한 데이터를 저장하고 관리합니다.\n\n')

    overview.add_run('1.2 주요 기능\n').bold = True
    overview.add_run('• 데이터 수집 이력 관리\n')
    overview.add_run('• 사용자 및 권한 관리\n')
    overview.add_run('• 시스템 설정 관리\n')
    overview.add_run('• 분석 데이터 저장\n\n')

    overview.add_run('1.3 데이터베이스 버전\n').bold = True
    overview.add_run('• PostgreSQL 13 이상\n')
    overview.add_run('• 타임존: Asia/Seoul\n')
    overview.add_run('• 캐릭터셋: UTF-8\n\n')

    # 2. 접속 정보
    doc.add_paragraph('2. 접속 정보', style='DBH1')
    conn_info = doc.add_paragraph()
    conn_info.add_run('2.1 기본 접속 정보\n').bold = True
    conn_info.add_run('• 호스트: 10.200.153.136\n')
    conn_info.add_run('• 포트: 22543\n')
    conn_info.add_run('• 데이터베이스: etl_db_dev\n')
    conn_info.add_run('• 사용자: etl_user\n')
    conn_info.add_run('• 비밀번호: etl_password\n\n')

    conn_info.add_run('2.2 접속 예시\n').bold = True
    conn_info.add_run('psql -h 10.200.153.136 -p 22543 -U etl_user -d etl_db_dev\n\n').font.name = 'Courier New'

    # 3. 테이블 스키마
    doc.add_paragraph('3. 테이블 스키마', style='DBH1')

    # 3.1 tb_con_hist
    doc.add_paragraph('3.1 tb_con_hist - 데이터 수집 이력', style='DBH2')
    hist_desc = doc.add_paragraph()
    hist_desc.add_run('데이터 수집 작업의 상세 이력을 저장하는 메인 테이블입니다.\n\n')

    hist_desc.add_run('컬럼 상세:\n').bold = True
    hist_table = doc.add_table(rows=1, cols=6)
    hist_table.style = 'Table Grid'
    hdr_cells = hist_table.rows[0].cells
    hdr_cells[0].text = '컬럼명'
    hdr_cells[1].text = '타입'
    hdr_cells[2].text = '길이'
    hdr_cells[3].text = 'NULL'
    hdr_cells[4].text = '기본값'
    hdr_cells[5].text = '설명'

    # 데이터 행 추가
    rows = [
        ['job_id', 'varchar', '250', 'NOT NULL', '', '작업 고유 식별자'],
        ['con_id', 'varchar', '250', 'NOT NULL', '', '연결 고유 식별자'],
        ['rqs_info', 'text', '', 'NULL', '', '요청 정보'],
        ['start_dt', 'timestamptz', '', 'NULL', '', '수집 시작 일시'],
        ['execution_dt', 'timestamptz', '', 'NULL', '', '실행 일시'],
        ['end_dt', 'timestamptz', '', 'NULL', '', '수집 종료 일시'],
        ['status', 'varchar', '20', 'NULL', '', '수집 상태'],
        ['trbl_hist_no', 'integer', '', 'NULL', '', '문제 이력 번호']
    ]

    for row_data in rows:
        row_cells = hist_table.add_row().cells
        for i, cell_data in enumerate(row_data):
            row_cells[i].text = cell_data

    hist_desc.add_run('\n\n상태 코드 값:\n').bold = True
    hist_desc.add_run('• CD901: 정상(성공) - 데이터 수집 성공\n')
    hist_desc.add_run('• CD902: 실패 - 데이터 수집 실패\n')
    hist_desc.add_run('• CD903: 미수집 - 예정된 수집 누락\n')
    hist_desc.add_run('• CD904: 측정중 - 수집 진행 중\n\n')

    hist_desc.add_run('제약조건:\n').bold = True
    hist_desc.add_run('• 프라이머리 키: (con_id, job_id)\n')
    hist_desc.add_run('• 트리거: trg_log_con_hist_changes (변경 로그 기록)\n\n')

    # 3.2 tb_user
    doc.add_paragraph('3.2 tb_user - 사용자 정보', style='DBH2')
    user_desc = doc.add_paragraph()
    user_desc.add_run('시스템 사용자의 계정 정보를 저장하는 테이블입니다.\n\n')

    user_table = doc.add_table(rows=1, cols=6)
    user_table.style = 'Table Grid'
    hdr_cells = user_table.rows[0].cells
    hdr_cells[0].text = '컬럼명'
    hdr_cells[1].text = '타입'
    hdr_cells[2].text = '길이'
    hdr_cells[3].text = 'NULL'
    hdr_cells[4].text = '기본값'
    hdr_cells[5].text = '설명'

    user_rows = [
        ['user_id', 'varchar', '50', 'NOT NULL', '', '사용자 고유 ID'],
        ['user_pwd', 'varchar', '255', 'NOT NULL', '', '해시된 비밀번호'],
        ['acc_sts', 'varchar', '20', 'NOT NULL', "'PENDING'", '계정 상태'],
        ['acc_cre_dt', 'timestamptz', '', 'NULL', 'CURRENT_TIMESTAMP', '계정 생성 일시'],
        ['acc_apr_dt', 'timestamptz', '', 'NULL', '', '계정 승인 일시']
    ]

    for row_data in user_rows:
        row_cells = user_table.add_row().cells
        for i, cell_data in enumerate(row_data):
            row_cells[i].text = cell_data

    user_desc.add_run('\n\n계정 상태 코드:\n').bold = True
    user_desc.add_run('• PENDING: 승인 대기 - 회원가입 후 관리자 승인 필요\n')
    user_desc.add_run('• APPROVED: 승인 완료 - 정상 사용 가능\n')
    user_desc.add_run('• REJECTED: 승인 거부 - 사용 불가\n')
    user_desc.add_run('• SUSPENDED: 정지 - 일시적 사용 제한\n\n')

    # 3.3 tb_con_mst
    doc.add_paragraph('3.3 tb_con_mst - 마스터 코드', style='DBH2')
    mst_desc = doc.add_paragraph()
    mst_desc.add_run('시스템에서 사용하는 각종 코드와 설정값을 저장하는 마스터 테이블입니다.\n\n')

    mst_table = doc.add_table(rows=1, cols=6)
    mst_table.style = 'Table Grid'
    hdr_cells = mst_table.rows[0].cells
    hdr_cells[0].text = '컬럼명'
    hdr_cells[1].text = '타입'
    hdr_cells[2].text = '길이'
    hdr_cells[3].text = 'NULL'
    hdr_cells[4].text = '기본값'
    hdr_cells[5].text = '설명'

    mst_rows = [
        ['cd_cl', 'varchar', '20', 'NOT NULL', '', '코드 분류'],
        ['cd', 'varchar', '20', 'NOT NULL', '', '코드 값'],
        ['cd_nm', 'varchar', '20', 'NULL', '', '코드명'],
        ['cd_desc', 'varchar', '50', 'NULL', '', '코드 설명'],
        ['item1', 'varchar', '50', 'NULL', '', '추가 항목 1'],
        ['item2', 'varchar', '50', 'NULL', '', '추가 항목 2'],
        ['item3', 'varchar', '150', 'NULL', '', '추가 항목 3'],
        ['item4', 'varchar', '50', 'NULL', '', '추가 항목 4'],
        ['item5', 'varchar', '50', 'NULL', '', '추가 항목 5'],
        ['item6', 'varchar', '50', 'NULL', '', '추가 항목 6'],
        ['item7', 'varchar', '50', 'NULL', '', '추가 항목 7'],
        ['item8', 'varchar', '50', 'NULL', '', '추가 항목 8'],
        ['item9', 'varchar', '400', 'NULL', '', '추가 항목 9'],
        ['item10', 'varchar', '400', 'NULL', '', '추가 항목 10'],
        ['update_dt', 'timestamptz', '', 'NULL', '', '수정 일시'],
        ['del_dt', 'timestamptz', '', 'NULL', '', '삭제 일시'],
        ['use_yn', 'char', '18', 'NULL', '', '사용 여부']
    ]

    for row_data in mst_rows:
        row_cells = mst_table.add_row().cells
        for i, cell_data in enumerate(row_data):
            row_cells[i].text = cell_data

    mst_desc.add_run('\n\n주요 코드 분류:\n').bold = True
    mst_desc.add_run('• CD900: 장애 코드 (item1: 영문명)\n')
    mst_desc.add_run('• CD100: 작업 유형 코드\n')
    mst_desc.add_run('• CD200: 상태 코드\n\n')

    # 3.4 tb_menu
    doc.add_paragraph('3.4 tb_menu - 메뉴 정보', style='DBH2')
    menu_desc = doc.add_paragraph()
    menu_desc.add_run('시스템의 메뉴 구조를 정의하는 테이블입니다.\n\n')

    menu_table = doc.add_table(rows=1, cols=6)
    menu_table.style = 'Table Grid'
    hdr_cells = menu_table.rows[0].cells
    hdr_cells[0].text = '컬럼명'
    hdr_cells[1].text = '타입'
    hdr_cells[2].text = '길이'
    hdr_cells[3].text = 'NULL'
    hdr_cells[4].text = '기본값'
    hdr_cells[5].text = '설명'

    menu_rows = [
        ['menu_id', 'varchar', '50', 'NOT NULL', '', '메뉴 고유 ID'],
        ['menu_nm', 'varchar', '100', 'NOT NULL', '', '메뉴 표시 이름'],
        ['menu_url', 'varchar', '255', 'NOT NULL', '', '메뉴 URL 경로'],
        ['menu_order', 'integer', '', 'NULL', '', '메뉴 표시 순서']
    ]

    for row_data in menu_rows:
        row_cells = menu_table.add_row().cells
        for i, cell_data in enumerate(row_data):
            row_cells[i].text = cell_data

    menu_desc.add_run('\n\n제약조건:\n').bold = True
    menu_desc.add_run('• 프라이머리 키: menu_id\n\n')

    # 3.5 tb_mngr_sett
    doc.add_paragraph('3.5 tb_mngr_sett - 관리자 설정', style='DBH2')
    sett_desc = doc.add_paragraph()
    sett_desc.add_run('Job별 모니터링 임계값과 표시 설정을 저장하는 테이블입니다.\n\n')

    sett_table = doc.add_table(rows=1, cols=6)
    sett_table.style = 'Table Grid'
    hdr_cells = sett_table.rows[0].cells
    hdr_cells[0].text = '컬럼명'
    hdr_cells[1].text = '타입'
    hdr_cells[2].text = '길이'
    hdr_cells[3].text = 'NULL'
    hdr_cells[4].text = '기본값'
    hdr_cells[5].text = '설명'

    sett_rows = [
        ['cd', 'varchar', '50', 'NOT NULL', '', '설정 대상 Job ID'],
        ['cnn_failr_thrs_val', 'integer', '', 'NULL', '3', '연속 실패 임계값'],
        ['cnn_warn_thrs_val', 'integer', '', 'NULL', '2', '연속 경고 임계값'],
        ['cnn_failr_icon_id', 'integer', '', 'NULL', '', '실패 상태 아이콘 ID'],
        ['cnn_failr_wrd_colr', 'varchar', '7', 'NULL', "'#FF0000'", '실패 상태 색상'],
        ['cnn_warn_icon_id', 'integer', '', 'NULL', '', '경고 상태 아이콘 ID'],
        ['cnn_warn_wrd_colr', 'varchar', '7', 'NULL', "'#FFA500'", '경고 상태 색상'],
        ['cnn_sucs_icon_id', 'integer', '', 'NULL', '', '성공 상태 아이콘 ID'],
        ['cnn_sucs_wrd_colr', 'varchar', '7', 'NULL', "'#008000'", '성공 상태 색상'],
        ['dly_sucs_rt_thrs_val', 'integer', '', 'NULL', '80', '일별 성공률 임계값(%)'],
        ['dd7_sucs_rt_thrs_val', 'integer', '', 'NULL', '75', '7일 성공률 임계값(%)'],
        ['mthl_sucs_rt_thrs_val', 'integer', '', 'NULL', '70', '월별 성공률 임계값(%)'],
        ['chrt_dsp_yn', 'boolean', '', 'NULL', 'true', '차트 표시 여부']
    ]

    for row_data in sett_rows:
        row_cells = sett_table.add_row().cells
        for i, cell_data in enumerate(row_data):
            row_cells[i].text = cell_data

    sett_desc.add_run('\n\n제약조건:\n').bold = True
    sett_desc.add_run('• 프라이머리 키: cd\n')
    sett_desc.add_run('• 외래키: cd → tb_con_mst.cd (Job ID 참조)\n\n')

    # 3.6 tb_data_spec
    doc.add_paragraph('3.6 tb_data_spec - 데이터 명세서', style='DBH2')
    spec_desc = doc.add_paragraph()
    spec_desc.add_run('외부 API의 명세 정보를 저장하는 테이블입니다.\n\n')

    spec_table = doc.add_table(rows=1, cols=6)
    spec_table.style = 'Table Grid'
    hdr_cells = spec_table.rows[0].cells
    hdr_cells[0].text = '컬럼명'
    hdr_cells[1].text = '타입'
    hdr_cells[2].text = '길이'
    hdr_cells[3].text = 'NULL'
    hdr_cells[4].text = '기본값'
    hdr_cells[5].text = '설명'

    spec_rows = [
        ['id', 'integer', '', 'NOT NULL', 'nextval()', '고유 식별자'],
        ['data_name', 'varchar', '255', 'NOT NULL', '', '데이터 명칭'],
        ['description', 'text', '', 'NULL', '', '데이터 설명'],
        ['api_url', 'varchar', '2048', 'NULL', '', 'API 엔드포인트 URL'],
        ['provider', 'varchar', '255', 'NULL', '', '데이터 제공 기관'],
        ['keywords', 'varchar', '1024', 'NULL', '', '검색 키워드'],
        ['reference_doc_url', 'varchar', '2048', 'NULL', '', '참고 문서 URL'],
        ['created_at', 'timestamptz', '', 'NULL', 'CURRENT_TIMESTAMP', '생성 일시'],
        ['updated_at', 'timestamptz', '', 'NULL', 'CURRENT_TIMESTAMP', '수정 일시']
    ]

    for row_data in spec_rows:
        row_cells = spec_table.add_row().cells
        for i, cell_data in enumerate(row_data):
            row_cells[i].text = cell_data

    spec_desc.add_run('\n\n제약조건:\n').bold = True
    spec_desc.add_run('• 프라이머리 키: id\n')
    spec_desc.add_run('• 시퀀스: tb_data_spec_id_seq\n\n')

    # 3.7 tb_icon
    doc.add_paragraph('3.7 tb_icon - 아이콘 정보', style='DBH2')
    icon_desc = doc.add_paragraph()
    icon_desc.add_run('시스템에서 사용하는 아이콘 정보를 저장하는 테이블입니다.\n\n')

    icon_table = doc.add_table(rows=1, cols=6)
    icon_table.style = 'Table Grid'
    hdr_cells = icon_table.rows[0].cells
    hdr_cells[0].text = '컬럼명'
    hdr_cells[1].text = '타입'
    hdr_cells[2].text = '길이'
    hdr_cells[3].text = 'NULL'
    hdr_cells[4].text = '기본값'
    hdr_cells[5].text = '설명'

    icon_rows = [
        ['icon_id', 'integer', '', 'NOT NULL', 'nextval()', '아이콘 고유 ID'],
        ['icon_cd', 'text', '', 'NOT NULL', '', '아이콘 코드'],
        ['icon_nm', 'varchar', '50', 'NOT NULL', '', '아이콘 이름'],
        ['icon_expl', 'varchar', '255', 'NULL', '', '아이콘 설명'],
        ['icon_cre_dt', 'timestamp', '', 'NULL', 'CURRENT_TIMESTAMP', '생성 일시'],
        ['icon_dsp_yn', 'boolean', '', 'NULL', 'true', '표시 여부']
    ]

    for row_data in icon_rows:
        row_cells = icon_table.add_row().cells
        for i, cell_data in enumerate(row_data):
            row_cells[i].text = cell_data

    icon_desc.add_run('\n\n제약조건:\n').bold = True
    icon_desc.add_run('• 프라이머리 키: icon_id\n')
    icon_desc.add_run('• 유니크: icon_cd\n')
    icon_desc.add_run('• 시퀀스: tb_icons_icon_id_seq\n\n')

    # 3.8 tb_user_auth_ctrl
    doc.add_paragraph('3.8 tb_user_auth_ctrl - 사용자 권한 제어', style='DBH2')
    auth_desc = doc.add_paragraph()
    auth_desc.add_run('사용자별 메뉴 접근 권한을 제어하는 테이블입니다.\n\n')

    auth_table = doc.add_table(rows=1, cols=6)
    auth_table.style = 'Table Grid'
    hdr_cells = auth_table.rows[0].cells
    hdr_cells[0].text = '컬럼명'
    hdr_cells[1].text = '타입'
    hdr_cells[2].text = '길이'
    hdr_cells[3].text = 'NULL'
    hdr_cells[4].text = '기본값'
    hdr_cells[5].text = '설명'

    auth_rows = [
        ['auth_id', 'integer', '', 'NOT NULL', 'nextval()', '권한 고유 ID'],
        ['user_id', 'varchar', '50', 'NOT NULL', '', '사용자 ID'],
        ['menu_id', 'varchar', '50', 'NOT NULL', '', '메뉴 ID'],
        ['auth_yn', 'boolean', '', 'NOT NULL', 'true', '권한 여부']
    ]

    for row_data in auth_rows:
        row_cells = auth_table.add_row().cells
        for i, cell_data in enumerate(row_data):
            row_cells[i].text = cell_data

    auth_desc.add_run('\n\n제약조건:\n').bold = True
    auth_desc.add_run('• 프라이머리 키: auth_id\n')
    auth_desc.add_run('• 유니크: (user_id, menu_id)\n')
    auth_desc.add_run('• 외래키: user_id → tb_user.user_id\n')
    auth_desc.add_run('• 외래키: menu_id → tb_menu.menu_id\n')
    auth_desc.add_run('• 시퀀스: tb_user_access_control_id_seq\n\n')

    # 3.9 tb_user_data_perm_auth_ctrl
    doc.add_paragraph('3.9 tb_user_data_perm_auth_ctrl - 데이터 접근 권한', style='DBH2')
    data_perm_desc = doc.add_paragraph()
    data_perm_desc.add_run('사용자별 Job ID 데이터 접근 권한을 제어하는 테이블입니다.\n\n')

    data_perm_table = doc.add_table(rows=1, cols=6)
    data_perm_table.style = 'Table Grid'
    hdr_cells = data_perm_table.rows[0].cells
    hdr_cells[0].text = '컬럼명'
    hdr_cells[1].text = '타입'
    hdr_cells[2].text = '길이'
    hdr_cells[3].text = 'NULL'
    hdr_cells[4].text = '기본값'
    hdr_cells[5].text = '설명'

    data_perm_rows = [
        ['perm_id', 'integer', '', 'NOT NULL', 'nextval()', '권한 고유 ID'],
        ['user_id', 'varchar', '50', 'NOT NULL', '', '사용자 ID'],
        ['job_id', 'varchar', '50', 'NOT NULL', '', 'Job ID'],
        ['perm_yn', 'boolean', '', 'NOT NULL', 'true', '접근 권한 여부']
    ]

    for row_data in data_perm_rows:
        row_cells = data_perm_table.add_row().cells
        for i, cell_data in enumerate(row_data):
            row_cells[i].text = cell_data

    data_perm_desc.add_run('\n\n제약조건:\n').bold = True
    data_perm_desc.add_run('• 프라이머리 키: perm_id\n')
    data_perm_desc.add_run('• 유니크: (user_id, job_id)\n')
    data_perm_desc.add_run('• 외래키: user_id → tb_user.user_id\n')
    data_perm_desc.add_run('• 외래키: job_id → tb_con_mst.cd\n\n')

    # 3.10 기타 주요 테이블 개요
    doc.add_paragraph('3.10 기타 주요 테이블 개요', style='DBH2')
    other_tables = doc.add_paragraph()
    other_tables.add_run('시스템에서 사용하는 기타 주요 테이블들:\n\n').bold = True

    other_tables.add_run('• tb_con_hist_evnt_log: 수집 이력 변경 로그\n')
    other_tables.add_run('• tb_con_trbl_hist: 문제 발생 이력\n')
    other_tables.add_run('• tb_data_clt_schd_sett: 수집 스케줄 설정\n')
    other_tables.add_run('• tb_data_spec_parm: 데이터 명세서 파라미터\n')
    other_tables.add_run('• tb_user_acs_log: 사용자 접근 로그\n')
    other_tables.add_run('• tb_col_mapp: 컬럼 매핑 정보\n\n')

    # 4. 주요 쿼리 패턴
    doc.add_paragraph('4. 주요 쿼리 패턴', style='DBH1')

    # 4.1 성공률 계산 쿼리
    doc.add_paragraph('4.1 성공률 계산 쿼리', style='DBH2')
    success_query = doc.add_paragraph()
    success_query.add_run('일별/Job별 성공률을 계산하는 표준 쿼리 패턴:\n\n')
    success_query.add_run('SELECT\n')
    success_query.add_run('    (start_dt AT TIME ZONE \'Asia/Seoul\')::date AS date,\n')
    success_query.add_run('    job_id,\n')
    success_query.add_run('    (SUM(CASE WHEN status = \'CD901\' THEN 1 ELSE 0 END) * 100.0 / NULLIF(COUNT(*), 0)) AS success_rate\n')
    success_query.add_run('FROM tb_con_hist\n')
    success_query.add_run('{where_clause}\n')
    success_query.add_run('GROUP BY (start_dt AT TIME ZONE \'Asia/Seoul\')::date, job_id\n')
    success_query.add_run('ORDER BY (start_dt AT TIME ZONE \'Asia/Seoul\')::date, job_id;\n\n').font.name = 'Courier New'

    success_query.add_run('계산 방식:\n').bold = True
    success_query.add_run('• 성공률 = (성공 건수 ÷ 전체 건수) × 100\n')
    success_query.add_run('• 성공 기준: status = \'CD901\'\n')
    success_query.add_run('• 시간대: Asia/Seoul 기준\n\n')

    # 4.2 집계 쿼리 패턴
    doc.add_paragraph('4.2 집계 쿼리 패턴', style='DBH2')
    agg_query = doc.add_paragraph()
    agg_query.add_run('다양한 기간별 통계를 계산하는 집계 쿼리:\n\n')
    agg_query.add_run('-- 일별 집계\n')
    agg_query.add_run('SELECT DATE_TRUNC(\'day\', start_dt AT TIME ZONE \'Asia/Seoul\') AS period,\n')
    agg_query.add_run('       COUNT(*) as total_count,\n')
    agg_query.add_run('       SUM(CASE WHEN status = \'CD901\' THEN 1 ELSE 0 END) as success_count\n')
    agg_query.add_run('FROM tb_con_hist\n')
    agg_query.add_run('WHERE start_dt >= \'2025-01-01\'\n')
    agg_query.add_run('GROUP BY DATE_TRUNC(\'day\', start_dt AT TIME ZONE \'Asia/Seoul\')\n')
    agg_query.add_run('ORDER BY period;\n\n').font.name = 'Courier New'

    # 5. 인덱스 및 최적화
    doc.add_paragraph('5. 인덱스 및 최적화', style='DBH1')
    index_desc = doc.add_paragraph()
    index_desc.add_run('5.1 주요 인덱스\n').bold = True
    index_desc.add_run('• tb_con_hist: (con_id, job_id) - 프라이머리 키\n')
    index_desc.add_run('• tb_con_hist: start_dt - 시간 범위 쿼리 최적화\n')
    index_desc.add_run('• tb_con_hist: (job_id, start_dt) - Job별 기간 쿼리\n')
    index_desc.add_run('• tb_user: user_id - 프라이머리 키\n')
    index_desc.add_run('• tb_con_mst: (cd_cl, cd) - 프라이머리 키\n\n')

    index_desc.add_run('5.2 쿼리 최적화 팁\n').bold = True
    index_desc.add_run('• 시간 범위 쿼리: start_dt 인덱스 활용\n')
    index_desc.add_run('• Job 필터링: job_id + start_dt 복합 인덱스\n')
    index_desc.add_run('• 대량 데이터: LIMIT과 OFFSET 적절히 사용\n')
    index_desc.add_run('• 파티셔닝: 대용량 테이블은 날짜별 파티셔닝 고려\n\n')

    # 6. 백업 및 복원
    doc.add_paragraph('6. 백업 및 복원', style='DBH1')
    backup_desc = doc.add_paragraph()
    backup_desc.add_run('6.1 전체 백업\n').bold = True
    backup_desc.add_run('pg_dump -h 10.200.153.136 -p 22543 -U etl_user -d etl_db_dev > backup.sql\n\n').font.name = 'Courier New'

    backup_desc.add_run('6.2 복원\n').bold = True
    backup_desc.add_run('psql -h 10.200.153.136 -p 22543 -U etl_user -d etl_db_dev < backup.sql\n\n').font.name = 'Courier New'

    backup_desc.add_run('6.3 자동 백업 스크립트\n').bold = True
    backup_desc.add_run('# crontab에 추가\n')
    backup_desc.add_run('0 2 * * * pg_dump -U etl_user etl_db_dev > /backup/msys_$(date +\\%Y\\%m\\%d).sql\n\n').font.name = 'Courier New'

    doc.save('MSYS_Database_Manual.docx')
    print('데이터베이스 매뉴얼 생성 완료: MSYS_Database_Manual.docx')

def create_cpu_monitor_manual():
    """CPU 모니터링 앱 설명서 생성"""
    doc = Document()

    # 스타일 설정
    title_style = doc.styles.add_style('CpuTitle', WD_STYLE_TYPE.PARAGRAPH)
    title_style.font.size = Pt(28)
    title_style.font.color.rgb = RGBColor(0, 0, 0)
    title_style.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    heading1_style = doc.styles.add_style('CpuH1', WD_STYLE_TYPE.PARAGRAPH)
    heading1_style.font.size = Pt(18)
    heading1_style.font.bold = True

    heading2_style = doc.styles.add_style('CpuH2', WD_STYLE_TYPE.PARAGRAPH)
    heading2_style.font.size = Pt(14)
    heading2_style.font.bold = True

    heading3_style = doc.styles.add_style('CpuH3', WD_STYLE_TYPE.PARAGRAPH)
    heading3_style.font.size = Pt(12)
    heading3_style.font.bold = True

    code_style = doc.styles.add_style('CpuCode', WD_STYLE_TYPE.PARAGRAPH)
    code_style.font.name = 'Courier New'
    code_style.font.size = Pt(10)

    # 표지
    title = doc.add_paragraph('MSYS CPU 모니터링 앱 설명서', style='CpuTitle')
    subtitle = doc.add_paragraph('시스템 리소스 모니터링 가이드', style='CpuTitle')
    version = doc.add_paragraph('버전 1.0', style='CpuTitle')
    version.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    doc.add_page_break()

    # 목차
    toc_title = doc.add_paragraph('목차', style='CpuH1')
    toc = doc.add_paragraph()
    toc.add_run('1. 개요 ........................ 3\n')
    toc.add_run('2. 설치 및 실행 ................ 4\n')
    toc.add_run('3. 모니터링 기능 ................ 6\n')
    toc.add_run('4. 설정 옵션 .................. 9\n')
    toc.add_run('5. 로그 및 결과 해석 ........... 11\n')
    toc.add_run('6. 문제 해결 .................. 14\n')
    toc.add_run('부록 A. 기술 사양 ................ 16\n')

    doc.add_page_break()

    # 1. 개요
    doc.add_paragraph('1. 개요', style='CpuH1')
    overview = doc.add_paragraph()
    overview.add_run('1.1 목적\n').bold = True
    overview.add_run('MSYS CPU 모니터링 앱은 시스템의 CPU 사용률을 실시간으로 모니터링하고, 임계값 초과 시 경고를 기록하는 도구입니다. 시스템 성능 저하의 조기 발견과 문제 해결을 지원합니다.\n\n')

    overview.add_run('1.2 주요 기능\n').bold = True
    overview.add_run('• 실시간 CPU 사용률 모니터링\n')
    overview.add_run('• 임계값 기반 경고 시스템\n')
    overview.add_run('• 급격한 CPU 사용량 증가 감지\n')
    overview.add_run('• 상위 CPU 사용 프로세스 식별\n')
    overview.add_run('• 자동 로그 기록 및 로테이션\n\n')

    overview.add_run('1.3 지원 플랫폼\n').bold = True
    overview.add_run('• Windows (pythonw.exe 사용)\n')
    overview.add_run('• Linux/macOS (nohup 사용)\n')
    overview.add_run('• 백그라운드 실행 지원\n\n')

    # 2. 설치 및 실행
    doc.add_paragraph('2. 설치 및 실행', style='CpuH1')

    # 2.1 요구사항
    doc.add_paragraph('2.1 시스템 요구사항', style='CpuH2')
    reqs = doc.add_paragraph()
    reqs.add_run('• Python 3.8 이상\n')
    reqs.add_run('• psutil 라이브러리 (pip install psutil)\n')
    reqs.add_run('• 로그 파일 쓰기 권한\n\n')

    # 2.2 설치
    doc.add_paragraph('2.2 설치', style='CpuH2')
    install = doc.add_paragraph()
    install.add_run('필수 라이브러리 설치:\n\n')
    install.add_run('pip install psutil\n\n').font.name = 'Courier New'

    # 2.3 실행
    doc.add_paragraph('2.3 실행 방법', style='CpuH2')
    run = doc.add_paragraph()
    run.add_run('2.3.1 백그라운드 실행 (권장)\n').bold = True
    run.add_run('python start_cpu_monitor.py\n\n').font.name = 'Courier New'

    run.add_run('2.3.2 직접 실행\n').bold = True
    run.add_run('python utils/cpu_monitor.py\n\n').font.name = 'Courier New'

    run.add_run('2.3.3 테스트 모드 실행\n').bold = True
    run.add_run('python utils/cpu_monitor.py test\n\n').font.name = 'Courier New'

    # 2.4 실행 확인
    doc.add_paragraph('2.4 실행 상태 확인', style='CpuH2')
    check = doc.add_paragraph()
    check.add_run('프로세스 확인 (Windows):\n\n')
    check.add_run('tasklist /FI "IMAGENAME eq python.exe"\n\n').font.name = 'Courier New'

    check.add_run('프로세스 확인 (Linux):\n\n')
    check.add_run('ps aux | grep python\n\n').font.name = 'Courier New'

    # 3. 모니터링 기능
    doc.add_paragraph('3. 모니터링 기능', style='CpuH1')

    # 3.1 기본 모니터링
    doc.add_paragraph('3.1 기본 CPU 모니터링', style='CpuH2')
    basic = doc.add_paragraph()
    basic.add_run('• 5초 간격으로 CPU 사용률 측정\n')
    basic.add_run('• 실시간 로그 기록\n')
    basic.add_run('• 백분율 단위 표시\n\n')

    # 3.2 임계값 모니터링
    doc.add_paragraph('3.2 임계값 기반 경고', style='CpuH2')
    threshold = doc.add_paragraph()
    threshold.add_run('기본 임계값:\n').bold = True
    threshold.add_run('• CPU 사용률 80% 초과 시 경고\n')
    threshold.add_run('• 이전 값 대비 20% 급격한 증가 시 경고\n\n')

    # 3.3 프로세스 모니터링
    doc.add_paragraph('3.3 상위 CPU 사용 프로세스', style='CpuH2')
    process = doc.add_paragraph()
    process.add_run('CPU 사용률이 임계값을 초과하면 상위 5개 프로세스 정보를 기록:\n').bold = True
    process.add_run('• 프로세스 ID (PID)\n')
    process.add_run('• 프로세스 이름\n')
    process.add_run('• CPU 사용률\n\n')

    # 3.4 급격한 변화 감지
    doc.add_paragraph('3.4 급격한 CPU 변화 감지', style='CpuH2')
    spike = doc.add_paragraph()
    spike.add_run('이전 측정값과 비교하여 급격한 증가를 감지:\n').bold = True
    spike.add_run('• 기본 임계값: 20% 증가\n')
    spike.add_run('• 스팸 방지: 10초 쿨다운\n\n')

    # 4. 설정 옵션
    doc.add_paragraph('4. 설정 옵션', style='CpuH1')

    # 4.1 모니터링 간격
    doc.add_paragraph('4.1 모니터링 간격 설정', style='CpuH2')
    interval = doc.add_paragraph()
    interval.add_run('interval 매개변수로 설정 (기본값: 5초)\n\n')
    interval.add_run('monitor_cpu(interval=10)  # 10초 간격\n\n').font.name = 'Courier New'

    # 4.2 임계값 설정
    doc.add_paragraph('4.2 임계값 설정', style='CpuH2')
    thresholds = doc.add_paragraph()
    thresholds.add_run('threshold: CPU 사용률 임계값 (기본값: 80%)\n')
    thresholds.add_run('spike_threshold: 급격한 증가 임계값 (기본값: 20%)\n\n')
    thresholds.add_run('monitor_cpu(threshold=90, spike_threshold=15)\n\n').font.name = 'Courier New'

    # 4.3 테스트 모드
    doc.add_paragraph('4.3 테스트 모드', style='CpuH2')
    test_mode = doc.add_paragraph()
    test_mode.add_run('test_mode=True로 설정 시 3번만 실행하고 종료\n\n')
    test_mode.add_run('monitor_cpu(test_mode=True)\n\n').font.name = 'Courier New'

    # 5. 로그 및 결과 해석
    doc.add_paragraph('5. 로그 및 결과 해석', style='CpuH1')

    # 5.1 로그 파일 위치
    doc.add_paragraph('5.1 로그 파일', style='CpuH2')
    log_location = doc.add_paragraph()
    log_location.add_run('• 기본 위치: log/cpu_monitor.log\n')
    log_location.add_run('• 로그 로테이션: 일별 자동 로테이션\n')
    log_location.add_run('• 압축: 이전 로그 자동 압축\n\n')

    # 5.2 로그 형식
    doc.add_paragraph('5.2 로그 형식', style='CpuH2')
    log_format = doc.add_paragraph()
    log_format.add_run('일반 로그:\n').bold = True
    log_format.add_run('2025-12-23 13:30:15 - INFO - 현재 CPU 사용률: 45.2%\n\n').font.name = 'Courier New'

    log_format.add_run('경고 로그:\n').bold = True
    log_format.add_run('2025-12-23 13:30:20 - WARNING - CPU 사용률 경고: 85.7%\n\n').font.name = 'Courier New'

    log_format.add_run('급격한 증가 로그:\n').bold = True
    log_format.add_run('2025-12-23 13:30:25 - WARNING - CPU 급격한 상승 감지: 45.2% → 85.7% (+40.5%)\n\n').font.name = 'Courier New'

    # 5.3 프로세스 정보 로그
    doc.add_paragraph('5.3 프로세스 정보 로그', style='CpuH2')
    proc_log = doc.add_paragraph()
    proc_log.add_run('상위 CPU 사용 프로세스:\n').bold = True
    proc_log.add_run('2025-12-23 13:30:30 - WARNING -   PID: 1234, 이름: python.exe, CPU: 45.2%\n')
    proc_log.add_run('2025-12-23 13:30:30 - WARNING -   PID: 5678, 이름: chrome.exe, CPU: 23.1%\n\n').font.name = 'Courier New'

    # 5.4 로그 분석
    doc.add_paragraph('5.4 로그 분석 방법', style='CpuH2')
    analysis = doc.add_paragraph()
    analysis.add_run('5.4.1 정상 패턴\n').bold = True
    analysis.add_run('• CPU 사용률 1-50%: 정상 작동\n')
    analysis.add_run('• 일정한 패턴 유지\n\n')

    analysis.add_run('5.4.2 경고 패턴\n').bold = True
    analysis.add_run('• CPU 사용률 80% 초과 지속\n')
    analysis.add_run('• 급격한 사용량 증가\n')
    analysis.add_run('• 상위 프로세스 확인 필요\n\n')

    analysis.add_run('5.4.3 문제 패턴\n').bold = True
    analysis.add_run('• CPU 사용률 95% 이상 지속\n')
    analysis.add_run('• 동일 프로세스의 반복적 고사용량\n')
    analysis.add_run('• 시스템 응답성 저하\n\n')

    # 6. 문제 해결
    doc.add_paragraph('6. 문제 해결', style='CpuH1')

    # 6.1 일반적인 문제
    doc.add_paragraph('6.1 일반적인 문제', style='CpuH2')
    issues = doc.add_paragraph()
    issues.add_run('6.1.1 모니터링이 시작되지 않음\n').bold = True
    issues.add_run('• psutil 라이브러리 설치 확인: pip install psutil\n')
    issues.add_run('• Python 경로 확인\n')
    issues.add_run('• 로그 디렉토리 쓰기 권한 확인\n\n')

    issues.add_run('6.1.2 로그 파일이 생성되지 않음\n').bold = True
    issues.add_run('• log 디렉토리 존재 확인\n')
    issues.add_run('• 파일 시스템 권한 확인\n')
    issues.add_run('• 디스크 공간 부족 확인\n\n')

    issues.add_run('6.1.3 CPU 사용률이 0%로 표시됨\n').bold = True
    issues.add_run('• 시스템 권한 확인 (관리자 권한 필요할 수 있음)\n')
    issues.add_run('• psutil 버전 호환성 확인\n\n')

    # 6.2 고급 문제 해결
    doc.add_paragraph('6.2 고급 문제 해결', style='CpuH2')
    advanced = doc.add_paragraph()
    advanced.add_run('6.2.1 백그라운드 프로세스 종료\n').bold = True
    advanced.add_run('Windows:\n')
    advanced.add_run('taskkill /F /IM python.exe\n\n').font.name = 'Courier New'
    advanced.add_run('Linux:\n')
    advanced.add_run('pkill -f cpu_monitor.py\n\n').font.name = 'Courier New'

    advanced.add_run('6.2.2 로그 분석 스크립트\n').bold = True
    advanced.add_run('grep "WARNING" log/cpu_monitor.log | tail -10\n\n').font.name = 'Courier New'

    # 부록 A. 기술 사양
    doc.add_paragraph('부록 A. 기술 사양', style='CpuH1')
    spec = doc.add_paragraph()
    spec.add_run('A.1 의존성 라이브러리\n').bold = True
    spec.add_run('• psutil >= 5.0.0: 시스템 정보 수집\n')
    spec.add_run('• Python >= 3.8: 런타임 환경\n\n')

    spec.add_run('A.2 시스템 요구사항\n').bold = True
    spec.add_run('• 메모리: 최소 50MB\n')
    spec.add_run('• 저장소: 최소 10MB (로그 포함)\n')
    spec.add_run('• CPU: 모니터링 오버헤드 최소\n\n')

    spec.add_run('A.3 제한사항\n').bold = True
    spec.add_run('• Windows: 관리자 권한 필요 (프로세스 정보 수집 시)\n')
    spec.add_run('• Linux: /proc 파일시스템 접근 권한 필요\n')
    spec.add_run('• macOS: 시스템 통합 권한 필요할 수 있음\n\n')

    doc.save('MSYS_CPU_Monitor_Manual.docx')
    print('CPU 모니터링 앱 설명서 생성 완료: MSYS_CPU_Monitor_Manual.docx')

def create_data_flow_debug_guide():
    """데이터 호출 흐름 디버깅 가이드 생성"""
    doc = Document()

    # 스타일 설정
    title_style = doc.styles.add_style('DebugTitle', WD_STYLE_TYPE.PARAGRAPH)
    title_style.font.size = Pt(28)
    title_style.font.color.rgb = RGBColor(0, 0, 0)
    title_style.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    heading1_style = doc.styles.add_style('DebugH1', WD_STYLE_TYPE.PARAGRAPH)
    heading1_style.font.size = Pt(18)
    heading1_style.font.bold = True

    heading2_style = doc.styles.add_style('DebugH2', WD_STYLE_TYPE.PARAGRAPH)
    heading2_style.font.size = Pt(14)
    heading2_style.font.bold = True

    heading3_style = doc.styles.add_style('DebugH3', WD_STYLE_TYPE.PARAGRAPH)
    heading3_style.font.size = Pt(12)
    heading3_style.font.bold = True

    code_style = doc.styles.add_style('DebugCode', WD_STYLE_TYPE.PARAGRAPH)
    code_style.font.name = 'Courier New'
    code_style.font.size = Pt(10)

    # 표지
    title = doc.add_paragraph('MSYS 데이터 호출 흐름 디버깅 가이드', style='DebugTitle')
    subtitle = doc.add_paragraph('HTML 중심 데이터 호출 추적 및 문제 해결 가이드', style='DebugTitle')
    version = doc.add_paragraph('버전 1.0', style='DebugTitle')
    version.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    doc.add_page_break()

    # 목차
    toc_title = doc.add_paragraph('목차', style='DebugH1')
    toc = doc.add_paragraph()
    toc.add_run('1. 개요 ........................ 3\n')
    toc.add_run('2. 메뉴 페이지별 카드 컴포넌트 데이터 호출 가이드 ................ 4\n')
    toc.add_run('   2.1 대시보드 페이지 .................. 4\n')
    toc.add_run('   2.2 데이터 분석 페이지 ................ 15\n')
    toc.add_run('   2.3 차트 분석 페이지 .................. 25\n')
    toc.add_run('   2.4 카드 요약 페이지 .................. 35\n')
    toc.add_run('   2.5 관리자 설정 페이지 ................ 45\n')
    toc.add_run('3. 인증 시스템 ................ 65\n')
    toc.add_run('   3.1 회원가입 프로세스 .................. 65\n')
    toc.add_run('   3.2 로그인 프로세스 .................. 75\n')
    toc.add_run('   3.3 비밀번호 변경 프로세스 ................ 85\n')
    toc.add_run('   3.4 권한 인증 및 검증 프로세스 ................ 95\n')
    toc.add_run('4. 디버깅 가이드 ................ 105\n')
    toc.add_run('부록 A. 데이터 흐름 요약 ................ 110\n')

    doc.add_page_break()

    # 1. 개요
    doc.add_paragraph('1. 개요', style='DebugH1')
    overview = doc.add_paragraph()
    overview.add_run('1.1 목적\n').bold = True
    overview.add_run('이 가이드는 각 메뉴 페이지의 카드(컴포넌트)를 기준으로 데이터 호출 과정을 TB→DAO→Service→Route→HTML 역순으로 추적하여 설명합니다. 각 단계별 권한 검증, 에러 처리, 예상 결과 등을 포함하여 문제가 발생했을 때 신속한 디버깅을 지원합니다.\n\n')

    overview.add_run('1.2 대상 독자\n').bold = True
    overview.add_run('• 개발자: 데이터 호출 흐름 이해 및 디버깅\n')
    overview.add_run('• 운영자: 시스템 문제 해결 및 모니터링\n')
    overview.add_run('• 테스터: 데이터 흐름 검증 및 테스트 케이스 작성\n\n')

    overview.add_run('1.3 사용 방법\n').bold = True
    overview.add_run('1. 문제가 발생한 메뉴 페이지와 카드 컴포넌트 확인\n')
    overview.add_run('2. 해당 섹션에서 HTML → JavaScript → Route → Service → DAO → DB 순으로 추적\n')
    overview.add_run('3. 각 단계별 권한 검증과 에러 처리 확인\n')
    overview.add_run('4. 문제 지점 식별 및 해결\n\n')

    # 2. 메뉴 페이지별 카드 컴포넌트 데이터 호출 가이드
    doc.add_paragraph('2. 메뉴 페이지별 카드 컴포넌트 데이터 호출 가이드', style='DebugH1')

    # 2.1 대시보드 페이지
    doc.add_paragraph('2.1 대시보드 페이지', style='DebugH2')

    # 요약 카드 컴포넌트
    doc.add_paragraph('2.1.1 요약 카드 컴포넌트 (dashboard_data)', style='DebugH3')
    dashboard_desc = doc.add_paragraph()
    dashboard_desc.add_run('HTML 템플릿 요구사항:\n').bold = True
    dashboard_desc.add_run('templates/dashboard.html에서 dashboard_data 리스트의 각 항목에 cd_nm(Job 이름), success_rate(성공률) 등이 필요\n\n')

    dashboard_desc.add_run('JavaScript API 호출:\n').bold = True
    dashboard_desc.add_run('loadDashboardData() 함수가 /api/dashboard/summary API 호출\n\n')

    dashboard_desc.add_run('권한 검증:\n').bold = True
    dashboard_desc.add_run('• 로그인 상태 확인\n')
    dashboard_desc.add_run('• dashboard 메뉴 권한 체크\n')
    dashboard_desc.add_run('• Job ID별 데이터 접근 권한 필터링\n\n')

    dashboard_desc.add_run('Route/Service/DAO/DB 처리:\n').bold = True
    dashboard_desc.add_run('dashboard_api.summary() → dashboard_service.get_summary() → dashboard_mapper.get_summary() → tb_con_hist + tb_mngr_sett\n\n')

    # 상세 테이블 카드 컴포넌트
    doc.add_paragraph('2.1.2 상세 테이블 카드 컴포넌트 (dashboard_table_data)', style='DebugH3')
    table_desc = doc.add_paragraph()
    table_desc.add_run('동일한 dashboard_data를 테이블 형식으로 표시하는 컴포넌트\n\n')

    # 2.2 데이터 분석 페이지
    doc.add_paragraph('2.2 데이터 분석 페이지', style='DebugH2')

    # 요약 통계 카드
    doc.add_paragraph('2.2.1 요약 통계 카드 컴포넌트 (summary_data)', style='DebugH3')
    analysis_desc = doc.add_paragraph()
    analysis_desc.add_run('HTML 템플릿 요구사항:\n').bold = True
    analysis_desc.add_run('total_count(총 건수), avg_success_rate(평균 성공률) 객체 필요\n\n')

    analysis_desc.add_run('JavaScript API 호출:\n').bold = True
    analysis_desc.add_run('fetchSummaryData() → /api/analysis/summary\n\n')

    analysis_desc.add_run('데이터 흐름:\n').bold = True
    analysis_desc.add_run('analysis_api.get_summary() → analysis_service.get_summary_data() → analysis_mapper.get_summary_stats() → tb_con_hist\n\n')

    # 추이 차트 카드
    doc.add_paragraph('2.2.2 추이 차트 카드 컴포넌트 (trend_data)', style='DebugH3')
    trend_desc = doc.add_paragraph()
    trend_desc.add_run('시간에 따른 데이터 변화를 차트로 표시하는 컴포넌트\n\n')

    # 원천 데이터 테이블 카드
    doc.add_paragraph('2.2.3 원천 데이터 테이블 카드 컴포넌트 (raw_data)', style='DebugH3')
    raw_desc = doc.add_paragraph()
    raw_desc.add_run('실제 수집 데이터를 테이블로 표시하는 컴포넌트\n\n')

    # 2.3 차트 분석 페이지
    doc.add_paragraph('2.3 차트 분석 페이지', style='DebugH2')

    # 성공률 추이 차트
    doc.add_paragraph('2.3.1 성공률 추이 차트 카드 (chart_trend_data)', style='DebugH3')
    chart_trend_desc = doc.add_paragraph()
    chart_trend_desc.add_run('시간에 따른 Job별 성공률 추이를 라인 차트로 표시\n\n')

    # 장애 코드 현황 차트
    doc.add_paragraph('2.3.2 장애 코드 현황 차트 카드 (error_stats_data)', style='DebugH3')
    error_desc = doc.add_paragraph()
    error_desc.add_run('장애 코드별 발생 빈도를 파이 차트로 표시\n\n')

    # 2.4 카드 요약 페이지
    doc.add_paragraph('2.4 카드 요약 페이지', style='DebugH2')

    # 그룹별 요약 카드
    doc.add_paragraph('2.4.1 그룹별 요약 카드 컴포넌트 (card_summary_data)', style='DebugH3')
    card_desc = doc.add_paragraph()
    card_desc.add_run('Job ID별 그룹화된 요약 정보를 카드 형태로 표시\n\n')

    # 2.5 관리자 설정 페이지
    doc.add_paragraph('2.5 관리자 설정 페이지', style='DebugH2')

    # 각 탭별 카드 컴포넌트들
    tabs = [
        ('2.5.1 기본 설정 탭 카드 (admin_basic_data)', '임계값 및 기본 설정 관리'),
        ('2.5.2 수집 스케줄 설정 탭 카드 (admin_schedule_data)', '데이터 수집 스케줄 설정'),
        ('2.5.3 아이콘 관리 탭 카드 (admin_icon_data)', '시스템 아이콘 관리'),
        ('2.5.4 차트/시각화 설정 탭 카드 (admin_chart_data)', '차트 색상 및 시각화 설정'),
        ('2.5.5 사용자 관리 탭 카드 (admin_user_data)', '사용자 계정 및 권한 관리'),
        ('2.5.6 데이터 접근 권한 탭 카드 (admin_permission_data)', 'Job별 데이터 접근 권한 설정'),
        ('2.5.7 통계 탭 카드 (admin_stats_data)', '시스템 사용 통계 조회')
    ]

    for tab_title, tab_desc in tabs:
        doc.add_paragraph(tab_title, style='DebugH3')
        tab_para = doc.add_paragraph()
        tab_para.add_run(f'{tab_desc}\n\n')

    # 3. 인증 시스템
    doc.add_paragraph('3. 인증 시스템', style='DebugH1')

    # 3.1 회원가입 프로세스
    doc.add_paragraph('3.1 회원가입 프로세스', style='DebugH2')
    register_desc = doc.add_paragraph()
    register_desc.add_run('HTML 템플릿 요구사항:\n').bold = True
    register_desc.add_run('사용자 ID, 비밀번호, 비밀번호 확인 입력 폼\n\n')

    register_desc.add_run('JavaScript API 호출:\n').bold = True
    register_desc.add_run('fetch(\'/api/auth/register\', POST) - 사용자 등록 요청\n\n')

    register_desc.add_run('데이터 흐름:\n').bold = True
    register_desc.add_run('auth_routes.register() → user_service.create_user() → user_dao.create_user() → tb_user\n\n')

    register_desc.add_run('권한 검증:\n').bold = True
    register_desc.add_run('• 비밀번호 정책 검증\n')
    register_desc.add_run('• 사용자 ID 중복 체크\n')
    register_desc.add_run('• PENDING 상태로 저장\n\n')

    # 3.2 로그인 프로세스
    doc.add_paragraph('3.2 로그인 프로세스', style='DebugH2')
    login_desc = doc.add_paragraph()
    login_desc.add_run('HTML 템플릿 요구사항:\n').bold = True
    login_desc.add_run('사용자 ID, 비밀번호 입력 폼\n\n')

    login_desc.add_run('JavaScript API 호출:\n').bold = True
    login_desc.add_run('fetch(\'/api/auth/login\', POST) - 인증 요청\n\n')

    login_desc.add_run('데이터 흐름:\n').bold = True
    login_desc.add_run('auth_routes.login() → user_service.authenticate_user() → user_dao.get_user_by_id() → tb_user\n\n')

    login_desc.add_run('권한 검증:\n').bold = True
    login_desc.add_run('• 계정 상태 확인 (APPROVED/PENDING/REJECTED/SUSPENDED)\n')
    login_desc.add_run('• 세션 생성 및 권한 정보 로드\n\n')

    # 3.3 비밀번호 변경 프로세스
    doc.add_paragraph('3.3 비밀번호 변경 프로세스', style='DebugH2')
    pw_desc = doc.add_paragraph()
    pw_desc.add_run('HTML 템플릿 요구사항:\n').bold = True
    pw_desc.add_run('현재 비밀번호, 새 비밀번호, 확인 입력 폼\n\n')

    pw_desc.add_run('JavaScript API 호출:\n').bold = True
    pw_desc.add_run('fetch(\'/api/auth/change_password\', POST) - 비밀번호 변경 요청\n\n')

    pw_desc.add_run('데이터 흐름:\n').bold = True
    pw_desc.add_run('auth_routes.change_password() → user_service.update_user_password() → user_dao.update_user_password() → tb_user\n\n')

    pw_desc.add_run('권한 검증:\n').bold = True
    pw_desc.add_run('• 현재 비밀번호 검증\n')
    pw_desc.add_run('• 새 비밀번호 정책 확인\n')
    pw_desc.add_run('• 세션 갱신\n\n')

    # 3.4 권한 인증 및 검증 프로세스
    doc.add_paragraph('3.4 권한 인증 및 검증 프로세스', style='DebugH2')
    auth_desc = doc.add_paragraph()
    auth_desc.add_run('미들웨어 적용:\n').bold = True
    auth_desc.add_run('@app.before_request - 모든 요청에 대한 권한 검증\n\n')

    auth_desc.add_run('검증 단계:\n').bold = True
    auth_desc.add_run('1. 로그인 상태 확인\n')
    auth_desc.add_run('2. 계정 상태 검증 (APPROVED)\n')
    auth_desc.add_run('3. 메뉴 권한 체크\n')
    auth_desc.add_run('4. 데이터 접근 권한 필터링\n\n')

    auth_desc.add_run('데이터 흐름:\n').bold = True
    auth_desc.add_run('auth_middleware.check_permissions() → user_service.get_user_permissions() → user_dao.get_user_menu_permissions() → tb_user_auth_ctrl\n\n')

    # 4. 디버깅 가이드
    doc.add_paragraph('4. 디버깅 가이드', style='DebugH1')

    # 4.1 데이터가 표시되지 않는 경우
    doc.add_paragraph('4.1 데이터가 표시되지 않는 경우', style='DebugH2')
    debug1 = doc.add_paragraph()
    debug1.add_run('1. 브라우저 개발자 도구 Network 탭에서 API 호출 상태 확인\n')
    debug1.add_run('2. 403/401 에러 시 권한 문제 검토\n')
    debug1.add_run('3. 로그 파일에서 권한 관련 메시지 확인\n')
    debug1.add_run('4. 해당 카드 컴포넌트의 HTML 템플릿 변수 검증\n\n')

    # 4.2 권한 문제 해결
    doc.add_paragraph('4.2 권한 문제 해결', style='DebugH2')
    debug2 = doc.add_paragraph()
    debug2.add_run('1. 사용자 관리 메뉴에서 권한 설정 확인\n')
    debug2.add_run('2. 데이터 접근 권한 탭에서 Job ID 권한 확인\n')
    debug2.add_run('3. 로그에서 "권한 없음" 메시지 확인\n')
    debug2.add_run('4. 미들웨어 권한 체크 로직 검증\n\n')

    # 4.3 데이터 이상 시
    doc.add_paragraph('4.3 데이터 이상 시', style='DebugH2')
    debug3 = doc.add_paragraph()
    debug3.add_run('1. DAO 메소드의 SQL 쿼리 확인\n')
    debug3.add_run('2. DB에서 직접 쿼리 실행하여 결과 확인\n')
    debug3.add_run('3. Service 레벨의 데이터 가공 로직 검증\n')
    debug3.add_run('4. Route 레벨의 파라미터 처리 확인\n\n')

    # 부록 A. 데이터 흐름 요약
    doc.add_paragraph('부록 A. 데이터 흐름 요약', style='DebugH1')
    summary = doc.add_paragraph()
    summary.add_run('A.1 일반적인 데이터 호출 패턴\n').bold = True
    summary.add_run('HTML (템플릿 변수) → JavaScript (API 호출) → Route (권한 체크) → Service (비즈니스 로직) → DAO (SQL 실행) → DB (데이터 저장소)\n\n')

    summary.add_run('A.2 권한 검증 포인트\n').bold = True
    summary.add_run('• Route 레벨: 메뉴 접근 권한\n')
    summary.add_run('• Service 레벨: 데이터 접근 권한 필터링\n')
    summary.add_run('• 미들웨어: 세션 및 계정 상태 검증\n\n')

    summary.add_run('A.3 에러 처리 패턴\n').bold = True
    summary.add_run('• 401: 로그인 필요\n')
    summary.add_run('• 403: 권한 부족\n')
    summary.add_run('• 500: 서버 내부 오류\n')
    summary.add_run('• 빈 배열: 데이터 접근 권한 없음\n\n')

    try:
        doc.save('MSYS_Data_Flow_Debug_Guide.docx')
        print('데이터 호출 흐름 디버깅 가이드 생성 완료: MSYS_Data_Flow_Debug_Guide.docx')
    except PermissionError:
        doc.save('MSYS_Data_Flow_Debug_Guide_v2.docx')
        print('데이터 호출 흐름 디버깅 가이드 생성 완료: MSYS_Data_Flow_Debug_Guide_v2.docx')

def create_complete_function_manual():
    """전체 MSYS 기능 매뉴얼 생성 - 각 페이지별 상세 기능 설명"""
    doc = Document()

    # 스타일 설정
    title_style = doc.styles.add_style('CompleteTitle', WD_STYLE_TYPE.PARAGRAPH)
    title_style.font.size = Pt(28)
    title_style.font.color.rgb = RGBColor(0, 0, 0)
    title_style.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    heading1_style = doc.styles.add_style('CompleteH1', WD_STYLE_TYPE.PARAGRAPH)
    heading1_style.font.size = Pt(18)
    heading1_style.font.bold = True

    heading2_style = doc.styles.add_style('CompleteH2', WD_STYLE_TYPE.PARAGRAPH)
    heading2_style.font.size = Pt(14)
    heading2_style.font.bold = True

    heading3_style = doc.styles.add_style('CompleteH3', WD_STYLE_TYPE.PARAGRAPH)
    heading3_style.font.size = Pt(12)
    heading3_style.font.bold = True

    code_style = doc.styles.add_style('CompleteCode', WD_STYLE_TYPE.PARAGRAPH)
    code_style.font.name = 'Courier New'
    code_style.font.size = Pt(10)

    # 표지
    title = doc.add_paragraph('MSYS 전체 기능 매뉴얼', style='CompleteTitle')
    subtitle = doc.add_paragraph('시스템 기능 및 사용자 가이드', style='CompleteTitle')
    version = doc.add_paragraph('버전 1.14.2', style='CompleteTitle')
    version.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    doc.add_page_break()

    # 목차
    toc_title = doc.add_paragraph('목차', style='CompleteH1')
    toc = doc.add_paragraph()
    toc.add_run('1. 시스템 개요 ........................ 3\n')
    toc.add_run('2. 메뉴 구조 ........................ 5\n')
    toc.add_run('3. 대시보드 ........................ 7\n')
    toc.add_run('   3.1 개요 ...................... 7\n')
    toc.add_run('   3.2 화면 구성 .................. 8\n')
    toc.add_run('   3.3 주요 기능 .................. 12\n')
    toc.add_run('   3.4 사용자 인터랙션 .............. 15\n')
    toc.add_run('4. 데이터 분석 ...................... 18\n')
    toc.add_run('   4.1 개요 ...................... 18\n')
    toc.add_run('   4.2 화면 구성 .................. 19\n')
    toc.add_run('   4.3 주요 기능 .................. 22\n')
    toc.add_run('   4.4 사용자 인터랙션 .............. 25\n')
    toc.add_run('5. 차트 분석 ...................... 28\n')
    toc.add_run('   5.1 개요 ...................... 28\n')
    toc.add_run('   5.2 화면 구성 .................. 29\n')
    toc.add_run('   5.3 주요 기능 .................. 32\n')
    toc.add_run('   5.4 사용자 인터랙션 .............. 35\n')
    toc.add_run('6. 잔디 현황 ...................... 38\n')
    toc.add_run('   6.1 개요 ...................... 38\n')
    toc.add_run('   6.2 화면 구성 .................. 39\n')
    toc.add_run('   6.3 주요 기능 .................. 42\n')
    toc.add_run('   6.4 사용자 인터랙션 .............. 45\n')
    toc.add_run('7. 매핑 관리 ...................... 48\n')
    toc.add_run('   7.1 개요 ...................... 48\n')
    toc.add_run('   7.2 화면 구성 .................. 49\n')
    toc.add_run('   7.3 주요 기능 .................. 52\n')
    toc.add_run('   7.4 사용자 인터랙션 .............. 55\n')
    toc.add_run('8. 데이터 명세서 .................... 58\n')
    toc.add_run('   8.1 개요 ...................... 58\n')
    toc.add_run('   8.2 화면 구성 .................. 59\n')
    toc.add_run('   8.3 주요 기능 .................. 62\n')
    toc.add_run('   8.4 사용자 인터랙션 .............. 65\n')
    toc.add_run('9. 관리자 설정 .................... 68\n')
    toc.add_run('   9.1 개요 ...................... 68\n')
    toc.add_run('   9.2 기본 설정 탭 ................ 69\n')
    toc.add_run('   9.3 수집 스케줄 설정 탭 ............ 75\n')
    toc.add_run('   9.4 아이콘 관리 탭 ................ 80\n')
    toc.add_run('   9.5 차트/시각화 설정 탭 ............ 85\n')
    toc.add_run('   9.6 사용자 관리 탭 ................ 90\n')
    toc.add_run('   9.7 데이터 접근 권한 탭 ............ 95\n')
    toc.add_run('   9.8 통계 탭 .................... 100\n')
    toc.add_run('10. 수집 일정 ..................... 105\n')
    toc.add_run('    10.1 개요 ..................... 105\n')
    toc.add_run('    10.2 화면 구성 ................... 106\n')
    toc.add_run('    10.3 주요 기능 ................... 109\n')
    toc.add_run('    10.4 사용자 인터랙션 ............... 112\n')
    toc.add_run('11. 카드 요약 ..................... 115\n')
    toc.add_run('    11.1 개요 ..................... 115\n')
    toc.add_run('    11.2 화면 구성 ................... 116\n')
    toc.add_run('    11.3 주요 기능 ................... 119\n')
    toc.add_run('    11.4 사용자 인터랙션 ............... 122\n')
    toc.add_run('12. 인증 시스템 ................... 125\n')
    toc.add_run('    12.1 로그인 ..................... 125\n')
    toc.add_run('    12.2 회원가입 ................... 130\n')
    toc.add_run('    12.3 비밀번호 변경 ................. 135\n')
    toc.add_run('부록 A. 단축키 및 팁 ................ 140\n')

    doc.add_page_break()

    # 1. 시스템 개요
    doc.add_paragraph('1. 시스템 개요', style='CompleteH1')
    overview = doc.add_paragraph()
    overview.add_run('1.1 MSYS 소개\n').bold = True
    overview.add_run('MSYS(Monitoring SYStem)는 데이터 수집 및 모니터링을 위한 웹 기반 시스템입니다. 다양한 데이터 소스로부터 주기적으로 데이터를 수집하고, 실시간 모니터링 및 분석 기능을 제공합니다.\n\n')

    overview.add_run('1.2 주요 특징\n').bold = True
    overview.add_run('• 실시간 데이터 수집 모니터링\n')
    overview.add_run('• 다양한 시각화 차트 및 그래프\n')
    overview.add_run('• 사용자 권한 기반 접근 제어\n')
    overview.add_run('• 자동화된 데이터 수집 스케줄링\n')
    overview.add_run('• 관리자 설정을 통한 시스템 커스터마이징\n\n')

    overview.add_run('1.3 시스템 요구사항\n').bold = True
    overview.add_run('• 웹 브라우저: Chrome, Firefox, Safari, Edge\n')
    overview.add_run('• 권장 해상도: 1920x1080 이상\n')
    overview.add_run('• JavaScript 활성화 필수\n\n')

    # 2. 메뉴 구조
    doc.add_paragraph('2. 메뉴 구조', style='CompleteH1')
    menu_structure = doc.add_paragraph()
    menu_structure.add_run('2.1 메인 메뉴\n').bold = True
    menu_structure.add_run('시스템은 다음과 같은 메뉴 구조로 구성되어 있습니다:\n\n')

    menu_table = doc.add_table(rows=1, cols=3)
    menu_table.style = 'Table Grid'
    hdr_cells = menu_table.rows[0].cells
    hdr_cells[0].text = '메뉴 이름'
    hdr_cells[1].text = 'URL 경로'
    hdr_cells[2].text = '설명'

    menus = [
        ('대시보드', '/dashboard', '시스템 현황 실시간 모니터링'),
        ('데이터 분석', '/data_analysis', '수집 데이터 심층 분석'),
        ('차트 분석', '/chart_analysis', '시각화된 데이터 추이 분석'),
        ('잔디 현황', '/jandi', '히트맵 형태의 활동 현황'),
        ('매핑 관리', '/mapping', '데이터베이스 컬럼 매핑 관리'),
        ('데이터 명세서', '/data_spec', '외부 API 명세서 관리'),
        ('관리자 설정', '/mngr_sett', '시스템 설정 및 관리'),
        ('수집 일정', '/collection_schedule', '데이터 수집 스케줄 모니터링'),
        ('카드 요약', '/card_summary', '그룹별 수집 현황 요약')
    ]

    for menu_name, menu_url, menu_desc in menus:
        row_cells = menu_table.add_row().cells
        row_cells[0].text = menu_name
        row_cells[1].text = menu_url
        row_cells[2].text = menu_desc

    menu_structure.add_run('\n\n2.2 권한별 메뉴 접근\n').bold = True
    menu_structure.add_run('각 사용자는 부여된 권한에 따라 접근 가능한 메뉴가 결정됩니다.\n\n')

    # 3. 대시보드
    doc.add_paragraph('3. 대시보드', style='CompleteH1')

    # 3.1 개요
    doc.add_paragraph('3.1 개요', style='CompleteH2')
    dash_overview = doc.add_paragraph()
    dash_overview.add_run('대시보드는 MSYS 시스템의 메인 페이지로, 전체 데이터 수집 현황을 실시간으로 모니터링할 수 있는 화면입니다. 각 Job별 성공률, 연속 실패 횟수, 현재 상태 등을 한눈에 파악할 수 있습니다.\n\n')

    dash_overview.add_run('주요 목적:\n').bold = True
    dash_overview.add_run('• 시스템 전체 현황 파악\n')
    dash_overview.add_run('• 문제 발생 Job 신속 식별\n')
    dash_overview.add_run('• 데이터 수집 상태 모니터링\n\n')

    # 3.2 화면 구성
    doc.add_paragraph('3.2 화면 구성', style='CompleteH2')
    dash_composition = doc.add_paragraph()
    dash_composition.add_run('3.2.1 레이아웃 구조\n').bold = True
    dash_composition.add_run('대시보드는 다음과 같은 영역으로 구성되어 있습니다:\n\n')

    dash_composition.add_run('1. 헤더 영역\n').bold = True
    dash_composition.add_run('• 시스템 로고 및 타이틀\n')
    dash_composition.add_run('• 사용자 정보 및 로그아웃 버튼\n')
    dash_composition.add_run('• 현재 시간 표시\n\n')

    dash_composition.add_run('2. 요약 카드 영역\n').bold = True
    dash_composition.add_run('• 전체 Job 개수 표시\n')
    dash_composition.add_run('• 총 수집 건수 표시\n')
    dash_composition.add_run('• 평균 성공률 표시\n\n')

    dash_composition.add_run('3. 상세 테이블 영역\n').bold = True
    dash_composition.add_run('• Job별 상세 정보 테이블\n')
    dash_composition.add_run('• 실시간 상태 표시\n')
    dash_composition.add_run('• 정렬 및 필터링 기능\n\n')

    dash_composition.add_run('4. 컨트롤 패널\n').bold = True
    dash_composition.add_run('• 기간 선택 (시작일/종료일)\n')
    dash_composition.add_run('• 조회 버튼\n')
    dash_composition.add_run('• 자동 새로고침 토글\n\n')

    # 3.3 주요 기능
    doc.add_paragraph('3.3 주요 기능', style='CompleteH2')
    dash_features = doc.add_paragraph()
    dash_features.add_run('3.3.1 실시간 모니터링\n').bold = True
    dash_features.add_run('• 5초 간격 자동 새로고침\n')
    dash_features.add_run('• 실시간 데이터 수집 상태 반영\n')
    dash_features.add_run('• 색상 코드를 통한 상태 구분\n\n')

    dash_features.add_run('3.3.2 상태 표시 시스템\n').bold = True
    dash_features.add_run('• 녹색: 정상 상태 (성공률 ≥ 임계값)\n')
    dash_features.add_run('• 노랑: 경고 상태 (성공률 80-95%)\n')
    dash_features.add_run('• 빨강: 위험 상태 (성공률 < 80%)\n\n')

    dash_features.add_run('3.3.3 데이터 필터링\n').bold = True
    dash_features.add_run('• 기간별 조회 (일/주/월 단위)\n')
    dash_features.add_run('• Job별 권한 기반 필터링\n')
    dash_features.add_run('• 실시간 검색 및 정렬\n\n')

    # 3.4 사용자 인터랙션
    doc.add_paragraph('3.4 사용자 인터랙션', style='CompleteH2')
    dash_interaction = doc.add_paragraph()
    dash_interaction.add_run('3.4.1 기본 조작\n').bold = True
    dash_interaction.add_run('1. 페이지 접속 시 자동으로 최신 데이터 로드\n')
    dash_interaction.add_run('2. 기간 선택 후 [조회] 버튼 클릭으로 데이터 필터링\n')
    dash_interaction.add_run('3. 테이블 헤더 클릭으로 정렬 변경\n')
    dash_interaction.add_run('4. 자동 새로고침 토글로 실시간 모니터링 제어\n\n')

    dash_interaction.add_run('3.4.2 고급 기능\n').bold = True
    dash_interaction.add_run('• 테이블 행 클릭으로 상세 정보 팝업\n')
    dash_interaction.add_run('• Job ID 클릭으로 해당 Job의 상세 분석 페이지 이동\n')
    dash_interaction.add_run('• 우측 클릭으로 컨텍스트 메뉴 표시\n\n')

    # 4. 데이터 분석
    doc.add_paragraph('4. 데이터 분석', style='CompleteH1')

    # 4.1 개요
    doc.add_paragraph('4.1 개요', style='CompleteH2')
    analysis_overview = doc.add_paragraph()
    analysis_overview.add_run('데이터 분석 페이지는 수집된 데이터를 심층적으로 분석하고 필터링하여 조회할 수 있는 기능을 제공합니다. 다양한 조건으로 데이터를 검색하고, 추이 분석 및 통계 정보를 확인할 수 있습니다.\n\n')

    analysis_overview.add_run('주요 목적:\n').bold = True
    analysis_overview.add_run('• 데이터 수집 결과 상세 분석\n')
    analysis_overview.add_run('• 문제 발생 패턴 식별\n')
    analysis_overview.add_run('• 데이터 품질 검증\n\n')

    # 4.2 화면 구성
    doc.add_paragraph('4.2 화면 구성', style='CompleteH2')
    analysis_composition = doc.add_paragraph()
    analysis_composition.add_run('4.2.1 탭 메뉴 구조\n').bold = True
    analysis_composition.add_run('데이터 분석 페이지는 다음과 같은 탭으로 구성되어 있습니다:\n\n')

    analysis_composition.add_run('• 요약 데이터: 전체 통계 요약\n').bold = True
    analysis_composition.add_run('• 추이 데이터: 시간별 변화 그래프\n')
    analysis_composition.add_run('• 원천 데이터: 실제 수집 데이터 목록\n')
    analysis_composition.add_run('• Job 정보: 각 Job의 상세 정보\n\n')

    analysis_composition.add_run('4.2.2 필터 컨트롤\n').bold = True
    analysis_composition.add_run('• 기간 선택: 시작일과 종료일 설정\n')
    analysis_composition.add_run('• Job ID 선택: 다중 선택 가능\n')
    analysis_composition.add_run('• 장애코드 필터: 특정 오류 유형 필터링\n')
    analysis_composition.add_run('• 상태 필터: 성공/실패 상태 필터링\n\n')

    # 4.3 주요 기능
    doc.add_paragraph('4.3 주요 기능', style='CompleteH2')
    analysis_features = doc.add_paragraph()
    analysis_features.add_run('4.3.1 다중 조건 필터링\n').bold = True
    analysis_features.add_run('• 기간, Job ID, 상태, 장애코드 등 복합 조건 검색\n')
    analysis_features.add_run('• 실시간 필터 적용 및 결과 표시\n')
    analysis_features.add_run('• 필터 조건 저장 및 재사용\n\n')

    analysis_features.add_run('4.3.2 데이터 내보내기\n').bold = True
    analysis_features.add_run('• CSV 형식 데이터 다운로드\n')
    analysis_features.add_run('• Excel 형식 보고서 생성\n')
    analysis_features.add_run('• 필터링된 결과만 선택적 내보내기\n\n')

    analysis_features.add_run('4.3.3 AI 분석 기능\n').bold = True
    analysis_features.add_run('• 자동화된 데이터 패턴 분석\n')
    analysis_features.add_run('• 이상 징후 자동 감지\n')
    analysis_features.add_run('• 분석 결과 텍스트 리포트 생성\n\n')

    # 4.4 사용자 인터랙션
    doc.add_paragraph('4.4 사용자 인터랙션', style='CompleteH2')
    analysis_interaction = doc.add_paragraph()
    analysis_interaction.add_run('4.4.1 기본 분석 절차\n').bold = True
    analysis_interaction.add_run('1. 분석할 기간 선택\n')
    analysis_interaction.add_run('2. 관심 있는 Job ID 선택\n')
    analysis_interaction.add_run('3. 필요한 필터 조건 설정\n')
    analysis_interaction.add_run('4. [조회] 버튼으로 데이터 로드\n')
    analysis_interaction.add_run('5. 각 탭에서 분석 결과 확인\n\n')

    analysis_interaction.add_run('4.4.2 고급 분석 기능\n').bold = True
    analysis_interaction.add_run('• [AI 분석] 버튼으로 자동 분석 실행\n')
    analysis_interaction.add_run('• 차트 위 마우스 오버로 상세 값 확인\n')
    analysis_interaction.add_run('• 데이터 행 클릭으로 상세 정보 팝업\n\n')

    # 5. 차트 분석
    doc.add_paragraph('5. 차트 분석', style='CompleteH1')

    # 5.1 개요
    doc.add_paragraph('5.1 개요', style='CompleteH2')
    chart_overview = doc.add_paragraph()
    chart_overview.add_run('차트 분석 페이지는 데이터 수집 성공률의 시간적 추이와 장애 코드별 현황을 시각적인 차트로 분석하는 기능을 제공합니다. 다양한 차트 유형을 통해 데이터 패턴을 직관적으로 파악할 수 있습니다.\n\n')

    chart_overview.add_run('주요 목적:\n').bold = True
    chart_overview.add_run('• 데이터 추이 시각화\n')
    chart_overview.add_run('• 장애 패턴 분석\n')
    chart_overview.add_run('• 성능 모니터링\n\n')

    # 5.2 화면 구성
    doc.add_paragraph('5.2 화면 구성', style='CompleteH2')
    chart_composition = doc.add_paragraph()
    chart_composition.add_run('5.2.1 차트 영역\n').bold = True
    chart_composition.add_run('• 성공률 추이 라인 차트\n')
    chart_composition.add_run('• 장애 코드 현황 파이 차트\n')
    chart_composition.add_run('• 인터랙티브 차트 컨트롤\n\n')

    chart_composition.add_run('5.2.2 컨트롤 패널\n').bold = True
    chart_composition.add_run('• 기간 선택 드롭다운\n')
    chart_composition.add_run('• Job ID 체크박스 그룹\n')
    chart_composition.add_run('• 차트 옵션 설정\n\n')

    # 5.3 주요 기능
    doc.add_paragraph('5.3 주요 기능', style='CompleteH2')
    chart_features = doc.add_paragraph()
    chart_features.add_run('5.3.1 다중 Job 비교\n').bold = True
    chart_features.add_run('• 여러 Job의 성공률 동시 비교\n')
    chart_features.add_run('• 범례 클릭으로 개별 Job 표시/숨김\n')
    chart_features.add_run('• 색상 코드를 통한 Job 구분\n\n')

    chart_features.add_run('5.3.2 인터랙티브 기능\n').bold = True
    chart_features.add_run('• 마우스 오버로 상세 값 표시\n')
    chart_features.add_run('• 줌 인/아웃 기능\n')
    chart_features.add_run('• 데이터 포인트 클릭으로 상세 정보\n\n')

    # 5.4 사용자 인터랙션
    doc.add_paragraph('5.4 사용자 인터랙션', style='CompleteH2')
    chart_interaction = doc.add_paragraph()
    chart_interaction.add_run('5.4.1 차트 분석 절차\n').bold = True
    chart_interaction.add_run('1. 분석할 기간 선택\n')
    chart_interaction.add_run('2. 비교할 Job ID 체크\n')
    chart_interaction.add_run('3. [조회] 버튼으로 차트 갱신\n')
    chart_interaction.add_run('4. 차트 위에서 마우스 조작으로 상호작용\n\n')

    # 6. 잔디 현황
    doc.add_paragraph('6. 잔디 현황', style='CompleteH1')

    # 6.1 개요
    doc.add_paragraph('6.1 개요', style='CompleteH2')
    jandi_overview = doc.add_paragraph()
    jandi_overview.add_run('잔디 현황 페이지는 Github의 잔디밭과 유사하게 각 Job의 일별 데이터 수집 현황을 히트맵 형태로 시각화하여 보여줍니다. 직관적인 색상 코드를 통해 장기간 데이터 수집 패턴을 한눈에 파악할 수 있습니다.\n\n')

    jandi_overview.add_run('주요 목적:\n').bold = True
    jandi_overview.add_run('• 장기간 활동 패턴 파악\n')
    jandi_overview.add_run('• 데이터 수집 일관성 모니터링\n')
    jandi_overview.add_run('• 휴일 및 특이일 패턴 분석\n\n')

    # 6.2 화면 구성
    doc.add_paragraph('6.2 화면 구성', style='CompleteH2')
    jandi_composition = doc.add_paragraph()
    jandi_composition.add_run('6.2.1 히트맵 영역\n').bold = True
    jandi_composition.add_run('• 월별 캘린더 형태 히트맵\n')
    jandi_composition.add_run('• 색상 강도로 활동량 표시\n')
    jandi_composition.add_run('• 날짜별 상세 정보 툴팁\n\n')

    jandi_composition.add_run('6.2.2 Job 목록\n').bold = True
    jandi_composition.add_run('• Job ID 및 이름 목록\n')
    jandi_composition.add_run('• 펼치기/접기 토글 버튼\n')
    jandi_composition.add_run('• 각 Job별 히트맵 표시\n\n')

    # 6.3 주요 기능
    doc.add_paragraph('6.3 주요 기능', style='CompleteH2')
    jandi_features = doc.add_paragraph()
    jandi_features.add_run('6.3.1 패턴 분석\n').bold = True
    jandi_features.add_run('• 주중/주말 패턴 구분\n')
    jandi_features.add_run('• 계절별 활동량 변화\n')
    jandi_features.add_run('• 휴일 및 공휴일 영향 분석\n\n')

    jandi_features.add_run('6.3.2 색상 범례\n').bold = True
    jandi_features.add_run('• 0건: 흰색 (미수집)\n')
    jandi_features.add_run('• 1-5건: 연한 녹색 (낮은 활동)\n')
    jandi_features.add_run('• 6-10건: 중간 녹색 (보통 활동)\n')
    jandi_features.add_run('• 11건+: 진한 녹색 (높은 활동)\n\n')

    # 6.4 사용자 인터랙션
    doc.add_paragraph('6.4 사용자 인터랙션', style='CompleteH2')
    jandi_interaction = doc.add_paragraph()
    jandi_interaction.add_run('6.4.1 히트맵 탐색\n').bold = True
    jandi_interaction.add_run('1. 관심 있는 Job 선택\n')
    jandi_interaction.add_run('2. [펼치기] 버튼으로 히트맵 표시\n')
    jandi_interaction.add_run('3. 날짜별 색상으로 활동량 파악\n')
    jandi_interaction.add_run('4. 마우스 오버로 구체적인 수치 확인\n\n')

    # 7. 매핑 관리
    doc.add_paragraph('7. 매핑 관리', style='CompleteH1')

    # 7.1 개요
    doc.add_paragraph('7.1 개요', style='CompleteH2')
    mapping_overview = doc.add_paragraph()
    mapping_overview.add_run('매핑 관리 페이지는 데이터베이스 테이블 간의 컬럼 매핑 정보를 관리하고 조회하는 기능을 제공합니다. 데이터 변환 규칙을 정의하고 관리할 수 있습니다.\n\n')

    mapping_overview.add_run('주요 목적:\n').bold = True
    mapping_overview.add_run('• 데이터 변환 규칙 관리\n')
    mapping_overview.add_run('• 테이블 간 관계 정의\n')
    mapping_overview.add_run('• 데이터 일관성 유지\n\n')

    # 7.2 화면 구성
    doc.add_paragraph('7.2 화면 구성', style='CompleteH2')
    mapping_composition = doc.add_paragraph()
    mapping_composition.add_run('7.2.1 매핑 테이블\n').bold = True
    mapping_composition.add_run('• 소스 테이블/컬럼 정보\n')
    mapping_composition.add_run('• 타겟 테이블/컬럼 정보\n')
    mapping_composition.add('• 매핑 유형 및 규칙\n\n')

    # 7.3 주요 기능
    doc.add_paragraph('7.3 주요 기능', style='CompleteH2')
    mapping_features = doc.add_paragraph()
    mapping_features.add_run('7.3.1 매핑 CRUD\n').bold = True
    mapping_features.add_run('• 새로운 매핑 관계 생성\n')
    mapping_features.add_run('• 기존 매핑 정보 수정\n')
    mapping_features.add_run('• 불필요한 매핑 삭제\n\n')

    # 7.4 사용자 인터랙션
    doc.add_paragraph('7.4 사용자 인터랙션', style='CompleteH2')
    mapping_interaction = doc.add_paragraph()
    mapping_interaction.add_run('7.4.1 매핑 관리 절차\n').bold = True
    mapping_interaction.add_run('1. [새 매핑 추가] 버튼 클릭\n')
    mapping_interaction.add_run('2. 소스 및 타겟 정보 입력\n')
    mapping_interaction.add_run('3. 매핑 규칙 설정\n')
    mapping_interaction.add_run('4. [저장]으로 매핑 생성\n\n')

    # 8. 데이터 명세서
    doc.add_paragraph('8. 데이터 명세서', style='CompleteH1')

    # 8.1 개요
    doc.add_paragraph('8.1 개요', style='CompleteH2')
    spec_overview = doc.add_paragraph()
    spec_overview.add_run('데이터 명세서 페이지는 시스템에서 사용하는 외부 데이터(API 등)의 명세를 관리하는 기능을 제공합니다. API 엔드포인트, 파라미터, 응답 형식 등을 체계적으로 관리할 수 있습니다.\n\n')

    spec_overview.add_run('주요 목적:\n').bold = True
    spec_overview.add_run('• 외부 API 명세 표준화\n')
    spec_overview.add_run('• 데이터 인터페이스 문서화\n')
    spec_overview.add_run('• API 변경 영향도 분석\n\n')

    # 8.2 화면 구성
    doc.add_paragraph('8.2 화면 구성', style='CompleteH2')
    spec_composition = doc.add_paragraph()
    spec_composition.add_run('8.2.1 명세서 목록\n').bold = True
    spec_composition.add_run('• 명세서 이름 및 설명\n')
    spec_composition.add_run('• API URL 및 제공자 정보\n')
    spec_composition.add_run('• 최종 수정일 및 상태\n\n')

    # 8.3 주요 기능
    doc.add_paragraph('8.3 주요 기능', style='CompleteH2')
    spec_features = doc.add_paragraph()
    spec_features.add_run('8.3.1 명세서 관리\n').bold = True
    spec_features.add_run('• 새로운 API 명세서 등록\n')
    spec_features.add_run('• 기존 명세서 정보 수정\n')
    spec_features.add_run('• 명세서 버전 관리\n\n')

    # 8.4 사용자 인터랙션
    doc.add_paragraph('8.4 사용자 인터랙션', style='CompleteH2')
    spec_interaction = doc.add_paragraph()
    spec_interaction.add_run('8.4.1 명세서 관리 절차\n').bold = True
    spec_interaction.add_run('1. [새 명세서] 버튼 클릭\n')
    spec_interaction.add_run('2. API 정보 입력\n')
    spec_interaction.add_run('3. 파라미터 및 응답 형식 정의\n')
    spec_interaction.add_run('4. [저장]으로 명세서 등록\n\n')

    # 9. 관리자 설정
    doc.add_paragraph('9. 관리자 설정', style='CompleteH1')

    # 9.1 개요
    doc.add_paragraph('9.1 개요', style='CompleteH2')
    admin_overview = doc.add_paragraph()
    admin_overview.add_run('관리자 설정 페이지는 시스템의 각종 설정을 관리하고 모니터링하는 기능을 제공합니다. Job별 임계값, 사용자 권한, 시스템 설정 등을 중앙 집중적으로 관리할 수 있습니다.\n\n')

    admin_overview.add_run('주요 목적:\n').bold = True
    admin_overview.add_run('• 시스템 설정 중앙 관리\n')
    admin_overview.add_run('• 사용자 권한 체계적 관리\n')
    admin_overview.add_run('• 모니터링 임계값 설정\n\n')

    # 9.2 기본 설정 탭
    doc.add_paragraph('9.2 기본 설정 탭', style='CompleteH2')
    basic_settings = doc.add_paragraph()
    basic_settings.add_run('9.2.1 기능 개요\n').bold = True
    basic_settings.add_run('Job별 모니터링 임계값과 표시 설정을 관리하는 탭입니다.\n\n')

    basic_settings.add_run('9.2.2 주요 설정 항목\n').bold = True
    basic_settings.add_run('• 연속 실패 임계값: 3회 (기본값)\n')
    basic_settings.add_run('• 일별 성공률 임계값: 80% (기본값)\n')
    basic_settings.add_run('• 주별 성공률 임계값: 75% (기본값)\n')
    basic_settings.add_run('• 월별 성공률 임계값: 70% (기본값)\n')
    basic_settings.add_run('• 상태별 아이콘 및 색상 설정\n\n')

    # 9.3 수집 스케줄 설정 탭
    doc.add_paragraph('9.3 수집 스케줄 설정 탭', style='CompleteH2')
    schedule_settings = doc.add_paragraph()
    schedule_settings.add_run('9.3.1 기능 개요\n').bold = True
    schedule_settings.add_run('데이터 수집 작업의 스케줄을 설정하고 관리하는 탭입니다.\n\n')

    schedule_settings.add_run('9.3.2 주요 설정 항목\n').bold = True
    schedule_settings.add_run('• 그룹 최소 개수: 3 (기본값)\n')
    schedule_settings.add_run('• 그룹 외곽선 스타일: solid (기본값)\n')
    schedule_settings.add_run('• 그룹 색상 기준: prgr/succ\n')
    schedule_settings.add_run('• 진행률/성공률 임계값 설정\n\n')

    # 9.4 아이콘 관리 탭
    doc.add_paragraph('9.4 아이콘 관리 탭', style='CompleteH2')
    icon_settings = doc.add_paragraph()
    icon_settings.add_run('9.4.1 기능 개요\n').bold = True
    icon_settings.add_run('시스템에서 사용하는 아이콘을 관리하는 탭입니다.\n\n')

    icon_settings.add_run('9.4.2 주요 기능\n').bold = True
    icon_settings.add_run('• 아이콘 코드 및 이름 관리\n')
    icon_settings.add_run('• 표시 여부 설정\n')
    icon_settings.add_run('• CSV 일괄 관리\n\n')

    # 9.5 차트/시각화 설정 탭
    doc.add_paragraph('9.5 차트/시각화 설정 탭', style='CompleteH2')
    chart_settings = doc.add_paragraph()
    chart_settings.add_run('9.5.1 기능 개요\n').bold = True
    chart_settings.add_run('차트 및 시각화 요소의 설정을 관리하는 탭입니다.\n\n')

    chart_settings.add_run('9.5.2 주요 설정 항목\n').bold = True
    chart_settings.add_run('• Job별 차트 색상 설정\n')
    chart_settings.add_run('• 잔디 차트 색상 범위 설정\n')
    chart_settings.add_run('• 시각화 옵션 관리\n\n')

    # 9.6 사용자 관리 탭
    doc.add_paragraph('9.6 사용자 관리 탭', style='CompleteH2')
    user_settings = doc.add_paragraph()
    user_settings.add_run('9.6.1 기능 개요\n').bold = True
    user_settings.add_run('시스템 사용자를 관리하는 탭입니다.\n\n')

    user_settings.add_run('9.6.2 주요 기능\n').bold = True
    user_settings.add_run('• 사용자 계정 상태 관리\n')
    user_settings.add_run('• 메뉴 권한 설정\n')
    user_settings.add_run('• 비밀번호 초기화\n\n')

    # 9.7 데이터 접근 권한 탭
    doc.add_paragraph('9.7 데이터 접근 권한 탭', style='CompleteH2')
    permission_settings = doc.add_paragraph()
    permission_settings.add_run('9.7.1 기능 개요\n').bold = True
    permission_settings.add_run('사용자별 데이터 접근 권한을 설정하는 탭입니다.\n\n')

    permission_settings.add_run('9.7.2 주요 기능\n').bold = True
    permission_settings.add_run('• Job ID별 권한 부여/해제\n')
    permission_settings.add_run('• 권한 템플릿 적용\n')
    permission_settings.add_run('• 권한 변경 이력 관리\n\n')

    # 9.8 통계 탭
    doc.add_paragraph('9.8 통계 탭', style='CompleteH2')
    stats_settings = doc.add_paragraph()
    stats_settings.add_run('9.8.1 기능 개요\n').bold = True
    stats_settings.add_run('시스템 사용 통계를 분석하는 탭입니다.\n\n')

    stats_settings.add_run('9.8.2 제공 통계\n').bold = True
    stats_settings.add_run('• 일별 메뉴 접근 통계\n')
    stats_settings.add_run('• 주별/월별 사용량 추이\n')
    stats_settings.add_run('• 사용자별 활동 분석\n\n')

    # 10. 수집 일정
    doc.add_paragraph('10. 수집 일정', style='CompleteH1')

    # 10.1 개요
    doc.add_paragraph('10.1 개요', style='CompleteH2')
    schedule_overview = doc.add_paragraph()
    schedule_overview.add_run('수집 일정 페이지는 데이터 수집 작업의 스케줄과 실행 현황을 모니터링하는 기능을 제공합니다. 예정된 작업과 실제 실행 결과를 비교하여 시스템의 안정성을 확인할 수 있습니다.\n\n')

    schedule_overview.add_run('주요 목적:\n').bold = True
    schedule_overview.add_run('• 스케줄 준수율 모니터링\n')
    schedule_overview.add_run('• 작업 실행 상태 실시간 확인\n')
    schedule_overview.add_run('• 장애 발생 지점 식별\n\n')

    # 10.2 화면 구성
    doc.add_paragraph('10.2 화면 구성', style='CompleteH2')
    schedule_composition = doc.add_paragraph()
    schedule_composition.add_run('10.2.1 뷰 모드\n').bold = True
    schedule_composition.add_run('• 주간 뷰: 7일간의 일정 표시\n')
    schedule_composition.add_run('• 월간 뷰: 한 달간의 일정 표시\n\n')

    schedule_composition.add_run('10.2.2 일정 테이블\n').bold = True
    schedule_composition.add_run('• Job ID 및 이름\n')
    schedule_composition.add_run('• 예정 실행 시간\n')
    schedule_composition.add_run('• 실제 실행 상태\n')
    schedule_composition.add_run('• 실행 결과 및 소요 시간\n\n')

    # 10.3 주요 기능
    doc.add_paragraph('10.3 주요 기능', style='CompleteH2')
    schedule_features = doc.add_paragraph()
    schedule_features.add_run('10.3.1 상태 모니터링\n').bold = True
    schedule_features.add_run('• 예정: 아직 실행 예정\n')
    schedule_features.add_run('• 수집중: 현재 실행 중\n')
    schedule_features.add_run('• 성공: 정상 완료\n')
    schedule_features.add_run('• 실패: 오류 발생\n')
    schedule_features.add_run('• 미수집: 예정 시간 초과\n\n')

    # 10.4 사용자 인터랙션
    doc.add_paragraph('10.4 사용자 인터랙션', style='CompleteH2')
    schedule_interaction = doc.add_paragraph()
    schedule_interaction.add_run('10.4.1 일정 모니터링 절차\n').bold = True
    schedule_interaction.add_run('1. 주간/월간 뷰 선택\n')
    schedule_interaction.add_run('2. 기간 설정\n')
    schedule_interaction.add_run('3. [조회] 버튼으로 데이터 로드\n')
    schedule_interaction.add_run('4. 상태별 색상으로 문제 식별\n\n')

    # 11. 카드 요약
    doc.add_paragraph('11. 카드 요약', style='CompleteH1')

    # 11.1 개요
    doc.add_paragraph('11.1 개요', style='CompleteH2')
    card_overview = doc.add_paragraph()
    card_overview.add_run('카드 요약 페이지는 오늘의 데이터 수집 현황을 그룹별로 요약하여 카드 형태로 표시하는 기능을 제공합니다. 각 카드는 Job ID의 첫 3자리로 그룹화되어 실시간으로 수집 상태를 집계하여 보여줍니다.\n\n')

    card_overview.add_run('주요 목적:\n').bold = True
    card_overview.add_run('• 그룹별 현황 파악\n')
    card_overview.add_run('• 실시간 상태 모니터링\n')
    card_overview.add_run('• 문제 발생 그룹 신속 식별\n\n')

    # 11.2 화면 구성
    doc.add_paragraph('11.2 화면 구성', style='CompleteH2')
    card_composition = doc.add_paragraph()
    card_composition.add_run('11.2.1 그룹 카드\n').bold = True
    card_composition.add_run('• CD100, CD200 등 그룹별 카드\n')
    card_composition.add_run('• 각 카드별 상태별 카운트\n')
    card_composition.add_run('• 실시간 업데이트\n\n')

    card_composition.add_run('11.2.2 표시 옵션\n').bold = True
    card_composition.add_run('• 명칭 표시: Job 이름 표시\n')
    card_composition.add_run('• 코드 표시: Job ID 표시\n')
    card_composition.add_run('• 명칭+코드: 둘 다 표시\n\n')

    # 11.3 주요 기능
    doc.add_paragraph('11.3 주요 기능', style='CompleteH2')
    card_features = doc.add_paragraph()
    card_features.add_run('11.3.1 실시간 집계\n').bold = True
    card_features.add_run('• 오늘의 모든 Job 상태 실시간 집계\n')
    card_features.add_run('• 5가지 상태별 분류 (성공/수집중/실패/미수집/예정)\n')
    card_features.add_run('• 권한 기반 데이터 필터링\n\n')

    # 11.4 사용자 인터랙션
    doc.add_paragraph('11.4 사용자 인터랙션', style='CompleteH2')
    card_interaction = doc.add_paragraph()
    card_interaction.add_run('11.4.1 카드 모니터링 절차\n').bold = True
    card_interaction.add_run('1. 표시 옵션 선택\n')
    card_interaction.add_run('2. 그룹별 카드 상태 파악\n')
    card_interaction.add_run('3. 문제 있는 카드 식별\n')
    card_interaction.add_run('4. 엑셀 양식 다운로드\n\n')

    # 12. 인증 시스템
    doc.add_paragraph('12. 인증 시스템', style='CompleteH1')

    # 12.1 로그인
    doc.add_paragraph('12.1 로그인', style='CompleteH2')
    login_desc = doc.add_paragraph()
    login_desc.add_run('12.1.1 기능 개요\n').bold = True
    login_desc.add_run('사용자 인증 및 시스템 접근을 위한 로그인 기능입니다.\n\n')

    login_desc.add_run('12.1.2 로그인 절차\n').bold = True
    login_desc.add_run('1. 사용자 ID 입력\n')
    login_desc.add_run('2. 비밀번호 입력\n')
    login_desc.add_run('3. [로그인] 버튼 클릭\n')
    login_desc.add_run('4. 인증 성공 시 대시보드 이동\n\n')

    # 12.2 회원가입
    doc.add_paragraph('12.2 회원가입', style='CompleteH2')
    register_desc = doc.add_paragraph()
    register_desc.add_run('12.2.1 기능 개요\n').bold = True
    register_desc.add_run('신규 사용자의 시스템 등록 기능입니다.\n\n')

    register_desc.add_run('12.2.2 회원가입 절차\n').bold = True
    register_desc.add_run('1. 사용자 ID 입력\n')
    register_desc.add_run('2. 비밀번호 및 확인 입력\n')
    register_desc.add_run('3. 비밀번호 정책 준수 확인\n')
    register_desc.add_run('4. [가입 신청] 버튼 클릭\n')
    register_desc.add_run('5. 관리자 승인 대기\n\n')

    # 12.3 비밀번호 변경
    doc.add_paragraph('12.3 비밀번호 변경', style='CompleteH2')
    pw_change_desc = doc.add_paragraph()
    pw_change_desc.add_run('12.3.1 기능 개요\n').bold = True
    pw_change_desc.add_run('로그인된 사용자의 비밀번호 변경 기능입니다.\n\n')

    pw_change_desc.add_run('12.3.2 변경 절차\n').bold = True
    pw_change_desc.add_run('1. 현재 비밀번호 입력\n')
    pw_change_desc.add_run('2. 새 비밀번호 입력\n')
    pw_change_desc.add_run('3. 새 비밀번호 확인 입력\n')
    pw_change_desc.add_run('4. [변경] 버튼 클릭\n\n')

    # 부록 A. 단축키 및 팁
    doc.add_paragraph('부록 A. 단축키 및 팁', style='CompleteH1')
    tips = doc.add_paragraph()
    tips.add_run('A.1 유용한 단축키\n').bold = True
    tips.add_run('• Ctrl+R: 페이지 새로고침\n')
    tips.add_run('• F12: 개발자 도구 열기\n')
    tips.add_run('• Ctrl+F: 페이지 내 검색\n\n')

    tips.add_run('A.2 사용 팁\n').bold = True
    tips.add_run('• 대시보드에서 자동 새로고침을 활용하여 실시간 모니터링\n')
    tips.add_run('• 데이터 분석 시 필터 조건을 저장하여 재사용\n')
    tips.add_run('• 차트 분석에서 범례를 클릭하여 특정 Job만 표시\n\n')

    try:
        doc.save('MSYS_Complete_Function_Manual.docx')
        print('전체 기능 매뉴얼 생성 완료: MSYS_Complete_Function_Manual.docx')
    except PermissionError:
        doc.save('MSYS_Complete_Function_Manual_v2.docx')
        print('전체 기능 매뉴얼 생성 완료: MSYS_Complete_Function_Manual_v2.docx')

def create_data_flow_debug_manual():
    """실제 소스 코드 기반 데이터 호출 흐름 디버깅 메뉴얼 생성"""
    doc = Document()

    # 스타일 설정
    title_style = doc.styles.add_style('DebugTitle', WD_STYLE_TYPE.PARAGRAPH)
    title_style.font.size = Pt(28)
    title_style.font.color.rgb = RGBColor(0, 0, 0)
    title_style.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    heading1_style = doc.styles.add_style('DebugH1', WD_STYLE_TYPE.PARAGRAPH)
    heading1_style.font.size = Pt(18)
    heading1_style.font.bold = True

    heading2_style = doc.styles.add_style('DebugH2', WD_STYLE_TYPE.PARAGRAPH)
    heading2_style.font.size = Pt(14)
    heading2_style.font.bold = True

    heading3_style = doc.styles.add_style('DebugH3', WD_STYLE_TYPE.PARAGRAPH)
    heading3_style.font.size = Pt(12)
    heading3_style.font.bold = True

    code_style = doc.styles.add_style('DebugCode', WD_STYLE_TYPE.PARAGRAPH)
    code_style.font.name = 'Courier New'
    code_style.font.size = Pt(10)

    # 표지
    title = doc.add_paragraph('MSYS 데이터 호출 흐름 디버깅 가이드', style='DebugTitle')
    subtitle = doc.add_paragraph('실제 소스 코드 기반 데이터 호출 추적 및 문제 해결 가이드', style='DebugTitle')
    version = doc.add_paragraph('버전 1.0', style='DebugTitle')
    version.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    doc.add_page_break()

    # 목차
    toc_title = doc.add_paragraph('목차', style='DebugH1')
    toc = doc.add_paragraph()
    toc.add_run('1. 개요 ........................ 3\n')
    toc.add_run('2. 메뉴 페이지별 데이터 호출 흐름 ................ 4\n')
    toc.add_run('   2.1 대시보드 페이지 .................. 4\n')
    toc.add_run('   2.2 데이터 분석 페이지 ................ 15\n')
    toc.add_run('   2.3 차트 분석 페이지 .................. 25\n')
    toc.add_run('   2.4 카드 요약 페이지 .................. 35\n')
    toc.add_run('   2.5 관리자 설정 페이지 ................ 45\n')
    toc.add_run('   2.6 인증 시스템 ................ 65\n')
    toc.add_run('3. 디버깅 가이드 ................ 105\n')
    toc.add_run('부록 A. 데이터 흐름 요약 ................ 110\n')

    doc.add_page_break()

    # 1. 개요
    doc.add_paragraph('1. 개요', style='DebugH1')
    overview = doc.add_paragraph()
    overview.add_run('1.1 목적\n').bold = True
    overview.add_run('이 가이드는 실제 MSYS 소스 코드를 분석하여 각 메뉴 페이지의 데이터 호출 과정을 추적하고 설명합니다. HTML 템플릿에서 어떤 변수를 사용하는지, JavaScript가 어떤 API를 호출하는지, 백엔드에서 어떻게 데이터를 처리하는지를 실제 코드 기반으로 상세히 설명합니다.\n\n')

    overview.add_run('1.2 대상 독자\n').bold = True
    overview.add_run('• 개발자: 데이터 호출 흐름 이해 및 디버깅\n')
    overview.add_run('• 운영자: 시스템 문제 해결 및 모니터링\n')
    overview.add_run('• 테스터: 데이터 흐름 검증 및 테스트 케이스 작성\n\n')

    overview.add_run('1.3 사용 방법\n').bold = True
    overview.add_run('1. 문제가 발생한 메뉴 페이지 확인\n')
    overview.add_run('2. 해당 섹션에서 HTML → JavaScript → API → Service → DAO → DB 순으로 추적\n')
    overview.add_run('3. 각 단계별 실제 코드와 권한 검증 확인\n')
    overview.add_run('4. 문제 지점 식별 및 해결\n\n')

    # 2. 메뉴 페이지별 데이터 호출 흐름
    doc.add_paragraph('2. 메뉴 페이지별 데이터 호출 흐름', style='DebugH1')

    # 2.1 대시보드 페이지
    doc.add_paragraph('2.1 대시보드 페이지', style='DebugH2')

    # 대시보드 데이터 호출 흐름 설명
    dashboard_desc = doc.add_paragraph()
    dashboard_desc.add_run('대시보드 페이지의 주요 데이터 호출 흐름:\n\n').bold = True

    # HTML 레벨
    dashboard_desc.add_run('HTML 레벨 (templates/dashboard.html):\n').bold = True
    dashboard_desc.add_run('• JavaScript 모듈 로드: /static/js/pages/dashboard.js\n')
    dashboard_desc.add_run('• 서버 사이드 변수: is_admin 권한 정보\n')
    dashboard_desc.add_run('• DOM 요소: totalJobsCount, totalCollectionsCount, daySuccessRate 등\n')
    dashboard_desc.add_run('• 날짜 범위 표시: minDateDisplay, maxDateDisplay, summary-min-date, summary-max-date, eventlog-min-date, eventlog-max-date\n')
    dashboard_desc.add_run('• 이벤트 로그: event-log-ul, eventLogPagination, eventLogStartDate, eventLogEndDate\n')
    dashboard_desc.add_run('• 차트 컨테이너: dashboardChartContainer\n')
    dashboard_desc.add_run('• 페이징 컨트롤: detailTablePagination\n\n')

    # JavaScript 레벨
    dashboard_desc.add_run('JavaScript 레벨:\n').bold = True
    dashboard_desc.add_run('• static/js/pages/dashboard.js: 진입점, 모듈 초기화\n')
    dashboard_desc.add_run('• static/js/modules/dashboard/events.js: 데이터 로드 및 UI 업데이트\n')
    dashboard_desc.add_run('• static/js/modules/dashboard/ui.js: UI 렌더링 및 업데이트\n')
    dashboard_desc.add_run('• static/js/modules/dashboard/eventLog.js: 이벤트 로그 관리\n')
    dashboard_desc.add_run('• static/js/modules/common/api/dashboard.js: API 호출 함수\n\n')

    # API 레벨
    dashboard_desc.add_run('API 레벨:\n').bold = True
    dashboard_desc.add_run('• /api/dashboard/summary: 대시보드 요약 데이터 API\n')
    dashboard_desc.add_run('• /api/dashboard/min-max-dates: 날짜 범위 조회 API\n')
    dashboard_desc.add_run('• /api/dashboard/event-log: 이벤트 로그 조회 API\n')
    dashboard_desc.add_run('• 권한 검증: @login_required, @check_password_change_required\n')
    dashboard_desc.add_run('• 파라미터: start_date, end_date, all_data\n\n')

    # Service 레벨
    dashboard_desc.add_run('Service 레벨 (service/dashboard_service.py):\n').bold = True
    dashboard_desc.add_run('• get_summary(): 메인 비즈니스 로직 메소드\n')
    dashboard_desc.add_run('• _fetch_manager_settings_with_icons(): 관리자 설정 로드\n')
    dashboard_desc.add_run('• _get_allowed_job_ids(): 권한 기반 Job ID 필터링\n')
    dashboard_desc.add_run('• _combine_historical_and_today_data(): 과거+오늘 데이터 결합\n\n')

    # DAO 레벨
    dashboard_desc.add_run('DAO 레벨 (mapper/dashboard_mapper.py):\n').bold = True
    dashboard_desc.add_run('• get_summary(): SQL 쿼리 실행 및 결과 변환\n')
    dashboard_desc.add_run('• DashboardSQL.get_dashboard_summary(): 동적 SQL 생성\n')
    dashboard_desc.add_run('• convert_to_new_columns(): 컬럼명 변환\n\n')

    # DB 레벨
    dashboard_desc.add_run('DB 레벨:\n').bold = True
    dashboard_desc.add_run('• tb_con_hist: 수집 이력 데이터\n')
    dashboard_desc.add_run('• tb_mngr_sett: 관리자 설정\n')
    dashboard_desc.add_run('• tb_con_hist_evnt_log: 이벤트 로그\n')
    dashboard_desc.add_run('• 복합 집계 쿼리: Job별 성공률, 상태별 카운트 등\n\n')

    # 대시보드 화면의 주요 표시 요소별 데이터 흐름 설명
    doc.add_paragraph('대시보드 화면의 주요 표시 요소별 데이터 흐름 설명:', style='DebugH3')

    # 대시보드 스크린샷 삽입
    try:
        doc.add_picture('scrennshot/dashboard_screenshot.PNG', width=Inches(6))
        doc.add_paragraph('그림: 대시보드 화면 스크린샷', style='DebugH3')
    except:
        doc.add_paragraph('[스크린샷: dashboard_screenshot.PNG를 찾을 수 없습니다]', style='DebugH3')

    doc.add_paragraph('', style='DebugH3')  # 빈 줄

    # 1. "총 Job ID 개수" 표시 요소
    doc.add_paragraph('1. "총 Job ID 개수" 표시 요소 (파란색 큰 숫자)', style='DebugH3')

    intro1 = doc.add_paragraph()
    intro1.add_run('화면 좌측 상단에 "총 Job ID 개수"라고 표시되는 값은 시스템에 등록된 데이터 수집 작업의 총 개수를 나타냅니다.\n\n')

    intro1.add_run('이 값의 계산 과정:\n').bold = True
    intro1.add_run('1. 데이터베이스에서 모든 Job의 기본 정보를 조회합니다.\n')
    intro1.add_run('2. 사용자의 접근 권한에 따라 허용된 Job만 필터링합니다.\n')
    intro1.add_run('3. 필터링된 Job 목록의 개수를 계산합니다.\n')
    intro1.add_run('4. 계산된 개수를 화면에 표시합니다.\n\n')

    # 단계별 코드 설명
    doc.add_paragraph('단계별 코드 흐름:', style='DebugH3')

    # DB 단계
    db_step = doc.add_paragraph(style='DebugCode')
    db_step.add_run('1. DB 단계 - mapper/dashboard_mapper.py의 get_summary() 함수 호출\n')
    db_step.add_run('SELECT DISTINCT job_id FROM tb_con_hist WHERE [조건들...]  -- results 배열 생성\n')
    db_step.add_run('results = [{"job_id": "CD101"}, {"job_id": "CD102"}, ...]  # 각 딕셔너리에 job_id 키를 가진 데이터 구조\n\n')

    # Service 단계
    service_step = doc.add_paragraph(style='DebugCode')
    service_step.add_run('2. Service 단계 - service/dashboard_service.py의 _get_allowed_job_ids() 함수 호출\n')
    service_step.add_run('allowed_job_ids = self._get_allowed_job_ids(user)  # ["CD101", "CD102"] 형태의 Job ID 목록 반환\n')
    service_step.add_run('filtered_data = [item for item in raw_data if item[\'job_id\'] in allowed_job_ids]  # 권한 필터링\n')
    service_step.add_run('job_count = len(filtered_data)  # filtered_data 리스트의 길이로 Job 개수 계산\n\n')

    # API 단계
    api_step = doc.add_paragraph(style='DebugCode')
    api_step.add_run('3. API 단계 - routes/api/dashboard_api.py의 get_dashboard_summary() 함수\n')
    api_step.add_run('summary_data = dashboard_service.get_summary(...)  # Service에서 계산된 데이터 수신\n')
    api_step.add_run('return jsonify({\'total_jobs\': job_count, \'job_list\': filtered_data})  # JSON 객체 형태로 반환\n\n')

    # JavaScript 단계
    js_step = doc.add_paragraph(style='DebugCode')
    js_step.add_run('4. JavaScript 단계 - static/js/modules/dashboard/ui.js의 updateSummaryCards() 함수\n')
    js_step.add_run('const totalJobsCount = summaryData.length;  # summaryData 배열의 길이\n')
    js_step.add_run('totalJobsElement.textContent = totalJobsCount;  // HTML 요소에 숫자 값 표시\n\n')

    # 2. "총 호출 건수" 표시 요소
    doc.add_paragraph('2. "총 호출 건수" 표시 요소 (보라색 큰 숫자)', style='DebugH3')

    intro2 = doc.add_paragraph()
    intro2.add_run('화면 좌측 중간에 "총 호출 건수"라고 표시되는 값은 선택된 기간 동안의 총 데이터 수집 시도 건수를 나타냅니다.\n\n')

    intro2.add_run('이 값의 계산 과정:\n').bold = True
    intro2.add_run('1. 각 Job별로 성공, 실패, 미수집 건수를 조회합니다.\n')
    intro2.add_run('2. 모든 Job의 건수를 합산합니다.\n')
    intro2.add_run('3. 천 단위 구분자(쉼표)를 추가하여 표시합니다.\n\n')

    # 단계별 코드 설명
    doc.add_paragraph('단계별 코드 흐름:', style='DebugH3')

    # DB 단계
    db_step2 = doc.add_paragraph(style='DebugCode')
    db_step2.add_run('1. DB 단계 - 각 Job별 통계 조회\n')
    db_step2.add_run('SELECT job_id,\n')
    db_step2.add_run('       SUM(CASE WHEN status = \'CD901\' THEN 1 ELSE 0 END) as success_count,\n')
    db_step2.add_run('       SUM(CASE WHEN status = \'CD902\' THEN 1 ELSE 0 END) as fail_count,\n')
    db_step2.add_run('       SUM(CASE WHEN status = \'CD903\' THEN 1 ELSE 0 END) as no_data_count\n')
    db_step2.add_run('FROM tb_con_hist GROUP BY job_id  -- 각 Job별 통계\n\n')

    # Service 단계
    service_step2 = doc.add_paragraph(style='DebugCode')
    service_step2.add_run('2. Service 단계 - 총합 계산\n')
    service_step2.add_run('total_collections = 0\n')
    service_step2.add_run('for job in job_list:\n')
    service_step2.add_run('    total_collections += job.success_count + job.fail_count + job.no_data_count\n\n')

    # JavaScript 단계
    js_step2 = doc.add_paragraph(style='DebugCode')
    js_step2.add_run('3. JavaScript 단계 - 천 단위 구분 및 표시\n')
    js_step2.add_run('function formatNumberWithKoreanUnits(num) {\n')
    js_step2.add_run('    return num.toLocaleString(\'ko-KR\');  // 1,234,567 형태로 변환\n')
    js_step2.add_run('}\n')
    js_step2.add_run('totalCollectionsElement.textContent = formatNumberWithKoreanUnits(total_collections);\n\n')

    # 3. "일간 성공률" 표시 요소
    doc.add_paragraph('3. "일간 성공률" 표시 요소 (하단 기간별 현황의 첫 번째)', style='DebugH3')

    intro3 = doc.add_paragraph()
    intro3.add_run('화면 하단의 기간별 현황에서 "일간"이라고 표시되는 성공률은 오늘 하루 동안의 데이터 수집 성공률을 나타냅니다.\n\n')

    intro3.add_run('이 값의 계산 과정:\n').bold = True
    intro3.add_run('1. 오늘 예정된 모든 수집 작업을 확인합니다.\n')
    intro3.add_run('2. 오늘 성공한 작업 수를 계산합니다.\n')
    intro3.add_run('3. (성공 건수 ÷ 예정 건수) × 100으로 성공률을 계산합니다.\n\n')

    # 단계별 코드 설명
    doc.add_paragraph('단계별 코드 흐름:', style='DebugH3')

    # Service 단계
    service_step3 = doc.add_paragraph(style='DebugCode')
    service_step3.add_run('1. Service 단계 - 오늘 데이터 집계\n')
    service_step3.add_run('today_scheduled = get_today_schedule_count()  # 오늘 예정된 작업 수\n')
    service_step3.add_run('today_success = get_today_success_count()  # 오늘 성공한 작업 수\n')
    service_step3.add_run('day_success_rate = (today_success / today_scheduled * 100) if today_scheduled > 0 else 0\n\n')

    # JavaScript 단계
    js_step3 = doc.add_paragraph(style='DebugCode')
    js_step3.add_run('2. JavaScript 단계 - 백분율 표시\n')
    js_step3.add_run('daySuccessRateElement.textContent = day_success_rate.toFixed(2) + \'%\';\n\n')

    # 4. Job별 상세 현황 테이블
    doc.add_paragraph('4. Job별 상세 현황 테이블 (화면 중앙의 표)', style='DebugH3')

    intro4 = doc.add_paragraph()
    intro4.add_run('화면 중앙의 큰 테이블은 각 Job별 상세 현황을 보여줍니다. Job ID, 데이터명, 주기, 성공률 등의 정보를 포함합니다.\n\n')

    intro4.add_run('이 테이블의 데이터 흐름:\n').bold = True
    intro4.add_run('1. DB에서 각 Job의 통계 정보를 조회합니다.\n')
    intro4.add_run('2. 권한에 따라 접근 가능한 Job만 필터링합니다.\n')
    intro4.add_run('3. 페이징을 적용하여 현재 페이지의 데이터만 표시합니다.\n')
    intro4.add_run('4. HTML 테이블 형태로 렌더링합니다.\n\n')

    # 단계별 코드 설명
    doc.add_paragraph('단계별 코드 흐름:', style='DebugH3')

    # DB 단계
    db_step4 = doc.add_paragraph(style='DebugCode')
    db_step4.add_run('1. DB 단계 - Job별 통계 조회\n')
    db_step4.add_run('SELECT job_id, cd_nm, frequency,\n')
    db_step4.add_run('       (SUM(CASE WHEN status = \'CD901\' THEN 1 ELSE 0 END) * 100.0 / COUNT(*)) as success_rate\n')
    db_step4.add_run('FROM tb_con_hist h JOIN tb_con_mst m ON h.job_id = m.cd\n')
    db_step4.add_run('WHERE [기간 조건] AND job_id IN (권한 있는 Job 목록)\n')
    db_step4.add_run('GROUP BY job_id, cd_nm, frequency\n')
    db_step4.add_run('ORDER BY job_id  -- job_stats 배열 생성\n\n')

    # JavaScript 단계
    js_step4 = doc.add_paragraph(style='DebugCode')
    js_step4.add_run('2. JavaScript 단계 - 테이블 렌더링\n')
    js_step4.add_run('function renderDashboardSummaryTable(jobStats) {\n')
    js_step4.add_run('    const tbody = document.getElementById(\'dashboardTableBody\');\n')
    js_step4.add_run('    tbody.innerHTML = jobStats.map(job => `\n')
    js_step4.add_run('        <tr>\n')
    js_step4.add_run('            <td>${job.job_id}</td>\n')
    js_step4.add_run('            <td>${job.cd_nm}</td>\n')
    js_step4.add_run('            <td>${job.frequency}</td>\n')
    js_step4.add_run('            <td>${job.success_rate.toFixed(2)}%</td>\n')
    js_step4.add_run('        </tr>\n')
    js_step4.add_run('    `).join(\'\');\n')
    js_step4.add_run('}\n\n')

    doc.add_paragraph('※ 위 설명은 대시보드 화면의 주요 표시 요소들을 중심으로 한 것입니다. 각 값들은 실시간으로 업데이트되며, 사용자의 권한에 따라 표시되는 데이터가 달라집니다.', style='DebugH3')

    # 실제 코드 예제 추가
    doc.add_paragraph('실제 API 호출 및 데이터 처리 코드:', style='DebugH3')

    # JavaScript API 호출 코드
    js_api = doc.add_paragraph(style='DebugCode')
    js_api.add_run('// static/js/modules/common/api/dashboard.js\n')
    js_api.add_run('export async function fetchDashboardSummary(startDate, endDate, allData) {\n')
    js_api.add_run('    const params = new URLSearchParams({\n')
    js_api.add_run('        start_date: startDate,\n')
    js_api.add_run('        end_date: endDate,\n')
    js_api.add_run('        all_data: allData\n')
    js_api.add_run('    });\n')
    js_api.add_run('    const url = `${BASE_URL}/api/dashboard/summary?${params.toString()}`;\n')
    js_api.add_run('    const response = await fetch(url);\n')
    js_api.add_run('    if (!response.ok) {\n')
    js_api.add_run('        const errorText = await response.text();\n')
    js_api.add_run('        throw new Error(`HTTP error! status: ${response.status}`);\n')
    js_api.add_run('    }\n')
    js_api.add_run('    const data = await response.json();\n')
    js_api.add_run('    return data;\n')
    js_api.add_run('}\n\n')

    # API Route 코드
    api_route = doc.add_paragraph(style='DebugCode')
    api_route.add_run('# routes/api/dashboard_api.py\n')
    api_route.add_run('@dashboard_api_bp.route(\'/summary\', methods=[\'GET\'])\n')
    api_route.add_run('@login_required\n')
    api_route.add_run('@check_password_change_required\n')
    api_route.add_run('def get_dashboard_summary():\n')
    api_route.add_run('    try:\n')
    api_route.add_run('        start_date_str = request.args.get(\'start_date\')\n')
    api_route.add_run('        end_date_str = request.args.get(\'end_date\')\n')
    api_route.add_run('        all_data_str = request.args.get(\'all_data\', \'false\')\n')
    api_route.add_run('        all_data = all_data_str.lower() == \'true\'\n')
    api_route.add_run('        \n')
    api_route.add_run('        user = session.get(\'user\')\n')
    api_route.add_run('        summary_data = dashboard_service.get_summary(\n')
    api_route.add_run('            start_date_str, end_date_str, all_data, user=user)\n')
    api_route.add_run('        \n')
    api_route.add_run('        return jsonify(summary_data), 200\n')
    api_route.add_run('    except Exception as e:\n')
    api_route.add_run('        return jsonify({"message": "데이터 조회 실패"}), 500\n\n')

    # Service 메소드 코드
    service_method = doc.add_paragraph(style='DebugCode')
    service_method.add_run('# service/dashboard_service.py\n')
    service_method.add_run('def get_summary(self, start_date, end_date, all_data, user):\n')
    service_method.add_run('    # 1. 권한 검증\n')
    service_method.add_run('    allowed_job_ids = self._get_allowed_job_ids(user)\n')
    service_method.add_run('    if allowed_job_ids is not None and not allowed_job_ids:\n')
    service_method.add_run('        return []  # 권한 없음\n')
    service_method.add_run('    \n')
    service_method.add_run('    # 2. 데이터 조회\n')
    service_method.add_run('    historical_data = self.dashboard_mapper.get_summary(\n')
    service_method.add_run('        start_date, end_date, all_data, allowed_job_ids)\n')
    service_method.add_run('    \n')
    service_method.add_run('    # 3. 오늘 데이터 결합\n')
    service_method.add_run('    today_data = self._get_today_schedule_data()\n')
    service_method.add_run('    \n')
    service_method.add_run('    # 4. 최종 데이터 가공\n')
    service_method.add_run('    return self._process_summary_data(historical_data, today_data)\n\n')

    # DAO 메소드 코드
    dao_method = doc.add_paragraph(style='DebugCode')
    dao_method.add_run('# mapper/dashboard_mapper.py\n')
    dao_method.add_run('def get_summary(self, start_date, end_date, all_data, job_ids):\n')
    dao_method.add_run('    query, params = DashboardSQL.get_dashboard_summary(\n')
    dao_method.add_run('        start_date, end_date, all_data, job_ids)\n')
    dao_method.add_run('    \n')
    dao_method.add_run('    with self.conn.cursor() as cur:\n')
    dao_method.add_run('        cur.execute(query, params)\n')
    dao_method.add_run('        results = [dict(zip([desc[0] for desc in cur.description], row)) \n')
    dao_method.add_run('                  for row in cur.fetchall()]\n')
    dao_method.add_run('    \n')
    dao_method.add_run('    return convert_to_new_columns(\'TB_CON_HIST\', results)\n\n')

    # 2.2 데이터 분석 페이지
    doc.add_paragraph('2.2 데이터 분석 페이지', style='DebugH2')

    # 데이터 분석 데이터 호출 흐름 설명
    analysis_desc = doc.add_paragraph()
    analysis_desc.add_run('데이터 분석 페이지의 주요 데이터 호출 흐름:\n\n').bold = True

    # HTML 레벨
    analysis_desc.add_run('HTML 레벨 (templates/data_analysis.html):\n').bold = True
    analysis_desc.add_run('• 탭 메뉴: 요약 데이터, 추이 데이터, 원천 데이터\n')
    analysis_desc.add_run('• 필터 컨트롤: 기간, Job ID, 상태 선택\n')
    analysis_desc.add_run('• JavaScript 모듈: /static/js/pages/data_analysis.js\n\n')

    # JavaScript 레벨
    analysis_desc.add_run('JavaScript 레벨:\n').bold = True
    analysis_desc.add_run('• static/js/modules/data_analysis/data.js: 데이터 조회 함수들\n')
    analysis_desc.add_run('• fetchSummaryData(), fetchTrendData(), fetchRawData()\n')
    analysis_desc.add_run('• API 호출 및 결과 처리\n\n')

    # API 레벨
    analysis_desc.add_run('API 레벨 (routes/api/analysis_api.py):\n').bold = True
    analysis_desc.add_run('• /api/analysis/summary: 요약 통계 데이터\n')
    analysis_desc.add_run('• /api/analysis/trend: 추이 차트 데이터\n')
    analysis_desc.add_run('• /api/analysis/raw: 원천 데이터\n\n')

    # Service/DAO 레벨
    analysis_desc.add_run('Service/DAO 레벨:\n').bold = True
    analysis_desc.add_run('• service/analysis_service.py: 분석 비즈니스 로직\n')
    analysis_desc.add_run('• dao/analytics_dao.py: 분석 데이터 조회\n')
    analysis_desc.add_run('• 복합 쿼리: 기간별 집계, 필터링 등\n\n')

    # 2.3 차트 분석 페이지
    doc.add_paragraph('2.3 차트 분석 페이지', style='DebugH2')

    # 차트 분석 데이터 호출 흐름 설명
    chart_desc = doc.add_paragraph()
    chart_desc.add_run('차트 분석 페이지의 주요 데이터 호출 흐름:\n\n').bold = True

    # HTML 레벨
    chart_desc.add_run('HTML 레벨 (templates/chart_analysis.html):\n').bold = True
    chart_desc.add_run('• Chart.js 기반 시각화 컴포넌트\n')
    chart_desc.add_run('• 성공률 추이 라인 차트\n')
    chart_desc.add_run('• 장애 코드 현황 파이 차트\n\n')

    # JavaScript 레벨
    chart_desc.add_run('JavaScript 레벨:\n').bold = True
    chart_desc.add_run('• static/js/modules/chart_analysis/chart.js\n')
    chart_desc.add_run('• Chart.js 라이브러리 연동\n')
    chart_desc.add_run('• 인터랙티브 차트 기능\n\n')

    # API/Service 레벨
    chart_desc.add_run('API/Service 레벨:\n').bold = True
    chart_desc.add_run('• /api/chart/trend: 성공률 추이 데이터\n')
    chart_desc.add_run('• /api/chart/errors: 장애 코드 현황 데이터\n')
    chart_desc.add_run('• 시계열 데이터 처리 및 집계\n\n')

    # 2.4 카드 요약 페이지
    doc.add_paragraph('2.4 카드 요약 페이지', style='DebugH2')

    # 카드 요약 데이터 호출 흐름 설명
    card_desc = doc.add_paragraph()
    card_desc.add_run('카드 요약 페이지의 주요 데이터 호출 흐름:\n\n').bold = True

    # HTML 레벨
    card_desc.add_run('HTML 레벨 (templates/card_summary.html):\n').bold = True
    card_desc.add_run('• 그룹별 카드 레이아웃 (CD100, CD200 등)\n')
    card_desc.add_run('• 실시간 상태 표시 (성공/수집중/실패/미수집/예정)\n')
    card_desc.add_run('• 표시 옵션 (명칭/코드 선택)\n\n')

    # JavaScript 레벨
    card_desc.add_run('JavaScript 레벨:\n').bold = True
    card_desc.add_run('• static/js/pages/card_summary.js\n')
    card_desc.add_run('• 실시간 데이터 업데이트\n')
    card_desc.add_run('• 그룹별 데이터 집계\n\n')

    # API/Service 레벨
    card_desc.add_run('API/Service 레벨:\n').bold = True
    card_desc.add_run('• collection_schedule_service: 오늘의 스케줄 데이터\n')
    card_desc.add_run('• 실시간 상태 계산 및 표시\n\n')

    # 2.5 관리자 설정 페이지
    doc.add_paragraph('2.5 관리자 설정 페이지', style='DebugH2')

    # 관리자 설정 데이터 호출 흐름 설명
    admin_desc = doc.add_paragraph()
    admin_desc.add_run('관리자 설정 페이지의 주요 데이터 호출 흐름:\n\n').bold = True

    # HTML 레벨
    admin_desc.add_run('HTML 레벨 (templates/mngr_sett.html):\n').bold = True
    admin_desc.add_run('• 탭 메뉴: 기본 설정, 수집 스케줄, 아이콘, 차트, 사용자, 권한, 통계\n')
    admin_desc.add_run('• 동적 폼 생성 및 설정 값 표시\n')
    admin_desc.add_run('• 실시간 설정 저장 및 적용\n\n')

    # JavaScript 레벨
    admin_desc.add_run('JavaScript 레벨:\n').bold = True
    admin_desc.add_run('• static/js/pages/mngr_sett.js\n')
    admin_desc.add_run('• 각 탭별 데이터 로드 및 저장\n')
    admin_desc.add_run('• 설정 값 검증 및 적용\n\n')

    # API/Service 레벨
    admin_desc.add_run('API/Service 레벨:\n').bold = True
    admin_desc.add_run('• /api/admin/basic: 기본 설정 CRUD\n')
    admin_desc.add_run('• /api/admin/schedule: 수집 스케줄 설정\n')
    admin_desc.add_run('• /api/admin/users: 사용자 관리\n')
    admin_desc.add_run('• tb_mngr_sett 테이블 조작\n\n')

    # 2.6 인증 시스템
    doc.add_paragraph('2.6 인증 시스템', style='DebugH2')

    # 인증 시스템 데이터 호출 흐름 설명
    auth_desc = doc.add_paragraph()
    auth_desc.add_run('인증 시스템의 주요 데이터 호출 흐름:\n\n').bold = True

    # HTML 레벨
    auth_desc.add_run('HTML 레벨:\n').bold = True
    auth_desc.add_run('• templates/login.html: 로그인 폼\n')
    auth_desc.add_run('• templates/change_password.html: 비밀번호 변경 폼\n')
    auth_desc.add_run('• 세션 기반 사용자 정보 표시\n\n')

    # JavaScript 레벨
    auth_desc.add_run('JavaScript 레벨:\n').bold = True
    auth_desc.add_run('• static/js/pages/login.js: 로그인 처리\n')
    auth_desc.add_run('• static/js/pages/change_password.js: 비밀번호 변경\n')
    auth_desc.add_run('• 폼 검증 및 API 호출\n\n')

    # API/Service 레벨
    auth_desc.add_run('API/Service 레벨:\n').bold = True
    auth_desc.add_run('• /api/auth/login: 사용자 인증\n')
    auth_desc.add_run('• /api/auth/register: 사용자 등록\n')
    auth_desc.add_run('• /api/auth/change_password: 비밀번호 변경\n')
    auth_desc.add_run('• tb_user 테이블 조작\n\n')

    # 3. 디버깅 가이드
    doc.add_paragraph('3. 디버깅 가이드', style='DebugH1')

    # 3.1 데이터가 표시되지 않는 경우
    doc.add_paragraph('3.1 데이터가 표시되지 않는 경우', style='DebugH2')
    debug1 = doc.add_paragraph()
    debug1.add_run('1. 브라우저 개발자 도구 Network 탭에서 API 호출 상태 확인\n')
    debug1.add_run('2. 403/401 에러 시 권한 문제 검토\n')
    debug1.add_run('3. 로그 파일에서 권한 관련 메시지 확인\n')
    debug1.add_run('4. 해당 페이지의 HTML 템플릿과 JavaScript 코드 검증\n\n')

    # 3.2 권한 문제 해결
    doc.add_paragraph('3.2 권한 문제 해결', style='DebugH2')
    debug2 = doc.add_paragraph()
    debug2.add_run('1. 사용자 관리 메뉴에서 권한 설정 확인\n')
    debug2.add_run('2. 데이터 접근 권한 탭에서 Job ID 권한 확인\n')
    debug2.add_run('3. 로그에서 "권한 없음" 메시지 확인\n')
    debug2.add_run('4. 미들웨어 권한 체크 로직 검증\n\n')

    # 3.3 데이터 이상 시
    doc.add_paragraph('3.3 데이터 이상 시', style='DebugH2')
    debug3 = doc.add_paragraph()
    debug3.add_run('1. DAO 메소드의 SQL 쿼리 확인\n')
    debug3.add_run('2. DB에서 직접 쿼리 실행하여 결과 확인\n')
    debug3.add_run('3. Service 레벨의 데이터 가공 로직 검증\n')
    debug3.add_run('4. Route 레벨의 파라미터 처리 확인\n\n')

    # 부록 A. 데이터 흐름 요약
    doc.add_paragraph('부록 A. 데이터 흐름 요약', style='DebugH1')
    summary = doc.add_paragraph()
    summary.add_run('A.1 일반적인 데이터 호출 패턴\n').bold = True
    summary.add_run('HTML (템플릿) → JavaScript (API 호출) → Route (권한 체크) → Service (비즈니스 로직) → DAO (SQL 실행) → DB (데이터 저장소)\n\n')

    summary.add_run('A.2 권한 검증 포인트\n').bold = True
    summary.add_run('• Route 레벨: 메뉴 접근 권한\n')
    summary.add_run('• Service 레벨: 데이터 접근 권한 필터링\n')
    summary.add_run('• 미들웨어: 세션 및 계정 상태 검증\n\n')

    summary.add_run('A.3 에러 처리 패턴\n').bold = True
    summary.add_run('• 401: 로그인 필요\n')
    summary.add_run('• 403: 권한 부족\n')
    summary.add_run('• 500: 서버 내부 오류\n')
    summary.add_run('• 빈 배열: 데이터 접근 권한 없음\n\n')

    try:
        doc.save('MSYS_Data_Flow_Debug_Guide.docx')
        print('실제 소스 코드 기반 데이터 호출 흐름 디버깅 가이드 생성 완료: MSYS_Data_Flow_Debug_Guide.docx')
    except PermissionError:
        doc.save('MSYS_Data_Flow_Debug_Guide_v3.docx')
        print('실제 소스 코드 기반 데이터 호출 흐름 디버깅 가이드 생성 완료: MSYS_Data_Flow_Debug_Guide_v3.docx')


def process_subsection(doc, subsection, code_style):
    """MD 서브섹션을 DOCX로 변환 - MD 파일의 모든 코드 블록을 포함"""
    lines = subsection.strip().split('\n')

    # 코드 블록과 일반 텍스트 분리하여 처리
    in_code_block = False
    current_para = None
    code_language = ""

    for line in lines:
        if line.strip().startswith('```'):
            if in_code_block:
                # 코드 블록 끝
                in_code_block = False
                if current_para:
                    current_para.add_run('\n')
                code_language = ""
            else:
                # 코드 블록 시작
                in_code_block = True
                code_language = line.strip()[3:]  # ```javascript 등에서 언어 추출
                current_para = doc.add_paragraph(style='DebugCode')
                # 코드 블록 시작 표시
                current_para.add_run(f'```{code_language}\n')
        elif in_code_block:
            # 코드 내용
            if current_para:
                current_para.add_run(line + '\n')
        else:
            # 일반 텍스트
            if line.strip():
                if '**' in line:
                    # 볼드 텍스트 처리
                    parts = line.split('**')
                    para = doc.add_paragraph()
                    for i, part in enumerate(parts):
                        if i % 2 == 1:  # 홀수 인덱스는 볼드
                            para.add_run(part).bold = True
                        else:
                            para.add_run(part)
                elif line.startswith('#### ') or line.startswith('### ') or line.startswith('## ') or line.startswith('# '):
                    # 헤딩 처리
                    heading_text = line.strip()
                    if heading_text.startswith('#### '):
                        para = doc.add_paragraph(heading_text[5:], style='DebugH3')
                    elif heading_text.startswith('### '):
                        para = doc.add_paragraph(heading_text[4:], style='DebugH2')
                    elif heading_text.startswith('## '):
                        para = doc.add_paragraph(heading_text[3:], style='DebugH1')
                    else:
                        para = doc.add_paragraph(heading_text[2:], style='DebugTitle')
                elif line.startswith('1. ') or line.startswith('2. ') or line.startswith('3. ') or line.startswith('4. ') or line.startswith('5. ') or line.startswith('6. ') or line.startswith('7. '):
                    # 번호 매기기 항목
                    para = doc.add_paragraph()
                    para.add_run(line)
                elif line.startswith('• ') or line.startswith('- '):
                    # 글머리 기호
                    para = doc.add_paragraph()
                    para.add_run(line)
                elif line.strip() == '---':
                    # 구분선
                    doc.add_paragraph('---')
                elif line.strip().startswith('**') and line.strip().endswith('**'):
                    # 전체 볼드 라인
                    para = doc.add_paragraph()
                    para.add_run(line.strip()[2:-2]).bold = True
                else:
                    # 일반 텍스트
                    para = doc.add_paragraph()
                    para.add_run(line)


def process_debug_guide(doc, debug_content, code_style):
    """디버깅 가이드 섹션 처리"""
    lines = debug_content.split('\n')

    for line in lines:
        if line.strip().startswith('### 1. 데이터가 표시되지 않는 경우'):
            doc.add_paragraph('3.1 데이터가 표시되지 않는 경우', style='DebugH2')
        elif line.strip().startswith('### 2. 권한 문제 해결'):
            doc.add_paragraph('3.2 권한 문제 해결', style='DebugH2')
        elif line.strip().startswith('### 3. 데이터 이상 시'):
            doc.add_paragraph('3.3 데이터 이상 시', style='DebugH2')
        elif line.strip():
            para = doc.add_paragraph()
            para.add_run(line)

if __name__ == '__main__':
    create_installation_manual()
    create_function_manual()
    create_operation_manual()
    create_database_manual()
    create_cpu_monitor_manual()
    create_data_flow_debug_guide()
    create_data_flow_debug_manual()
    print('모든 메뉴얼 생성이 완료되었습니다.')
